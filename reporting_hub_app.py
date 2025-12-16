from __future__ import annotations

import io
import re
import os
import base64
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Optional, List, Dict

import numpy as np
import pandas as pd
import streamlit as st

# bcrypt is optional – if missing, we disable login
try:
    import bcrypt  # type: ignore[import]
except ModuleNotFoundError:
    bcrypt = None

# ----- PDF / ReportLab imports -----
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, landscape
from reportlab.pdfbase import pdfmetrics
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import Table, TableStyle
from reportlab.lib.utils import ImageReader


# ========================================
# Config
# ========================================

APP_DIR = Path(__file__).resolve().parent
DEFAULT_PARQUET_DIR = APP_DIR / "parquet_db"

st.set_page_config(page_title="📊 Reporting Hub", layout="wide")
st.title("📊 Reporting Hub — Consolidated Maintenance & Cost Report")


# ============== Auth & access helpers ==============

def get_app_config():
    """Return app_config from st.secrets, or empty dict if missing."""
    try:
        return st.secrets["app_config"]
    except Exception:
        return {}


APP_CONFIG = get_app_config()


def require_login():
    """
    Simple username/password login using bcrypt hashes from app_config.

    If bcrypt is not available, login is **disabled** and user is auto-logged in
    as 'local_dev'.
    """
    # If bcrypt is missing, skip login entirely (dev mode)
    if bcrypt is None:
        if "user" not in st.session_state:
            st.session_state["user"] = {
                "username": "local_dev",
                "name": "Local Dev",
            }
        return

    # Already logged in?
    if "user" in st.session_state:
        return

    creds = APP_CONFIG.get("credentials", {}).get("usernames", {})
    if not creds:
        st.error("Auth configuration missing. Please contact the app admin.")
        st.stop()

    with st.sidebar.form("login_form", clear_on_submit=False):
        st.markdown("### Login")
        username = st.text_input("Username")
        password = st.text_input("Password", type="password")
        submitted = st.form_submit_button("Sign in")

    if submitted:
        user_cfg = creds.get(username)
        if user_cfg is None:
            st.sidebar.error("Invalid username or password.")
        else:
            stored_hash = user_cfg["password"]
            try:
                ok = bcrypt.checkpw(password.encode("utf-8"), stored_hash.encode("utf-8"))
            except Exception:
                ok = False
            if ok:
                st.session_state["user"] = {
                    "username": username,
                    "name": user_cfg.get("name", username),
                }
                st.rerun()
            else:
                st.sidebar.error("Invalid username or password.")

    if "user" not in st.session_state:
        st.stop()


# Force login (or auto-dev) before showing any app content
require_login()

st.sidebar.markdown(
    f"**User:** {st.session_state['user']['name']} "
    f"(`{st.session_state['user']['username']}`)"
)



# --- Shared PARQUET_DIR ---
PARQUET_DIR = Path(
    st.session_state.get("parquet_dir", DEFAULT_PARQUET_DIR)
)
globals()["PARQUET_DIR"] = PARQUET_DIR  # keep as Path, not str

st.sidebar.markdown("### Data Source")
st.sidebar.write(f"Parquet folder: `{PARQUET_DIR}`")

DATA_FILES: Dict[str, str] = {
    "costs_trends": "Workorders.parquet",
    "transactions": "TRANSACTIONS-FILTER.parquet",
    "parts": "Parts_Master.parquet",      # for inventory analysis
    "expected": "Expected.parquet",       # for Expected Service tab
}

# ---------------------------------------------------
# Optional Users.parquet for assignee dropdowns (WO tab)
# ---------------------------------------------------
_parq_dir = Path(globals().get("PARQUET_DIR", Path.cwd() / "parquet_db"))

USERS_PATH: Optional[Path] = None
for _p in [
    _parq_dir / "Users.parquet",
    Path(__file__).resolve().parent / "Users.parquet",
]:
    if _p.is_file():
        USERS_PATH = _p
        break


# ========================================
# Helpers
# ========================================

@st.cache_data(show_spinner=False)
def load_parquet(path: Path, columns=None) -> pd.DataFrame:
    if not path.exists():
        return pd.DataFrame()
    return pd.read_parquet(path, columns=columns)


def _collect_parent_locations_from_locdb(parquet_dir: Path) -> pd.Series:
    """
    Pull parent locations ONLY from Loc_DB.parquet, column 'All Parents'
    (or close variants), and ignore everything else.
    """
    import pandas as pd

    path = parquet_dir / "Loc_DB.parquet"
    if not path.exists():
        # Fallback: empty, so the app still runs
        return pd.Series([], dtype=str)

    try:
        df_loc = pd.read_parquet(path)
    except Exception:
        return pd.Series([], dtype=str)

    # Try a few name variants, but prefer 'All Parents'
    candidates = ["All Parents", "All Parent Locations", "All_Parents"]
    col = None
    for c in candidates:
        if c in df_loc.columns:
            col = c
            break

    if not col:
        # If the column isn't there, return empty
        return pd.Series([], dtype=str)

    return (
        df_loc[col]
        .dropna()
        .astype(str)
        .drop_duplicates()
        .sort_values()
    )


def _find(df: pd.DataFrame, *cands: str) -> Optional[str]:
    """
    Case-insensitive, space/underscore-insensitive column finder.

    Looks for explicit candidate names only (no fuzzy 'status' searching).
    Returns the actual column name from df.columns, or None if not found.
    """
    if df is None or df.empty:
        return None

    # map: lowercased column name -> real column name
    lower_map = {str(c).strip().lower(): c for c in df.columns}

    # 1) Try exact lowercase matches first
    for cand in cands:
        if not cand:
            continue
        key = cand.strip().lower()
        if key in lower_map:
            return lower_map[key]

    # 2) Try ignoring spaces / underscores
    norm_map = {
        str(c).strip().lower().replace(" ", "").replace("_", ""): c
        for c in df.columns
    }
    for cand in cands:
        if not cand:
            continue
        key = cand.strip().lower().replace(" ", "").replace("_", "")
        if key in norm_map:
            return norm_map[key]

    return None


def _get_last_parquet_update(parquet_dir: Path):
    """
    Return the most recent modification datetime of any .parquet file
    in parquet_dir. If none found or error, return None.
    """
    try:
        latest_ts = None
        for p in parquet_dir.glob("*.parquet"):
            if not p.is_file():
                continue
            ts = p.stat().st_mtime
            if latest_ts is None or ts > latest_ts:
                latest_ts = ts
        if latest_ts is None:
            return None
        return datetime.fromtimestamp(latest_ts)
    except Exception:
        return None


# Show approximate last update time based on parquet file mtimes
_last_upd = _get_last_parquet_update(PARQUET_DIR)
if _last_upd:
    st.sidebar.caption(f"🕒 DB last updated (approx): {_last_upd:%Y-%m-%d %H:%M}")
else:
    st.sidebar.caption("🕒 DB last updated: n/a")

    
# =========================
# PDF KPI HELPERS
# =========================

# ----- PDF KPI HELPERS -----
def _get_inv_kpis_for_pdf():
    """
    Pull inventory KPI state from session (populated by _render_inventory_analysis).

    Returns a 4-tuple:
        (inv_kpis_dict, monthly_rollup_df, location_label, period_mode_str)
    Any missing pieces fall back to safe defaults.
    """
    import pandas as pd
    import streamlit as st

    # first try the newer consolidated state, then older names (for backward compat)
    state = getattr(st.session_state, "inv_kpi_state", None)
    if isinstance(state, dict):
        kpis = state.get("kpis", {})
        monthly = state.get("monthly_rollup", pd.DataFrame())
        loc_label = state.get("location_label", "« All »")
        period_mode = state.get("period_mode", "YTD → month")
        return (kpis, monthly, loc_label, period_mode)

    # fallback to older individual keys if they exist
    kpis = getattr(st.session_state, "hub_inv_kpis", {})
    monthly = getattr(st.session_state, "hub_inv_monthly_rollup", pd.DataFrame())
    loc_label = getattr(st.session_state, "hub_inv_location", "« All »")
    period_mode = getattr(st.session_state, "hub_inv_period_mode", "YTD → month")

    # ensure types
    if not isinstance(kpis, dict):
        kpis = {}
    if not isinstance(monthly, pd.DataFrame):
        monthly = pd.DataFrame()

    return (kpis, monthly, loc_label, period_mode)


    return inv_kpis, inv_monthly, inv_loc, inv_period_mode


def _compute_inv_kpis_for_pdf(filtered_dfs, start_date, end_date, locations):
    """
    Backwards-compat shim.

    Old PDF code called `_compute_inv_kpis_for_pdf(...)`.
    Newer code reads KPIs from `st.session_state` via `_get_inv_kpis_for_pdf`.
    """
    return _get_inv_kpis_for_pdf()



def _compute_wo_kpis_for_pdf(
    start_date: date,
    end_date: date,
    locations: Optional[List[str]] = None,
) -> dict:
    """Compute WO KPI aggregates for the PDF summary page."""
    from pathlib import Path
    import pandas as pd
    import streamlit as st

    # Locate parquet dir
    parq_dir = Path(
        globals().get("PARQUET_DIR", st.session_state.get("parquet_dir", ""))
        or (Path.cwd() / "parquet_db")
    )

    # Try common WO parquet filenames
    wo_candidates = [
        parq_dir / "WO_Trans_non_DB.parquet",
        parq_dir / "WO-Trans_non_DB.parquet",
        parq_dir / "WO-Trans_no_DB.parquet",
    ]
    wo_path = None
    for p in wo_candidates:
        if p.is_file():
            wo_path = p
            break

    empty_result = {
        "completed_count": 0,
        "overdue_count": 0,
        "overdue_avg_days": np.nan,
        "com_overdue_count": 0,
        "com_overdue_avg_days": np.nan,
        "open_count": 0,
        "open_avg_days": np.nan,
        "coming_due_count": 0,
    }

    if wo_path is None:
        return empty_result

    df = pd.read_parquet(wo_path)
    if df.empty:
        return empty_result

    # Use the GLOBAL `_find` helper defined near the top of the file
    loc_col              = _find(df, "Location", "location", "Location2")
    comp_col             = _find(df, "COMPLETED ON", "Completed On", "Completed on")
    wo_col               = _find(df, "WORKORDER", "Workorder", "workorder")
    is_overdue_col       = _find(df, "IsOverDue", "IsOverdue")
    days_overdue_col     = _find(df, "DaysOverDue")
    com_overdue_col      = _find(df, "ComOverDue")
    days_com_overdue_col = _find(df, "DaysOverDue_Com")
    is_open_col          = _find(df, "IsOpen")
    days_open_col        = _find(df, "DaysOpen")

    # Location filter (use the `locations` argument, not selected_locations)
    if loc_col and locations:
        allowed = {str(x).strip() for x in locations}
        df = df[df[loc_col].astype(str).str.strip().isin(allowed)]

    if df.empty:
        return empty_result

    if comp_col:
        df[comp_col] = pd.to_datetime(df[comp_col], errors="coerce")

    # Helper: truthy mask
    def _bool_mask(series: pd.Series | None) -> pd.Series:
        if series is None:
            return pd.Series(False, index=df.index)
        s = series
        if pd.api.types.is_bool_dtype(s):
            return s.fillna(False)
        num = pd.to_numeric(s, errors="coerce")
        if num.notna().any():
            return num.fillna(0).astype(int).astype(bool)
        sval = s.astype(str).str.strip().str.lower()
        return sval.isin({"true", "t", "yes", "y", "1"})

    # Completed in date range: unique WORKORDER with Completed date in [start,end]
    if comp_col and wo_col and start_date and end_date:
        mask_comp = df[comp_col].dt.date.between(start_date, end_date)
        completed_count = (
            df.loc[mask_comp, wo_col]
            .dropna()
            .astype(str)
            .nunique()
        )
    else:
        completed_count = 0

    # Overdue
    m_overdue = _bool_mask(df[is_overdue_col] if is_overdue_col else None)
    if days_overdue_col:
        days_over = pd.to_numeric(df.loc[m_overdue, days_overdue_col], errors="coerce")
        overdue_avg = float(days_over.mean()) if not days_over.empty else np.nan
    else:
        overdue_avg = np.nan
    overdue_count = int(m_overdue.sum())

    # Completed Overdue
    m_com_over = _bool_mask(df[com_overdue_col] if com_overdue_col else None)
    if days_com_overdue_col:
        days_com = pd.to_numeric(df.loc[m_com_over, days_com_overdue_col], errors="coerce")
        com_overdue_avg = float(days_com.mean()) if not days_com.empty else np.nan
    else:
        com_overdue_avg = np.nan
    com_overdue_count = int(m_com_over.sum())

    # Open
    m_open = _bool_mask(df[is_open_col] if is_open_col else None)
    if days_open_col:
        days_open = pd.to_numeric(df.loc[m_open, days_open_col], errors="coerce")
        open_avg = float(days_open.mean()) if not days_open.empty else np.nan
    else:
        open_avg = np.nan
    open_count = int(m_open.sum())

    # Coming due: DaysOpen < 0 (any rows, not just open ones)
    coming_due_count = 0
    if days_open_col:
        days_open_all = pd.to_numeric(df[days_open_col], errors="coerce")
        coming_due_count = int((days_open_all < 0).sum())

    return {
        "completed_count": int(completed_count),
        "overdue_count": overdue_count,
        "overdue_avg_days": overdue_avg,
        "com_overdue_count": com_overdue_count,
        "com_overdue_avg_days": com_overdue_avg,
        "open_count": open_count,
        "open_avg_days": open_avg,
        "coming_due_count": coming_due_count,
    }
    
def _compute_wo_tables_for_pdf(
    start_date: date,
    end_date: date,
    selected_locations: list[str] | None = None,
) -> dict[str, pd.DataFrame]:
    """
    Build detail tables for the WO KPI section in the PDF.

    Returns a dict with DataFrames:
      - 'completed'    : WOs completed in the date window
      - 'overdue'      : IsOverDue = True
      - 'com_overdue'  : ComOverDue = True
      - 'open'         : IsOpen = True
      - 'coming_due'   : DaysOpen < 0
    """
    from pathlib import Path
    import pandas as pd
    import numpy as np

    # ---------- locate parquet ----------
    PARQUET_DIR = Path(globals().get("PARQUET_DIR", "")) or Path.cwd() / "parquet_db"
    candidates = [
        PARQUET_DIR / "WO-Trans_non_DB.parquet",
        PARQUET_DIR / "WO-Trans_no_DB.parquet",
        PARQUET_DIR / "WO-Trans.parquet",
    ]
    path = None
    for p in candidates:
        if p.is_file():
            path = p
            break
    if path is None:
        return {}

    df = pd.read_parquet(path)

    # ---------- column mapping ----------
    def _num(series):
        return pd.to_numeric(series, errors="coerce") if series is not None else None

    loc_col   = _find(df, "Location", "location", "Location2")
    wo_col    = _find(df, "WORKORDER", "WorkOrder", "WO")
    title_col = _find(df, "TITLE", "Title", "Subject")
    desc_col  = _find(df, "DESCRIPTION", "Description")
    po_col    = _find(df, "PO", "PO Number", "PO_NUMBER", "PO #")
    pn_col    = _find(df, "P/N", "Part", "Item", "Part Number", "Item Name")
    comp_col  = _find(df, "Completed On", "completed_on", "CompletedOn")
    due_col   = _find(df, "DueDate", "Due Date")
    is_over_col  = _find(df, "IsOverDue", "IsOverdue")
    days_over_col = _find(df, "DaysOverDue", "DaysOverdue")
    com_over_col  = _find(df, "ComOverDue", "CompletedOverDue", "Completed Overdue")
    days_com_col  = _find(df, "DaysOverDue_Com", "DaysOverDueCom")
    is_open_col   = _find(df, "IsOpen")
    days_open_col = _find(df, "DaysOpen")

    assignee_col  = _find(df, "Assignee", "Assigned To", "Assigned_to")
    team_col      = _find(df, "Team Assigned", "team_assigned")

    # ---------- location filter ----------
    if loc_col and selected_locations:
        df = df[df[loc_col].astype(str).isin(map(str, selected_locations))]

    # ---------- coerce dates / numbers ----------
    if comp_col:
        df[comp_col] = pd.to_datetime(df[comp_col], errors="coerce")
    if due_col:
        df[due_col] = pd.to_datetime(df[due_col], errors="coerce")

    days_over = _num(df[days_over_col]) if days_over_col else None
    days_com  = _num(df[days_com_col]) if days_com_col else None
    days_open = _num(df[days_open_col]) if days_open_col else None

    # ---------- helper: truthy mask ----------
    def _bool_mask(series: pd.Series | None) -> pd.Series:
        if series is None:
            return pd.Series(False, index=df.index)
        s = series
        if pd.api.types.is_bool_dtype(s):
            return s.fillna(False)
        num = pd.to_numeric(s, errors="coerce")
        if num.notna().any():
            return num.fillna(0).astype(int).astype(bool)
        sval = s.astype(str).str.strip().str.lower()
        return sval.isin({"true", "t", "yes", "y", "1"})

    # ---------- date window for "completed" ----------
    if comp_col and start_date and end_date:
        m_comp = df[comp_col].between(
            pd.Timestamp(start_date),
            pd.Timestamp(end_date) + pd.Timedelta(days=1),
            inclusive="left",
        )
    else:
        m_comp = pd.Series(False, index=df.index)

    # ---------- masks ----------
    m_completed    = m_comp
    m_overdue      = _bool_mask(df[is_over_col] if is_over_col else None)
    m_com_overdue  = _bool_mask(df[com_over_col] if com_over_col else None)
    m_open         = _bool_mask(df[is_open_col] if is_open_col else None)
    m_coming_due   = days_open.lt(0) if days_open is not None else pd.Series(False, index=df.index)

    # ---------- build a slim display DF ----------
    disp_cols: list[tuple[str, str | None]] = [
        ("WorkOrder", wo_col),
        ("Title", title_col),
        ("PO", po_col),
        ("P/N", pn_col),
        ("Assignee", assignee_col),
        ("Team", team_col),
        ("DueDate", due_col),
        ("Completed On", comp_col),
        ("IsOverDue", is_over_col),
        ("DaysOverDue", days_over_col),
        ("ComOverDue", com_over_col),
        ("DaysOverDue_Com", days_com_col),
        ("IsOpen", is_open_col),
        ("DaysOpen", days_open_col),
    ]

    cols_final = [label for label, src in disp_cols if src is not None]
    if not cols_final:
        return {}

    df_slim = pd.DataFrame(index=df.index)
    for label, src in disp_cols:
        if src is not None:
            df_slim[label] = df[src]

    # format dates as date only for readability
    for dcol in ["DueDate", "Completed On"]:
        if dcol in df_slim.columns:
            s = pd.to_datetime(df_slim[dcol], errors="coerce")
            df_slim[dcol] = s.dt.date.where(~s.isna(), df_slim[dcol])

    # ---------- extract tables ----------
    tables: dict[str, pd.DataFrame] = {}

    def _tbl(mask, key):
        sub = df_slim.loc[mask, cols_final].copy()
        if not sub.empty:
            tables[key] = sub.reset_index(drop=True)

    _tbl(m_completed, "completed")
    _tbl(m_overdue, "overdue")
    _tbl(m_com_overdue, "com_overdue")
    _tbl(m_open, "open")
    _tbl(m_coming_due, "coming_due")

    return tables




def _compute_expected_kpis_for_pdf(df_expected: pd.DataFrame):
    """
    From Expected.parquet (already filtered by main page), count:
      - Overdue
      - Needs Service
      - New Reading Needed
    """
    import numpy as np
    import pandas as pd

    if df_expected is None or df_expected.empty:
        return {
            "overdue": 0,
            "needs_service": 0,
            "new_reading_needed": 0,
        }

    # Re-use Status column as in render_expected_service
    if "Status" not in df_expected.columns:
        return {
            "overdue": 0,
            "needs_service": 0,
            "new_reading_needed": 0,
        }

    status = (
        df_expected["Status"]
        .astype(str)
        .str.strip()
        .str.casefold()
    )

    return {
        "overdue": int(status.eq("overdue").sum()),
        "needs_service": int(status.eq("needs service").sum()),
        "new_reading_needed": int(status.eq("new reading needed").sum()),
    }



def _guess_date_range(all_dfs: Dict[str, pd.DataFrame]) -> Tuple[date, date]:
    """
    Guess a global date range across all datasets for the sidebar.
    """
    cols = (
        "Completed On", "Completed on",
        "Date",
        "Created On", "Created on",
        "Service date",
        "Trans Date",
    )
    dates = []
    for df in all_dfs.values():
        if df.empty:
            continue
        for c in cols:
            if c in df.columns:
                d = pd.to_datetime(df[c], errors="coerce")
                dates.append(d)
                break
    if not dates:
        today = date.today()
        return today.replace(day=1), today
    cat = pd.concat(dates).dropna()
    if cat.empty:
        today = date.today()
        return today.replace(day=1), today
    return cat.min().date(), cat.max().date()
    
    
@st.cache_data(show_spinner=True)
def load_data(path: Path | str) -> tuple[pd.DataFrame, Optional[str]]:
    """
    Load the WO-Trans parquet and normalize the Completed On date column.

    Returns:
      df       -> DataFrame
      date_col -> name of the 'Completed On' column (or None)
    """
    df = pd.read_parquet(path)

    # Find a completed-on style column WITHOUT using _find
    candidates = [
        "COMPLETED ON",
        "Completed On",
        "Completed on",
        "completed_on",
    ]

    date_col: Optional[str] = None
    # build a lowercase -> real name map
    lower_map = {str(c).strip().lower(): c for c in df.columns}

    for cand in candidates:
        key = cand.strip().lower()
        if key in lower_map:
            date_col = lower_map[key]
            break

    if date_col:
        df[date_col] = pd.to_datetime(df[date_col], errors="coerce")

    return df, date_col
    
    
if "load_users" not in globals():
    @st.cache_data(show_spinner=False)
    def load_users(path: Optional[Path] | str) -> pd.DataFrame:
        """Load Users.parquet if it exists; otherwise return empty DF."""
        if not path:
            return pd.DataFrame()
        p = Path(path)
        if not p.is_file():
            return pd.DataFrame()
        return pd.read_parquet(p)
        
        
if "_extract_unique_teams" not in globals():
    def _extract_unique_teams(series: pd.Series) -> List[str]:
        """
        Split multi-team strings (comma/semicolon) into unique trimmed names.
        """
        teams = set()
        for v in series.dropna():
            text = str(v).replace(";", ",")
            for part in text.split(","):
                t = part.strip()
                if t:
                    teams.add(t)
        return sorted(teams)

# ---------- Generic XLSX helper ----------
def _to_xlsx_bytes(df: pd.DataFrame, sheet_name: str = "Sheet1") -> bytes:
    """
    Convert a DataFrame to XLSX bytes for Streamlit download buttons.
    """
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False, sheet_name=sheet_name)
    output.seek(0)
    return output.getvalue()
    
    
# =========================
# PDF HELPERS
# =========================

def _df_to_rl_table(df: pd.DataFrame, max_rows: int = 25) -> Table:
    """
    Convert a DataFrame to a simple reportlab Table with basic styling.
    Truncates to `max_rows` data rows to keep the page readable.
    """
    if df is None or df.empty:
        data = [["(no rows)"]]
        tbl = Table(data)
        tbl.setStyle(
            TableStyle(
                [
                    ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                ]
            )
        )
        return tbl

    df_show = df.head(max_rows).copy()
    data = [list(map(str, df_show.columns.tolist()))] + df_show.astype(str).values.tolist()
    tbl = Table(data, repeatRows=1)

    tbl.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e5e7eb")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, 0), 9),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 1), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.gray),
                ("ALIGN", (0, 0), (-1, 0), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
        )
    )
    return tbl
    
# =========================
# PDF PAGES: Cover + KPI Summary
# =========================

def _pdf_cover_page(c, start_date: date | None, end_date: date | None, locations: list[str] | None):
    """Simple cover page for the Reporting Hub PDF."""
    from reportlab.lib.pagesizes import letter, landscape
    from reportlab.lib.units import inch

    width, height = letter

    c.setFont("Helvetica-Bold", 24)
    c.drawCentredString(width / 2.0, height - 1.5 * inch, "Reporting Hub Summary Report")

    c.setFont("Helvetica", 12)

    # Date range
    if start_date and end_date:
        date_text = f"Date Range: {start_date:%Y-%m-%d} to {end_date:%Y-%m-%d}"
    else:
        date_text = "Date Range: (not specified)"
    c.drawCentredString(width / 2.0, height - 2.0 * inch, date_text)

    # Locations
    if locations:
        loc_text = ", ".join(map(str, locations))
    else:
        loc_text = "All locations"
    c.drawCentredString(width / 2.0, height - 2.4 * inch, f"Locations: {loc_text}")

    c.setFont("Helvetica-Oblique", 10)
    c.drawCentredString(width / 2.0, 1.0 * inch, f"Generated {datetime.now():%Y-%m-%d %H:%M}")

    c.showPage()


def _pdf_kpi_page(
    c: canvas.Canvas,
    inv_kpis: dict,
    inv_monthly: pd.DataFrame,
    wo_kpis: dict,
    wo_tables: dict[str, pd.DataFrame],
    expected_kpis: dict,
    start_date: date,
    end_date: date,
    locations: List[str],
    title_left: str = "KPI Summary",
):
    """
    Draw KPI summary page:
      - Inventory KPIs + full Monthly Roll-up table
      - Work Order KPIs + their detail tables
      - Expected Service KPIs
    """
    page_w, page_h = landscape(letter)
    margin = 0.5 * inch
    x_left = margin
    x_right = page_w - margin
    y = page_h - margin

    def _ensure_space(min_height: float) -> float:
        nonlocal y
        if y - min_height < margin:
            c.showPage()
            y = page_h - margin
        return y

    # --------- TITLE + meta ---------
    c.setFont("Helvetica-Bold", 18)
    c.drawString(x_left, y, title_left)
    c.setFont("Helvetica", 9)
    loc_txt = ", ".join(map(str, locations)) if locations else "All locations"
    c.drawRightString(
        x_right,
        y,
        f"Period: {start_date:%Y-%m-%d} → {end_date:%Y-%m-%d}  •  {loc_txt}",
    )
    y -= 26

    # =====================================================================
    # 1) INVENTORY KPIs
    # =====================================================================
    c.setFont("Helvetica-Bold", 14)
    c.drawString(x_left, y, "Inventory KPIs")
    y -= 16

    # inv_kpis keys: ytd_growth_pct, avg_mom_pct, coverage_net_pct,
    #                coverage_abs_pct, gap_share_net_pct, avg_error_pct,
    #                reopen_pct_of_wo
    def _fmt_pct(v):
        return "" if v is None or (isinstance(v, float) and np.isnan(v)) else f"{v:.2f}%"

    # Two lines summary
    c.setFont("Helvetica", 9)
    line1 = (
        f"YTD / Period Growth: {_fmt_pct(inv_kpis.get('ytd_growth_pct'))}  •  "
        f"Avg MoM Δ: {_fmt_pct(inv_kpis.get('avg_mom_pct'))}"
    )
    line2 = (
        f"PO/WO Coverage (NET): {_fmt_pct(inv_kpis.get('coverage_net_pct'))}  •  "
        f"PO/WO Coverage (ABS): {_fmt_pct(inv_kpis.get('coverage_abs_pct'))}  •  "
        f"Adj/Unmatched (NET): {_fmt_pct(inv_kpis.get('gap_share_net_pct'))}  •  "
        f"Avg Error Rate: {_fmt_pct(inv_kpis.get('avg_error_pct'))}"
    )
    c.drawString(x_left, y, line1)
    y -= 12
    c.drawString(x_left, y, line2)
    y -= 18

    # --------- Inventory Monthly Roll-up (FULL TABLE) ----------
    if inv_monthly is not None and not inv_monthly.empty:
        _ensure_space(140)
        c.setFont("Helvetica-Bold", 11)
        c.drawString(x_left, y, "Inventory Monthly Roll-up")
        y -= 14

        # prepare a copy with formatted numbers
        tbl_df = inv_monthly.copy()

        money_cols = [
            col
            for col in tbl_df.columns
            if any(
                k in str(col)
                for k in [
                    "Inventory Value",
                    "In $",
                    "Out $",
                    "Net (In−Out)",
                    "trans_in $",
                    "trans_out $",
                    "Gap (Net IO − Net Trans)",
                ]
            )
        ]
        pct_cols = [
            col
            for col in tbl_df.columns
            if any(
                k in str(col)
                for k in [
                    "%Δ Inventory vs prior",
                    "% Impact on Inv",
                    "Error Rate vs Inv Value %",
                ]
            )
        ]

        for col in money_cols:
            if col in tbl_df.columns:
                tbl_df[col] = pd.to_numeric(tbl_df[col], errors="coerce").map(
                    lambda v: "" if pd.isna(v) else f"${v:,.0f}"
                )

        for col in pct_cols:
            if col in tbl_df.columns:
                tbl_df[col] = pd.to_numeric(tbl_df[col], errors="coerce").map(
                    lambda v: "" if pd.isna(v) else f"{v:.2f}%"
                )

        cols = list(tbl_df.columns)
        col_count = len(cols)
        if col_count > 0:
            max_table_width = page_w - 2 * margin
            col_width = max_table_width / col_count
            row_h = 10
            header_h = 12

            # header
            c.setFont("Helvetica-Bold", 7)
            for j, col in enumerate(cols):
                c.drawString(
                    x_left + j * col_width + 2,
                    y,
                    str(col)[:20],
                )
            y -= header_h
            c.setLineWidth(0.5)
            c.line(x_left, y + 2, x_left + col_count * col_width, y + 2)

            # rows
            c.setFont("Helvetica", 7)
            for _, row in tbl_df.iterrows():
                _ensure_space(row_h + 12)
                for j, col in enumerate(cols):
                    txt = str(row[col])
                    if len(txt) > 18:
                        txt = txt[:15] + "..."
                    c.drawString(
                        x_left + j * col_width + 2,
                        y,
                        txt,
                    )
                y -= row_h

            y -= 10  # spacing after table

    # =====================================================================
    # 2) WORK ORDER KPIs + DETAIL TABLES
    # =====================================================================
    _ensure_space(60)
    c.setFont("Helvetica-Bold", 14)
    c.drawString(x_left, y, "Work Order KPIs")
    y -= 16
    c.setFont("Helvetica", 9)

    # metrics summary
    c.drawString(
        x_left,
        y,
        f"Completed WOs in period: {wo_kpis.get('completed_count', 0):,}",
    )
    y -= 12
    c.drawString(
        x_left,
        y,
        f"Overdue WOs (IsOverDue=True): {wo_kpis.get('overdue_count', 0):,}  •  "
        f"Avg DaysOverDue: {wo_kpis.get('overdue_avg', float('nan')):.1f}"
        if not np.isnan(wo_kpis.get("overdue_avg", np.nan))
        else f"Overdue WOs (IsOverDue=True): {wo_kpis.get('overdue_count', 0):,}",
    )
    y -= 12
    c.drawString(
        x_left,
        y,
        f"Completed Overdue (ComOverDue=True): {wo_kpis.get('com_overdue_count', 0):,}  •  "
        f"Avg DaysOverDue_Com: {wo_kpis.get('com_overdue_avg', float('nan')):.1f}"
        if not np.isnan(wo_kpis.get("com_overdue_avg", np.nan))
        else f"Completed Overdue (ComOverDue=True): {wo_kpis.get('com_overdue_count', 0):,}",
    )
    y -= 12
    c.drawString(
        x_left,
        y,
        f"Open WOs (IsOpen=True): {wo_kpis.get('open_count', 0):,}  •  "
        f"Avg DaysOpen: {wo_kpis.get('open_avg', float('nan')):.1f}"
        if not np.isnan(wo_kpis.get("open_avg", np.nan))
        else f"Open WOs (IsOpen=True): {wo_kpis.get('open_count', 0):,}",
    )
    y -= 12
    c.drawString(
        x_left,
        y,
        f"Coming Due (DaysOpen < 0): {wo_kpis.get('coming_due_count', 0):,}",
    )
    y -= 18

    # shared small table renderer for WO detail
    def _draw_small_table(lbl: str, key: str):
        nonlocal y
        df_tbl = wo_tables.get(key)
        if df_tbl is None or df_tbl.empty:
            return
        _ensure_space(60)
        c.setFont("Helvetica-Bold", 10)
        c.drawString(x_left, y, lbl)
        y -= 12

        # limit to first 25 rows for readability
        df_view = df_tbl.head(25).copy()
        cols = list(df_view.columns)
        col_count = len(cols)
        max_width = page_w - 2 * margin
        col_width = max_width / col_count
        row_h = 9
        header_h = 11

        c.setFont("Helvetica-Bold", 6)
        for j, col in enumerate(cols):
            c.drawString(
                x_left + j * col_width + 2,
                y,
                str(col)[:18],
            )
        y -= header_h
        c.setLineWidth(0.4)
        c.line(x_left, y + 2, x_left + col_count * col_width, y + 2)

        c.setFont("Helvetica", 6)
        for _, row in df_view.iterrows():
            _ensure_space(row_h + 10)
            for j, col in enumerate(cols):
                txt = "" if pd.isna(row[col]) else str(row[col])
                if len(txt) > 18:
                    txt = txt[:15] + "..."
                c.drawString(
                    x_left + j * col_width + 2,
                    y,
                    txt,
                )
            y -= row_h

        y -= 8

    # draw each WO table under its KPI
    _draw_small_table("Overdue WOs (IsOverDue=True)", "overdue")
    _draw_small_table("Completed Overdue (ComOverDue=True)", "com_overdue")
    _draw_small_table("Open WOs (IsOpen=True)", "open")
    _draw_small_table("Coming Due (DaysOpen < 0)", "coming_due")
    _draw_small_table("Completed WOs in Period", "completed")

    # =====================================================================
    # 3) EXPECTED SERVICE KPIs (no tables yet, just counts)
    # =====================================================================
    _ensure_space(50)
    c.setFont("Helvetica-Bold", 14)
    c.drawString(x_left, y, "Expected Service KPIs")
    y -= 16
    c.setFont("Helvetica", 9)

    exp_total = expected_kpis.get("total", 0)
    exp_over  = expected_kpis.get("overdue_count", 0)
    exp_needs = expected_kpis.get("needs_service_count", 0)
    exp_newr  = expected_kpis.get("new_reading_needed_count", 0)
    exp_due   = expected_kpis.get("due_count", 0)

    c.drawString(
        x_left,
        y,
        f"Total Expected rows (current filter): {exp_total:,}",
    )
    y -= 12
    c.drawString(
        x_left,
        y,
        f"Overdue: {exp_over:,}  •  Needs Service: {exp_needs:,}  •  "
        f"New Reading Needed: {exp_newr:,}  •  Due: {exp_due:,}",
    )
    y -= 18


    c.showPage()



def _filter_by_locations(df: pd.DataFrame, selected_locations: Optional[List[str]]) -> pd.DataFrame:
    """
    Generic helper for Location/Location2 filtering using the same logic
    as the rest of the hub (uses _find if available).
    """
    if df is None or df.empty or not selected_locations:
        return df

    loc_col = None
    if "_find" in globals():
        loc_col = _find(df, "Location2", "Location", "Company", "Site")
    else:
        for c in ["Location2", "Location", "Company", "Site"]:
            if c in df.columns:
                loc_col = c
                break

    if not loc_col:
        return df

    loc_norm = df[loc_col].astype(str).str.strip()
    allowed = {str(s).strip() for s in selected_locations}
    mask = loc_norm.isin(allowed)
    return df[mask]


def _safe_num_series(s: pd.Series) -> pd.Series:
    return pd.to_numeric(s, errors="coerce").fillna(0.0)

from typing import Optional, List  # make sure Optional is imported at the top

from pathlib import Path
import pandas as pd
import streamlit as st

# make sure PARQUET_DIR already exists somewhere above;
# if not, this will create a sensible default:
PARQUET_DIR = Path(
    globals().get("PARQUET_DIR", st.session_state.get("parquet_dir", "")) 
    or (Path.cwd() / "parquet_db")
)

@st.cache_data
def _load_expected_for_pdf() -> pd.DataFrame:
    """Load Expected.parquet for the Reporting Hub PDF."""
    path = PARQUET_DIR / "Expected.parquet"
    if path.exists():
        return pd.read_parquet(path)
    return pd.DataFrame()

def build_reporting_hub_pdf(
    filtered_dfs,
    date_range,
    locations,
    inv_kpis=None,
    svc_kpis=None,
    df_expected: pd.DataFrame | None = None,
):
    """
    Build the Reporting Hub PDF in the following order:

      1. Cover page
      2. Index
      3. Costs & Trends — Location & YTD
      4. Inventory Analysis — KPIs & Monthly Rollup
      5. Transactions — Filtered Detail
      6–9. Work Orders — KPI/detail pages (4 sections)
      10–? Expected Services — All, Expected (coming due), Overdue, Needs Service, New Reading Needed
    """
    import math
    import calendar
    from datetime import datetime, date
    import textwrap as _tw
    import pandas as pd
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter, landscape
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import Table, TableStyle

    # ---------------------------
    # Normalize inputs
    # ---------------------------
    if date_range:
        start_date, end_date = date_range
    else:
        start_date = end_date = None

    locations = locations or []

    # Expected services DF (prefer filtered, fallback to loader)
    if df_expected is None:
        try:
            df_expected = filtered_dfs.get("expected", None)
        except Exception:
            df_expected = None
    if df_expected is None or getattr(df_expected, "empty", True):
        try:
            df_expected = _load_expected_for_pdf()
        except Exception:
            df_expected = pd.DataFrame()

    # Inventory KPI state from session
    inv_kpis_state, inv_monthly, inv_loc, inv_period_mode = _get_inv_kpis_for_pdf()
    if inv_kpis_state is None or not isinstance(inv_kpis_state, dict):
        inv_kpis_state = {}
    if inv_monthly is None or not isinstance(inv_monthly, pd.DataFrame):
        inv_monthly = pd.DataFrame()

    # ------------ Transactions base DF ------------
    df_tx = filtered_dfs.get("transactions", None)
    if df_tx is None:
        df_tx = pd.DataFrame()

    # Find date + location columns (case-insensitive)
    date_col_tx = _find(
        df_tx,
        "Completed On", "Completed on", "Date",
        "TransDate", "Created On",
    )
    loc_col_tx = _find(
        df_tx,
        "Location", "Location2", "NS Location", "loc",
    )

    # Start from the raw transactions that came into the PDF
    df_tx_scoped = df_tx.copy()

    # Apply LOCATION filter from main page
    if loc_col_tx and locations:
        loc_strs = {str(x).strip() for x in locations}
        df_tx_scoped = df_tx_scoped[
            df_tx_scoped[loc_col_tx].astype(str).str.strip().isin(loc_strs)
        ]

    # Apply DATE filter from main page (global date_range)
    if date_col_tx and start_date and end_date:
        df_tx_scoped[date_col_tx] = pd.to_datetime(df_tx_scoped[date_col_tx], errors="coerce")
        df_tx_scoped = df_tx_scoped[
            (df_tx_scoped[date_col_tx] >= pd.to_datetime(start_date)) &
            (df_tx_scoped[date_col_tx] <= pd.to_datetime(end_date))
        ]


    # Parts DF (for future use)
    df_parts = filtered_dfs.get("parts", None)
    if df_parts is None:
        df_parts = pd.DataFrame()
        
    # Work Orders DF for asset YTD summary
    df_wo = filtered_dfs.get("Workorders", None)
    if df_wo is None:
        df_wo = pd.DataFrame()


    # ---------------------------
    # Setup canvas + helpers
    # ---------------------------
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter  # initial page portrait

    def _title(text: str, y: float, size: int = 20):
        c.setFont("Helvetica-Bold", size)
        c.drawString(0.75 * inch, y, text)

    def _sub(text: str, y: float, size: int = 12):
        c.setFont("Helvetica", size)
        c.drawString(0.75 * inch, y, text)

    def _draw_wrapped_text(x: float, y_start: float, text: str, max_chars: int = 80, line_height: float = 0.18 * inch):
        """Simple text wrapper for long strings (like locations)."""
        if not text:
            return y_start
        lines = _tw.wrap(str(text), max_chars)
        y = y_start
        for line in lines:
            c.drawString(x, y, line)
            y -= line_height
        return y

    def _wrap_cell(val: object, max_chars: int) -> str:
        """Return a string with newline breaks so long values don't spill into adjacent columns."""
        s = "" if val is None else str(val)
        if len(s) <= max_chars:
            return s
        return _tw.fill(s, width=max_chars)

    def _draw_table(
        df: pd.DataFrame,
        x: float,
        y: float,
        max_width: float,
        max_height: float,
        font_size: int = 7,
    ):
        """
        Render a DataFrame as a compact table, truncating rows/cols if needed
        to fit the available space. This is a single-page draw; pagination is
        handled by _draw_table_paged.
        """
        if df is None or df.empty:
            return

        df2 = df.copy().astype(str)

        # keep all columns but cap at ~14 so it doesn't get unreadable
        max_cols = min(len(df2.columns), 14)
        if len(df2.columns) > max_cols:
            df2 = df2.iloc[:, :max_cols]

        # ---- ROW LIMIT (more conservative) ----
        # Bigger height per row to account for wrapped text and padding.
        approx_row_height = 0.30 * inch
        max_rows = int(max_height / approx_row_height)
        if max_rows < 3:
            max_rows = 3

        # reserve at least one row worth of space beyond our guess
        if len(df2) > max_rows - 1:
            df2 = df2.head(max_rows - 1)

        # estimate characters that fit into each column for wrapping
        if len(df2.columns) > 0:
            max_chars = max(8, int(80 / len(df2.columns)))
        else:
            max_chars = 20

        # wrap headers and cells so they don’t spill into next column
        wrapped_cols = [_wrap_cell(col, max_chars) for col in df2.columns]
        data = [wrapped_cols] + [
            [_wrap_cell(v, max_chars) for v in row]
            for row in df2.itertuples(index=False, name=None)
        ]

        col_width = max_width / (len(df2.columns) if len(df2.columns) else 1)
        col_widths = [col_width] * (len(df2.columns) if len(df2.columns) else 1)

        tbl = Table(data, colWidths=col_widths)
        tbl.setStyle(TableStyle([
            ("FONT", (0, 0), (-1, 0), "Helvetica-Bold", font_size),
            ("FONT", (0, 1), (-1, -1), "Helvetica", font_size),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f2937")),  # dark header
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
            ("ALIGN", (0, 0), (-1, -1), "LEFT"),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("GRID", (0, 0), (-1, -1), 0.35, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.white]),
        ]))
        tbl.wrapOn(c, max_width, max_height)
        tbl.drawOn(c, x, y - tbl._height)


    def _draw_table_paged(
        df: pd.DataFrame,
        x: float,
        top_y: float,
        max_width: float,
        bottom_margin: float = 0.7 * inch,
        font_size: int = 7,
        page_title: str | None = None,
    ):
        """
        Draw a potentially long table across multiple landscape pages.
        """
        nonlocal width, height

        if df is None or df.empty:
            return

        # Match the more conservative estimate in _draw_table
        approx_row_height = 0.30 * inch

        max_height = top_y - bottom_margin
        max_rows = int(max_height / approx_row_height)

        # leave extra buffer so we never clip at the bottom
        if max_rows < 5:
            rows_per_page = max_rows - 1
        else:
            rows_per_page = max_rows - 3

        if rows_per_page < 1:
            rows_per_page = 1

        start_idx = 0
        while start_idx < len(df):
            chunk = df.iloc[start_idx:start_idx + rows_per_page].copy()
            _draw_table(chunk, x, top_y, max_width, max_height, font_size=font_size)
            start_idx += rows_per_page
            if start_idx >= len(df):
                break

            # new page for remaining rows
            c.showPage()
            c.setPageSize(landscape(letter))
            width, height = landscape(letter)
            if page_title:
                _title(page_title + " (cont.)", height - 0.7 * inch, 16)
            top_y = height - 1.2 * inch
            max_height = top_y - bottom_margin




    def _draw_table(
        df: pd.DataFrame,
        x: float,
        y: float,
        max_width: float,
        max_height: float,
        font_size: int = 7,
        max_rows: int | None = None,
    ):

        """
        Render a DataFrame as a compact table, truncating rows/cols if needed
        to fit the available space. This is a single-page draw; pagination is
        handled by _draw_table_paged.
        """
        if df is None or df.empty:
            return

        df2 = df.copy().astype(str)

        # keep all columns but cap at ~14 so it doesn't get unreadable
        max_cols = min(len(df2.columns), 14)
        if len(df2.columns) > max_cols:
            df2 = df2.iloc[:, :max_cols]

        # limit rows so table doesn’t overflow vertically
        approx_row_height = 0.18 * inch
        if max_rows is None:
            max_rows = int(max_height / approx_row_height)
        if max_rows < 3:
            max_rows = 3
        if len(df2) > max_rows - 1:
            df2 = df2.head(max_rows - 1)

        # estimate characters that fit into each column for wrapping
        if len(df2.columns) > 0:
            max_chars = max(8, int(80 / len(df2.columns)))
        else:
            max_chars = 20

        # wrap headers and cells so they don’t spill into next column
        wrapped_cols = [_wrap_cell(col, max_chars) for col in df2.columns]
        data = [wrapped_cols] + [
            [_wrap_cell(v, max_chars) for v in row]
            for row in df2.itertuples(index=False, name=None)
        ]

        col_width = max_width / (len(df2.columns) if len(df2.columns) else 1)
        col_widths = [col_width] * (len(df2.columns) if len(df2.columns) else 1)

        tbl = Table(data, colWidths=col_widths)
        tbl.setStyle(TableStyle([
            ("FONT", (0, 0), (-1, 0), "Helvetica-Bold", font_size),
            ("FONT", (0, 1), (-1, -1), "Helvetica", font_size),
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f2937")),  # dark header
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
            ("ALIGN", (0, 0), (-1, -1), "LEFT"),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("GRID", (0, 0), (-1, -1), 0.35, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.white]),
        ]))
        tbl.wrapOn(c, max_width, max_height)
        tbl.drawOn(c, x, y - tbl._height)

    def _draw_table_paged(
        df: pd.DataFrame,
        x: float,
        top_y: float,
        max_width: float,
        bottom_margin: float = 0.7 * inch,
        font_size: int = 7,
        page_title: str | None = None,
    ):
        """
        Draw a potentially long table across multiple landscape pages.
        """
        nonlocal width, height

        if df is None or df.empty:
            return

        approx_row_height = 0.18 * inch
        max_height = top_y - bottom_margin
        max_rows = int(max_height / approx_row_height)

        # leave an extra row or two of space so we don't clip at the bottom
        if max_rows < 4:
            rows_per_page = max_rows - 1
        else:
            rows_per_page = max_rows - 2

        start_idx = 0
        while start_idx < len(df):
            chunk = df.iloc[start_idx:start_idx + rows_per_page].copy()
            _draw_table(
                chunk,
                x,
                top_y,
                max_width,
                max_height,
                font_size=font_size,
                max_rows=rows_per_page + 1,  # header + rows_per_page
            )

            start_idx += rows_per_page
            if start_idx >= len(df):
                break

            # new page for remaining rows
            c.showPage()
            c.setPageSize(landscape(letter))
            width, height = landscape(letter)
            if page_title:
                _title(page_title + " (cont.)", height - 0.7 * inch, 16)
            top_y = height - 1.2 * inch
            max_height = top_y - bottom_margin



    def _tx_ytd_pivot(df: pd.DataFrame, index_col: str, cost_col: str) -> pd.DataFrame:
        """
        Simple YTD-style pivot by index_col and month for the PDF.
        Assumes df is already filtered to the desired window.
        """
        if df is None or df.empty or index_col not in df.columns or cost_col not in df.columns:
            return pd.DataFrame()

        date_col = _find(df, "Completed On", "Completed on", "Date", "TransDate", "Created On")
        if not date_col:
            return pd.DataFrame()

        df2 = df[[index_col, cost_col, date_col]].copy()
        df2[date_col] = pd.to_datetime(df2[date_col], errors="coerce")
        df2 = df2.dropna(subset=[date_col])
        if df2.empty:
            return pd.DataFrame()

        df2["Mon"] = df2[date_col].dt.month
        g = df2.groupby([index_col, "Mon"], as_index=False)[cost_col].sum()

        df2["MonLbl"] = df2["Mon"].map(
            lambda m: calendar.month_abbr[int(m)] if pd.notna(m) else ""
        )
        g["MonLbl"] = g["Mon"].map(
            lambda m: calendar.month_abbr[int(m)] if pd.notna(m) else ""
        )
        pv = g.pivot(index=index_col, columns="MonLbl", values=cost_col).reset_index()

        mon_desc = [calendar.month_abbr[m] for m in range(12, 0, -1)]
        month_cols = [c for c in mon_desc if c in pv.columns]
        for c in month_cols:
            pv[c] = pd.to_numeric(pv[c], errors="coerce").fillna(0.0)

        if month_cols:
            pv["YTD Total"] = pv[month_cols].sum(axis=1, skipna=True)
        else:
            pv["YTD Total"] = 0.0

        ordered_cols = [index_col, "YTD Total"] + month_cols
        pv = pv[ordered_cols].sort_values("YTD Total", ascending=False)

        return pv

    def _fmt_currency(df: pd.DataFrame, skip_first: bool = True) -> pd.DataFrame:
        if df is None or df.empty:
            return df
        out = df.copy()
        start_idx = 1 if skip_first else 0
        for col in out.columns[start_idx:]:
            out[col] = pd.to_numeric(out[col], errors="coerce").map(
                lambda v: "" if pd.isna(v) else f"${float(v):,.2f}"
            )
        return out

    def _get_inv_kpi_val(keys: list[str]) -> float | None:
        for k in keys:
            if k in inv_kpis_state:
                try:
                    v = float(inv_kpis_state[k])
                    return v
                except Exception:
                    continue
        return None

    KPI_THRESHOLDS = {
        "ytd_growth": {"type": "band", "green_min": 0.0, "green_max": 10.0, "warn_min": -5.0, "warn_max": 20.0},
        "avg_mom": {"type": "max_good", "goal": 2.0, "warn_max": 4.0},
        "coverage": {"type": "min_good", "goal": 95.0, "warn_min": 90.0},
        "gap_share": {"type": "max_good", "goal": 15.0, "warn_max": 20.0},
        "avg_err": {"type": "max_good", "goal": 0.5, "warn_max": 1.0},
        "reopen": {"type": "max_good", "goal": 5.0, "warn_max": 10.0},
    }

    def _classify_kpi(metric_key: str, val: float | None) -> tuple[str, str]:
        """
        Return (color, arrow_shape) for a KPI.
        color in {green,yellow,red,grey}; shape in {up,down,flat}
        """
        if val is None or (isinstance(val, float) and math.isnan(val)):
            return ("grey", "flat")
        cfg = KPI_THRESHOLDS.get(metric_key)
        if not cfg:
            return ("grey", "flat")
        t = cfg["type"]
        if t == "band":
            gmin, gmax = cfg["green_min"], cfg["green_max"]
            wmin, wmax = cfg["warn_min"], cfg["warn_max"]
            if gmin <= val <= gmax:
                return ("green", "flat")
            if wmin <= val <= wmax:
                return ("yellow", "flat")
            return ("red", "flat")
        if t == "min_good":
            goal, wmin = cfg["goal"], cfg["warn_min"]
            if val >= goal:
                return ("green", "up")
            if val >= wmin:
                return ("yellow", "up")
            return ("red", "up")
        if t == "max_good":
            goal, wmax = cfg["goal"], cfg["warn_max"]
            if val <= goal:
                return ("green", "down")
            if val <= wmax:
                return ("yellow", "down")
            return ("red", "down")
        return ("grey", "flat")

    def _truthy_mask(series: pd.Series | None, index) -> pd.Series:
        """
        Return a boolean mask for numeric-ish 'in/out' logic.
        None → all False.
        """
        if series is None:
            return pd.Series(False, index=index)
        s = pd.to_numeric(series, errors="coerce")
        return (s.notna()) & (s != 0)

    # ---------------------------------------------------
    # 1) COVER PAGE
    # ---------------------------------------------------
    _title("Reporting Hub Summary", height - 1.0 * inch, 24)
    c.setFont("Helvetica", 12)

    if start_date and end_date:
        win_text = f"Reporting Window: {start_date:%Y-%m-%d} → {end_date:%Y-%m-%d}"
    else:
        win_text = "Reporting Window: (all available data)"

    c.drawString(0.75 * inch, height - 1.6 * inch, win_text)

    if locations:
        loc_text = ", ".join(str(x) for x in locations)
    else:
        loc_text = "All selected locations"

    _draw_wrapped_text(0.75 * inch, height - 1.9 * inch, f"Locations: {loc_text}", max_chars=90)

    c.drawString(0.75 * inch, height - 2.4 * inch, f"Generated: {datetime.now():%Y-%m-%d %H:%M}")
    c.showPage()

    # ---------------------------------------------------
    # 2) INDEX PAGE
    # ---------------------------------------------------
    c.setPageSize(letter)
    width, height = letter
    _title("Report Index", height - 1.0 * inch, 20)
    c.setFont("Helvetica", 11)
    y = height - 1.7 * inch

    index_lines = [
        "1. Cover Page",
        "2. Index (this page)",
        "3. Costs & Trends — Location & YTD",
        "4. Inventory Analysis — KPIs & Monthly Rollup",
        "5. Transactions — Filtered Detail",
        "6–9. Work Orders — KPIs & Detail (4 sections)",
        "10–14. Expected Services — All, Expected (Coming Due), Overdue, Needs Service, New Reading Needed",
    ]
    for line in index_lines:
        c.drawString(0.75 * inch, y, f"• {line}")
        y -= 0.28 * inch

    c.showPage()

    # ---------------------------------------------------
    # 3) COSTS & TRENDS PAGE — Workorders (Workorders.parquet) ONLY
    # ---------------------------------------------------
    c.setPageSize(landscape(letter))
    width, height = landscape(letter)
    _title("Costs & Trends — Summary", height - 0.7 * inch, 18)
    c.setFont("Helvetica", 10)

    # Base DF from filtered_dfs
    df_wo = filtered_dfs.get("workorders", pd.DataFrame()).copy()

    # last resort: session cache (only if present)
    if (df_wo is None) or df_wo.empty:
        try:
            import streamlit as st
            df_wo = st.session_state.get("rhub_costs_source_rows", pd.DataFrame()).copy()
        except Exception:
            df_wo = pd.DataFrame()

    # If still empty, render n/a and keep going (so the PDF doesn't crash)
    if df_wo is None or df_wo.empty:
        c.drawString(0.75 * inch, height - 1.3 * inch, "Window total: n/a (no Workorders data)")
        c.showPage()

    else:
        # --- Find key columns in Workorders ---
        loc_col  = _find(df_wo, "Location", "Location2", "NS Location", "location")
        date_col = _find(
            df_wo,
            "Completed On", "Completed on", "COMPLETED ON",
            "Created On", "Created on", "Date", "Service date"
        )

        item_col = _find(df_wo, "TOTAL ITEM COST", "Total Item Cost", "total item cost", "TotalItemCost")
        totl_col = _find(df_wo, "Total cost", "Total Cost", "TOTAL COST", "total cost", "TotalCost")
        cost_fallback_col = _find(df_wo, "_Cost", "_COST", "__cost", "Cost")

        # Build true cost:
        # Prefer TOTAL ITEM COST + Total cost, else fallback to _Cost
        if item_col and totl_col:
            df_wo["__pdf_cost"] = (
                pd.to_numeric(df_wo[item_col], errors="coerce").fillna(0.0)
                + pd.to_numeric(df_wo[totl_col], errors="coerce").fillna(0.0)
            ).astype(float)
        elif cost_fallback_col:
            df_wo["__pdf_cost"] = pd.to_numeric(df_wo[cost_fallback_col], errors="coerce").fillna(0.0).astype(float)
        else:
            df_wo["__pdf_cost"] = 0.0

        # --- Apply location + date filters (window) ---
        df_win = df_wo.copy()

        if loc_col and locations:
            allowed_locs = {str(x).strip() for x in locations}
            df_win = df_win[df_win[loc_col].astype(str).str.strip().isin(allowed_locs)]

        if date_col and date_col in df_win.columns:
            df_win[date_col] = pd.to_datetime(df_win[date_col], errors="coerce")
            if start_date and end_date:
                df_win = df_win[
                    (df_win[date_col] >= pd.to_datetime(start_date)) &
                    (df_win[date_col] <= pd.to_datetime(end_date))
                ]

        total_cost_window = float(df_win["__pdf_cost"].fillna(0.0).sum()) if not df_win.empty else 0.0

        # --- last month total within the current window ---
        last_month_total = 0.0
        if (end_date is not None) and (date_col is not None) and (not df_win.empty):
            m = (
                (df_win[date_col].dt.year == end_date.year) &
                (df_win[date_col].dt.month == end_date.month)
            )
            last_month_total = float(df_win.loc[m, "__pdf_cost"].fillna(0.0).sum())

        # Draw window totals
        if total_cost_window:
            label_total = "YTD total (window)" if (
                start_date and end_date and
                start_date.year == end_date.year and
                start_date.month == 1 and
                start_date.day == 1
            ) else "Window total"
            c.drawString(0.75 * inch, height - 1.3 * inch, f"{label_total}: ${total_cost_window:,.2f}")
        else:
            c.drawString(0.75 * inch, height - 1.3 * inch, "Window total: n/a")

        if last_month_total:
            c.drawString(0.75 * inch, height - 1.6 * inch, f"Last month in window: ${last_month_total:,.2f}")

        # --- Build YTD frame (Jan 1 -> end_date) for same locations ---
        df_ytd = df_wo.copy()

        if loc_col and locations:
            allowed_locs = {str(x).strip() for x in locations}
            df_ytd = df_ytd[df_ytd[loc_col].astype(str).str.strip().isin(allowed_locs)]

        if date_col and end_date:
            df_ytd[date_col] = pd.to_datetime(df_ytd[date_col], errors="coerce")
            year_start = datetime(end_date.year, 1, 1)
            df_ytd = df_ytd[
                (df_ytd[date_col] >= year_start) &
                (df_ytd[date_col] <= pd.to_datetime(end_date))
            ]

        # --- YTD Summary by Location (from Workorders only) ---
        ytd_loc = pd.DataFrame()
        if (df_ytd is not None) and (not df_ytd.empty) and loc_col:
            tmp = df_ytd[[loc_col, "__pdf_cost"]].copy()
            tmp["__Month"] = pd.to_datetime(df_ytd[date_col], errors="coerce").dt.month if date_col else np.nan

            p = tmp.pivot_table(
                index=loc_col,
                columns="__Month",
                values="__pdf_cost",
                aggfunc="sum",
                fill_value=0.0,
            )

            for mm in range(1, 13):
                if mm not in p.columns:
                    p[mm] = 0.0
            p = p[[mm for mm in range(1, 13)]]

            p["YTD Total"] = p.sum(axis=1)
            p = p.reset_index()

            month_names = {
                1: "Jan", 2: "Feb", 3: "Mar", 4: "Apr", 5: "May", 6: "Jun",
                7: "Jul", 8: "Aug", 9: "Sep", 10: "Oct", 11: "Nov", 12: "Dec"
            }
            p = p.rename(columns=month_names)
            ytd_loc = p

        # Reorder columns to: Location, YTD Total, current month, previous month, ...
        if ytd_loc is not None and not ytd_loc.empty and ("YTD Total" in ytd_loc.columns):
            import calendar as _cal
            cur_m = int(end_date.month) if end_date else int(date.today().month)
            mon_order = [_cal.month_abbr[m] for m in range(cur_m, 0, -1)] + \
                        [_cal.month_abbr[m] for m in range(12, cur_m, -1)]

            base_cols = [loc_col, "YTD Total"]
            month_cols = [m for m in mon_order if m in ytd_loc.columns]
            remaining = [c for c in ytd_loc.columns if c not in (base_cols + month_cols)]
            ytd_loc = ytd_loc[base_cols + month_cols + remaining]

            y = pd.to_numeric(ytd_loc["YTD Total"], errors="coerce").fillna(0.0)
            ytd_loc = ytd_loc.loc[y.sort_values(ascending=False).index].reset_index(drop=True)

        if ytd_loc is not None and not ytd_loc.empty:
            ytd_loc = _fmt_currency(ytd_loc, skip_first=True)

        y_top = height - 2.1 * inch
        _sub("YTD Summary by Location", y_top, 12)
        y_loc = y_top - 0.3 * inch
        bottom_margin = 0.7 * inch
        _draw_table(
            ytd_loc,
            x=0.75 * inch,
            y=y_loc,
            max_width=width - 1.5 * inch,
            max_height=y_loc - bottom_margin,
            font_size=6,
        )

        c.showPage()

    # ---------------------------------------------------
    # 3b) YTD SUMMARY BY ASSET — OWN PAGE (Workorders only)
    #     Columns: Asset, YTD Total, current month, prev month, ...
    # ---------------------------------------------------
    c.setPageSize(landscape(letter))
    width, height = landscape(letter)
    _title("YTD Summary by Asset", height - 0.7 * inch, 18)
    c.setFont("Helvetica", 10)

    import calendar as _cal
    import numpy as np
    import pandas as pd

    # Reuse df_ytd + __pdf_cost + date_col from section (3)
    asset_col = _find(df_ytd, "Asset", "Asset Name", "ASSET", "Name", "Equipment", "Equipment #")

    df_asset = pd.DataFrame()

    if (df_ytd is not None) and (not df_ytd.empty) and asset_col:
        tmpa = df_ytd[[asset_col, "__pdf_cost"]].copy()
        tmpa["__Month"] = pd.to_datetime(df_ytd[date_col], errors="coerce").dt.month if date_col else np.nan

        p = tmpa.pivot_table(
            index=asset_col,
            columns="__Month",
            values="__pdf_cost",
            aggfunc="sum",
            fill_value=0.0,
        )

        for mm in range(1, 13):
            if mm not in p.columns:
                p[mm] = 0.0
        p = p[[mm for mm in range(1, 13)]]

        p["YTD Total"] = p.sum(axis=1)
        p = p.reset_index()

        month_names = {m: _cal.month_abbr[m] for m in range(1, 13)}
        df_asset = p.rename(columns=month_names)

    if df_asset is not None and not df_asset.empty and ("YTD Total" in df_asset.columns):
        cur_m = int(end_date.month) if end_date else int(datetime.today().month)
        mon_order = (
            [_cal.month_abbr[m] for m in range(cur_m, 0, -1)]
            + [_cal.month_abbr[m] for m in range(12, cur_m, -1)]
        )

        base_cols = [asset_col, "YTD Total"]
        month_cols = [m for m in mon_order if m in df_asset.columns]
        remaining = [c for c in df_asset.columns if c not in (base_cols + month_cols)]
        df_asset = df_asset[base_cols + month_cols + remaining]

        y = pd.to_numeric(df_asset["YTD Total"], errors="coerce").fillna(0.0)
        df_asset = df_asset.loc[y.sort_values(ascending=False).index].reset_index(drop=True)

    if df_asset is None or df_asset.empty:
        c.drawString(0.75 * inch, height - 1.3 * inch, "No YTD asset summary rows to display.")
        c.showPage()
    else:
        # Format AFTER ordering
        df_asset = _fmt_currency(df_asset, skip_first=True).reset_index(drop=True)

        top_y_asset = height - 1.8 * inch

        MAX_ROWS_PER_PAGE = 12  # <-- tune as needed (does NOT count header)

        total_rows = len(df_asset)
        start_i = 0
        page_n = 1

        while start_i < total_rows:
            chunk = df_asset.iloc[start_i:start_i + MAX_ROWS_PER_PAGE].copy()

            _draw_table_paged(
                chunk,
                x=0.75 * inch,
                top_y=top_y_asset,
                max_width=width - 1.5 * inch,
                bottom_margin=0.7 * inch,
                font_size=6,
                page_title=f"YTD Summary by Asset (p{page_n})",
            )

            start_i += MAX_ROWS_PER_PAGE
            page_n += 1

            if start_i < total_rows:
                c.showPage()
                c.setPageSize(landscape(letter))
                width, height = landscape(letter)
                _title("YTD Summary by Asset (cont.)", height - 0.7 * inch, 18)
                c.setFont("Helvetica", 10)

        c.showPage()


    # ---------------------------------------------------
    # 5) TRANSACTIONS — FILTERED DETAIL  (paged)
    # ---------------------------------------------------
    c.setPageSize(landscape(letter))
    width, height = landscape(letter)
    _title("Transactions — Filtered Detail", height - 0.7 * inch, 18)
    c.setFont("Helvetica", 9)

    base_tx = df_tx_scoped.copy() if df_tx_scoped is not None else pd.DataFrame()

    col_in = _find(base_tx, "trans_in", "Trans In", "In")
    col_out = _find(base_tx, "trans_out", "Trans Out", "Out")

    rows_count = len(base_tx)

    cost_col = _find(
        base_tx,
        "_COST", "Cost", "total_cost", "TOTAL COST",
        "Ext Cost", "Extended Cost", "Value",
    )

    in_sum = out_sum = 0.0
    if cost_col and not base_tx.empty:
        in_mask_win = _truthy_mask(
            base_tx[col_in] if (col_in and col_in in base_tx.columns) else None,
            base_tx.index,
        )
        out_mask_win = _truthy_mask(
            base_tx[col_out] if (col_out and col_out in base_tx.columns) else None,
            base_tx.index,
        )

        in_sum = float(pd.to_numeric(base_tx.loc[in_mask_win, cost_col], errors="coerce").fillna(0).sum())
        out_sum = float(pd.to_numeric(base_tx.loc[out_mask_win, cost_col], errors="coerce").fillna(0).sum())

    net_sum = in_sum + out_sum

    c.drawString(0.75 * inch, height - 1.3 * inch, f"Rows: {rows_count:,}")
    c.drawString(3.0 * inch, height - 1.3 * inch, f"IN $ (window): ${in_sum:,.2f}")
    c.drawString(6.0 * inch, height - 1.3 * inch, f"OUT $ (window): ${out_sum:,.2f}")
    c.drawString(9.0 * inch, height - 1.3 * inch, f"IN+OUT $ (window): ${net_sum:,.2f}")

    # Trim some columns to shorten the Transactions table
    if not base_tx.empty:
        drop_targets = [
            "part type",
            "unit cost",
            "total cost",
            "trans_in",
            "trans in",
            "trans_out",
            "trans out",
            "_cost",
        ]
        lower_map = {str(c).casefold(): c for c in base_tx.columns}
        cols_to_drop = [lower_map[k] for k in drop_targets if k in lower_map]
        if cols_to_drop:
            base_tx = base_tx.drop(columns=cols_to_drop, errors="ignore")

    top_y = height - 1.8 * inch

    MAX_TX_ROWS_PER_PAGE = 10
    # <--- your target (does NOT count header)

    if base_tx is None or base_tx.empty:
        c.drawString(0.75 * inch, top_y, "No transactions in this window.")
        c.showPage()
    else:
        base_tx = base_tx.reset_index(drop=True)
        total_rows = len(base_tx)
        start_i = 0
        page_n = 1

        while start_i < total_rows:
            chunk = base_tx.iloc[start_i:start_i + MAX_TX_ROWS_PER_PAGE].copy()

            _draw_table_paged(
                chunk,
                x=0.5 * inch,
                top_y=top_y,
                max_width=width - 1.0 * inch,
                bottom_margin=0.7 * inch,
                font_size=6,
                page_title=f"Transactions — Filtered Detail (p{page_n})",
            )

            start_i += MAX_TX_ROWS_PER_PAGE
            page_n += 1

            if start_i < total_rows:
                c.showPage()
                c.setPageSize(landscape(letter))
                width, height = landscape(letter)
                _title("Transactions — Filtered Detail (cont.)", height - 0.7 * inch, 18)
                c.setFont("Helvetica", 9)

        c.showPage()

    # ---------------------------------------------------
    # 6) WORK ORDERS — KPI PAGES (one per table)  [PAGED]
    # ---------------------------------------------------
    wo_kpis = _compute_wo_kpis_for_pdf(start_date, end_date, locations)
    wo_tables = _compute_wo_tables_for_pdf(start_date, end_date, locations)

    wo_blocks = [
        ("IsOverDue (Open)", "overdue", "overdue_count", "overdue_avg_days"),
        ("Completed Overdue", "com_overdue", "com_overdue_count", "com_overdue_avg_days"),
        ("Open", "open", "open_count", "open_avg_days"),
        ("Completed in Window", "completed", "completed_count", None),
    ]

    MAX_ROWS_PER_PAGE = 10  # <-- tune as needed (does NOT count header)

    for label, tbl_key, count_key, avg_key in wo_blocks:
        c.setPageSize(landscape(letter))
        width, height = landscape(letter)
        _title("Work Orders — " + label, height - 0.7 * inch, 18)
        c.setFont("Helvetica", 9)

        count_val = wo_kpis.get(count_key, 0)
        line = f"Count: {int(count_val):,}"
        if avg_key:
            avg_val = wo_kpis.get(avg_key, 0.0)
            line += f"   •   Avg days: {avg_val:.2f}"
        c.drawString(0.75 * inch, height - 1.3 * inch, line)

        tbl_df = wo_tables.get(tbl_key, pd.DataFrame())

        # For Completed Overdue, force filter by Completed On within the main date range
        if label == "Completed Overdue" and start_date and end_date and not tbl_df.empty:
            col_done = _find(tbl_df, "Completed On", "Completed on", "Completed Date", "Date Completed")
            if col_done:
                tbl_df[col_done] = pd.to_datetime(tbl_df[col_done], errors="coerce")
                mask = (
                    (tbl_df[col_done] >= pd.to_datetime(start_date)) &
                    (tbl_df[col_done] <= pd.to_datetime(end_date))
                )
                tbl_df = tbl_df.loc[mask].reset_index(drop=True)

        # Drop helper/technical columns so the WO tables are cleaner
        if not tbl_df.empty:
            helper_names = {
                "IsOverDue", "IsOverdue", "IsComingDue",
                "IsOpen", "IsCompleted",
                "DaysOverDue", "DaysOverdue", "Days Overdue",
                "DueBucket", "Bucket",
                "ComOverDue",
            }
            clean_cols = []
            for col in tbl_df.columns:
                if str(col).startswith("_"):
                    continue
                if col in helper_names:
                    continue
                clean_cols.append(col)
            if clean_cols:
                tbl_df = tbl_df[clean_cols].reset_index(drop=True)

        top_y = height - 1.8 * inch

        if tbl_df is None or tbl_df.empty:
            c.drawString(0.75 * inch, top_y, "No rows to display.")
            c.showPage()
            continue

        # --- Paged draw ---
        total_rows = len(tbl_df)
        start_i = 0
        page_n = 1

        while start_i < total_rows:
            chunk = tbl_df.iloc[start_i:start_i + MAX_ROWS_PER_PAGE].copy()

            _draw_table_paged(
                chunk,
                x=0.5 * inch,
                top_y=top_y,
                max_width=width - 1.0 * inch,
                bottom_margin=0.7 * inch,
                font_size=6,
                page_title=f"Work Orders — {label} (p{page_n})",
            )

            start_i += MAX_ROWS_PER_PAGE
            page_n += 1

            if start_i < total_rows:
                c.showPage()
                c.setPageSize(landscape(letter))
                width, height = landscape(letter)
                _title("Work Orders — " + label + " (cont.)", height - 0.7 * inch, 18)
                c.setFont("Helvetica", 9)
                c.drawString(0.75 * inch, height - 1.3 * inch, line)

        c.showPage()


    # ---------------------------------------------------
    # 7) SERVICES / EXPECTED — All + KPI PAGES
    # ---------------------------------------------------
    from typing import Dict

    exp_kpis = _compute_expected_kpis_for_pdf(df_expected)

    # filter expected by locations (if a location column exists)
    if df_expected is not None and not df_expected.empty:
        loc_col_exp = _find(df_expected, "Location", "Location2", "loc", "NS Location")
        if loc_col_exp and locations:
            allowed = {str(x).strip() for x in locations}
            df_expected = df_expected[
                df_expected[loc_col_exp].astype(str).str.strip().isin(allowed)
            ]

    if df_expected is None or df_expected.empty:
        # No Expected data — skip Expected pages, but DO NOT return
        df_expected = pd.DataFrame()

    # ---- Build a slim Expected matrix similar to the page view ----
    dfE = df_expected.copy()

    def _pick_col(df_in: pd.DataFrame, candidates):
        lowmap = {str(c).lower(): c for c in df_in.columns}
        for c in candidates:
            if c in df_in.columns:
                return c
            lc = str(c).lower()
            if lc in lowmap:
                return lowmap[lc]
        return None

    col_asset      = _pick_col(dfE, ["Asset Name", "Asset", "Name"])
    col_loc        = _pick_col(dfE, ["Location", "Location2", "NS Location", "loc"])
    col_status     = "Status" if "Status" in dfE.columns else None
    col_meter      = _pick_col(dfE, ["Meter Type", "MeterType", "Meter"])
    col_ser        = _pick_col(dfE, ["SerType", "Service Type", "Type"])
    col_date       = _pick_col(dfE, ["Date", "Date of Last service", "Last Service Date"])
    col_due        = _pick_col(dfE, ["DueDate", "Due Date", "Next Due", "Next Service Date"])
    col_next       = _pick_col(dfE, ["Next Service", "NextService", "Next Service Due"])
    col_hours      = _pick_col(dfE, ["Hours", "Meter Reading", "Current Reading"])
    col_interval   = _pick_col(dfE, ["Interval", "Schedule", "Scedule", "Schedule Days"])
    col_in_date    = _pick_col(dfE, ["InDate", "In Date"])
    col_is_overdue = _pick_col(dfE, ["IsOverDue", "IsOverdue"])
    col_remaining  = _pick_col(dfE, ["Remaining", "Remain", "Remaining Hours", "Remaining Miles", "Remaining Days"])
    col_needs_serv = _pick_col(dfE, ["NeedsService", "Needs Service"])
    col_new_read   = _pick_col(dfE, ["NewReadingNeeded", "New Reading Needed"])

    exp_disp_cols: list[tuple[str, str | None]] = [
        ("Asset", col_asset),
        ("Location", col_loc),
        ("Status", col_status),
        ("Meter Type", col_meter),
        ("SerType", col_ser),
        ("Date", col_date),
        ("DueDate", col_due),
        ("Next Service", col_next),
        ("Hours", col_hours),
        ("Interval", col_interval),
        ("InDate", col_in_date),
        ("IsOverDue", col_is_overdue),
        ("Remaining", col_remaining),
        ("NeedsService", col_needs_serv),
        ("NewReadingNeeded", col_new_read),
    ]

    cols_final = [lab for lab, src in exp_disp_cols if src is not None]
    exp_slim = pd.DataFrame(index=dfE.index)
    for lab, src in exp_disp_cols:
        if src is not None and src in dfE.columns:
            exp_slim[lab] = dfE[src]
    exp_slim = exp_slim[cols_final] if cols_final else exp_slim

    if "Date" in exp_slim.columns:
        s = pd.to_datetime(exp_slim["Date"], errors="coerce")
        exp_slim["Date"] = s.dt.date.where(~s.isna(), exp_slim["Date"])

    # --- 7a: All Assets in Expected Matrix ---
    c.setPageSize(landscape(letter))
    width, height = landscape(letter)
    _title("Expected Services — All Assets in Matrix", height - 0.7 * inch, 18)
    c.setFont("Helvetica", 9)
    c.drawString(0.75 * inch, height - 1.3 * inch, f"Rows: {len(exp_slim):,}")

    top_y = height - 1.8 * inch
    _draw_table_paged(
        exp_slim,
        x=0.75 * inch,
        top_y=top_y,
        max_width=width - 1.5 * inch,
        bottom_margin=0.7 * inch,
        font_size=6,
        page_title="Expected Services — All Assets in Matrix",
    )

    c.showPage()

    exp_tables: Dict[str, pd.DataFrame] = {}
    if "Status" in dfE.columns:
        tmp = dfE.copy()
        tmp["_status_cf"] = tmp["Status"].astype(str).str.strip().str.casefold()
        for key, label in {
            "overdue": "overdue",
            "needs_service": "needs service",
            "new_reading_needed": "new reading needed",
            "expected": "expected",
        }.items():
            sub = tmp[tmp["_status_cf"] == label].drop(columns=["_status_cf"])
            if not sub.empty:
                exp_tables[key] = sub.reset_index(drop=True)

    def _expected_slim(tbl: pd.DataFrame) -> pd.DataFrame:
        if tbl is None or tbl.empty:
            return pd.DataFrame(columns=cols_final)
        out = pd.DataFrame(index=tbl.index)
        for lab, src in exp_disp_cols:
            if src is not None and src in tbl.columns:
                out[lab] = tbl[src]
        out = out[cols_final] if cols_final else out
        if "Date" in out.columns:
            s = pd.to_datetime(out["Date"], errors="coerce")
            out["Date"] = s.dt.date.where(~s.isna(), out["Date"])
        return out.reset_index(drop=True)

    blocks_exp = [
        ("Expected (Coming Due)", "expected", None),
        ("Overdue", "overdue", "overdue"),
        ("Needs Service", "needs_service", "needs_service"),
        ("New Reading Needed", "new_reading_needed", "new_reading_needed"),
    ]

    for label, tbl_key, kpi_key in blocks_exp:
        tbl_df_raw = exp_tables.get(tbl_key, pd.DataFrame())
        tbl_df = _expected_slim(tbl_df_raw)

        c.setPageSize(landscape(letter))
        width, height = landscape(letter)
        _title("Expected Services — " + label, height - 0.7 * inch, 18)
        c.setFont("Helvetica", 9)

        count_val = exp_kpis.get(kpi_key, 0) if kpi_key else len(tbl_df)

        c.drawString(0.75 * inch, height - 1.3 * inch, f"Count: {int(count_val):,}")

        top_y = height - 1.8 * inch
        _draw_table_paged(
            tbl_df,
            x=0.5 * inch,
            top_y=top_y,
            max_width=width - 1.0 * inch,
            bottom_margin=0.7 * inch,
            font_size=6,
            page_title="Expected Services — " + label,
        )

        c.showPage()

    # Finish
    c.save()
    return buf.getvalue()


    
# ---------- Shared helpers for WO report (safe to add once) ----------

if "_safe_str" not in globals():
    def _safe_str(val) -> str:
        """Return '' for None / NaN / NA, else string."""
        if val is None:
            return ""
        try:
            if pd.isna(val):
                return ""
        except TypeError:
            pass
        return str(val)


if "_wrap_text" not in globals():
    from reportlab.pdfbase import pdfmetrics

    def _wrap_text(text: str, max_width: float,
                   font_name: str = "Helvetica",
                   font_size: float = 10) -> list[str]:
        """
        Wrap a string into multiple lines so each line fits within max_width.
        Used for PDF text cells.
        """
        if not text:
            return []
        words = str(text).split()
        lines: list[str] = []
        current = ""

        for w in words:
            trial = f"{current} {w}".strip()
            width = pdfmetrics.stringWidth(trial, font_name, font_size)
            if width <= max_width:
                current = trial
            else:
                if current:
                    lines.append(current)
                current = w
        if current:
            lines.append(current)
        return lines


if "_format_date" not in globals():
    import numpy as np
    from datetime import datetime

    def _format_date(val) -> str:
        """
        Nice date formatting; handles pandas Timestamps, strings, NaN, 0, 1/0/1900, etc.

        Output: mm/dd/yyyy or '' if it's effectively blank.
        """

        # Treat obvious nulls as blank
        if val is None:
            return ""

        # NaN / NA / NaT
        try:
            if pd.isna(val):
                return ""
        except TypeError:
            pass

        # If it's already a Timestamp or datetime-like
        if isinstance(val, (pd.Timestamp, datetime, np.datetime64)):
            ts = pd.to_datetime(val, errors="coerce")
            if ts is pd.NaT:
                return ""
            # Treat anything before year 2000 as "no real date" in this system
            if ts.year < 2000:
                return ""
            return ts.strftime("%m/%d/%Y")

        # Numbers (Excel/placeholder 0, etc.) -> treat as blank
        if isinstance(val, (int, float)):
            # zero / negative => blank
            if val <= 0:
                return ""
            # If you ever want to decode real serials, handle here; for now bail out
            return ""

        # Strings: clean and try to parse
        s = str(val).strip()
        if not s:
            return ""
        if s in {"NaT", "0", "1/0/1900"}:
            return ""

        ts = pd.to_datetime(s, errors="coerce")
        if ts is not pd.NaT and ts.year >= 2000:
            return ts.strftime("%m/%d/%Y")

        # Fallback: just return the string as-is if parsing fails
        return s







# ========================================
# Load datasets
# ========================================

dfs: Dict[str, pd.DataFrame] = {}
for key, fname in DATA_FILES.items():
    dfs[key] = load_parquet(PARQUET_DIR / fname)
    
    



# ========================================
# Global filters (shared by pages)
# ========================================

all_locations = _collect_parent_locations_from_locdb(PARQUET_DIR)

st.sidebar.markdown("### Global Filters")

# Guess overall min/max from data
guess_start, guess_end = _guess_date_range(dfs)

today = date.today()
twelve_months_ago = today - timedelta(days=365)

# Safety fallbacks
if guess_start is None:
    guess_start = twelve_months_ago
if guess_end is None:
    guess_end = today

# Default: last 12 months, but never earlier than guessed min
default_start = max(guess_start, twelve_months_ago)
# Default end: min(guessed max, today)
default_end = guess_end if guess_end <= today else today

start_date = st.sidebar.date_input("Start date", value=default_start)
end_date = st.sidebar.date_input("End date", value=default_end)

if start_date > end_date:
    st.sidebar.error("Start date cannot be after end date.")
    st.stop()

# Locations:
# • UI: start with none selected
# • Logic: your page code already treats empty as "all"
if len(all_locations) > 0:
    selected_locations = st.sidebar.multiselect(
        "Locations",
        options=list(all_locations),
        default=[],
        help="If no locations are selected, all locations are included.",
    )
else:
    selected_locations = []



# ========================================
# PAGE: 💵 Costs & Trends
# ========================================
def _render_costs_trends(df_all: pd.DataFrame,
                         global_start: date,
                         global_end: date,
                         selected_locations):
    """
    Costs & Trends page:
    - Uses Workorders.parquet (df_all)
    - Master scope:
        * Year = global_end.year (full Jan–Dec)
        * Location = selected_locations
    """
    import calendar
    import pandas as pd
    import streamlit as st
    import numpy as np
    from io import BytesIO

    st.markdown("### 💵 Costs & Trends — Completed WO / PO / Transactions")

    if df_all.empty:
        st.info("Workorders.parquet is empty or not found.")
        return

    # ---------- Small helper: df -> XLSX bytes ----------
    def _df_to_xlsx_bytes(df: pd.DataFrame, sheet_name: str = "Sheet1") -> bytes:
        if df is None or df.empty:
            return b""
        buffer = BytesIO()
        with pd.ExcelWriter(buffer, engine="xlsxwriter") as writer:
            df.to_excel(writer, index=False, sheet_name=sheet_name)
        buffer.seek(0)
        return buffer.getvalue()

    df = df_all.copy()

    # ---------- Basic cleanup ----------
    for c in df.columns:
        if df[c].dtype == object:
            df[c] = df[c].astype("string").str.strip()

    # ---------- Column detection ----------
    def _first_present(df: pd.DataFrame, candidates):
        for c in candidates:
            if c in df.columns:
                return c
        return None

    date_col   = _first_present(df, ["Completed on","COMPLETED ON","Date","Service date","Created on","Due date","Started on"])
    loc_col    = _first_present(df, ["Location","NS Location","location","Location2"])
    asset_col  = _first_present(df, ["Asset","ASSET","asset"])
    wo_col     = _first_present(df, ["WORKORDER","Workorder","Work Order","WO_ID","ID"])
    title_col  = _first_present(df, ["TITLE","Title"])
    po_col     = _first_present(df, ["PO","Po","po","Purchase Order"])
    sort_col   = _first_present(df, ["Sort","SORT","sort"])

    if not date_col:
        st.error("Could not find a completion/date column in Workorders.parquet.")
        st.dataframe(df.head(50), use_container_width=True)
        return

    # ---------- Types & helper cols ----------
    df[date_col] = pd.to_datetime(df[date_col], errors="coerce")

    if sort_col:
        df[sort_col] = pd.to_numeric(df[sort_col], errors="coerce").fillna(1).astype(int)
    else:
        df["Sort"] = 1
        sort_col = "Sort"

    # Cost column
    cost_col = None
    if "_Cost" in df.columns:
        cost_col = "_Cost"
    elif "__cost" in df.columns:
        cost_col = "__cost"
    else:
        def _num(s): return pd.to_numeric(s, errors="coerce").fillna(0.0)
        item_col = _first_present(df, ["Total Item Cost","total item cost","item total","item_total","TotalItemCost"])
        totl_col = _first_present(df, ["Total Cost","total cost","cost total","cost_total","TotalCost"])
        item_cost = _num(df[item_col]) if item_col else 0.0
        totl_cost = _num(df[totl_col]) if totl_col else 0.0
        df["__cost"] = (item_cost + totl_cost).astype(float)
        cost_col = "__cost"

    df[cost_col] = pd.to_numeric(df[cost_col], errors="coerce").fillna(0.0)

    # Row index + year/month/day
    df["__row"]   = np.arange(len(df))
    df["__Year"]  = pd.to_numeric(df[date_col].dt.year, errors="coerce").astype("Int64")
    df["__Month"] = df[date_col].dt.month
    df["__Day"]   = pd.to_datetime(df[date_col].dt.date)

    # Group id
    if wo_col:
        g = df[wo_col].astype("string").str.strip().replace("", pd.NA).ffill()
        df["__grp"] = g.fillna((df[sort_col].eq(1)).cumsum().astype(str)).astype(str)
    else:
        df["__grp"] = (df[sort_col].eq(1)).cumsum().astype(str)

    # ---------- YEAR scope (full Jan–Dec for master year) ----------
    year_i = int(global_end.year)
    df_year = df[df["__Year"] == year_i].copy()

    # Apply master location filter
    if loc_col and selected_locations:
        loc_strs = [str(x) for x in selected_locations]
        df_year = df_year[df_year[loc_col].astype(str).isin(loc_strs)]

    if df_year.empty:
        st.info(f"No records found for year {year_i} and selected locations.")
        return

    # ---------- KPIs ----------
    total_cost = df_year[cost_col].sum()
    num_rows   = len(df_year)
    loc_label  = ", ".join(map(str, selected_locations)) if selected_locations else "All Locations"

    st.caption(f"Year: {year_i}  •  Locations: {loc_label}")
    top_row, = st.columns([1])
    top_row.metric("YTD Cost (Jan–Dec)", f"${total_cost:,.2f}", help=f"Rows: {num_rows:,}")

    # ---------- Helper: YTD pivot ----------
    def _ytd_month_pivot(df_in: pd.DataFrame, index_col: str, cost_col_in: str) -> pd.DataFrame:
        g = df_in.groupby([index_col, "__Month"], as_index=False)[cost_col_in].sum()
        g["Mon"] = g["__Month"].map(lambda m: calendar.month_abbr[int(m)] if pd.notna(m) else "")
        pv = g.pivot(index=index_col, columns="Mon", values=cost_col_in).reset_index()

        # Order columns: Dec, Nov, ..., Jan
        mon_desc = [calendar.month_abbr[m] for m in range(12, 0, -1)]
        month_cols = [c for c in mon_desc if c in pv.columns]

        for c in month_cols:
            pv[c] = pd.to_numeric(pv[c], errors="coerce").fillna(0.0)

        pv["YTD Total"] = pv[month_cols].sum(axis=1, skipna=True)

        ordered_cols = [index_col, "YTD Total"] + month_cols
        pv = pv[ordered_cols].sort_values("YTD Total", ascending=False)
        return pv

    # ---------- LEFT: YTD tables ----------
    left_col, right_col = st.columns([7, 5])

    with left_col:
        index_col = loc_col if loc_col else asset_col
        if not index_col:
            st.markdown("#### YTD Summary")
            st.info("No Location or Asset column to summarize by.")
            by_left = pd.DataFrame()
        else:
            by_left = _ytd_month_pivot(df_year, index_col, cost_col)

        st.markdown("#### YTD Summary by Location")
        if by_left.empty:
            st.info("No summary available.")
            # Save empty for PDF
            st.session_state["rhub_costs_ytd_loc"] = pd.DataFrame()
        else:
            # Caption + small XLSX button
            cap_col, btn_col = st.columns([4, 1])
            with cap_col:
                st.caption(f"Rows: {len(by_left)}  •  Total: ${by_left['YTD Total'].sum():,.2f}")
            with btn_col:
                xlsx_loc = _df_to_xlsx_bytes(by_left, sheet_name="YTD_Location")
                st.download_button(
                    "⬇️ XLSX",
                    data=xlsx_loc,
                    file_name=f"Costs_YTD_by_Location_{year_i}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key="dl_ytd_loc",
                    use_container_width=True,
                )

            st.dataframe(by_left, use_container_width=True, hide_index=True, height=360)
            # Save exact table for PDF
            st.session_state["rhub_costs_ytd_loc"] = by_left.copy()

        # --- YTD Summary by Asset (for report + optional display) ---
        if asset_col:
            by_asset = _ytd_month_pivot(df_year, asset_col, cost_col)
            st.session_state["rhub_costs_ytd_asset"] = by_asset.copy()

            st.markdown("#### YTD Summary by Asset")
            if by_asset.empty:
                st.info("No asset summary available.")
                st.session_state["rhub_costs_ytd_asset"] = pd.DataFrame()
            else:
                cap_col2, btn_col2 = st.columns([4, 1])
                with cap_col2:
                    st.caption(
                        f"Rows: {len(by_asset)}  •  Total: ${by_asset['YTD Total'].sum():,.2f}"
                    )
                with btn_col2:
                    xlsx_asset = _df_to_xlsx_bytes(by_asset, sheet_name="YTD_Asset")
                    st.download_button(
                        "⬇️ XLSX",
                        data=xlsx_asset,
                        file_name=f"Costs_YTD_by_Asset_{year_i}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        key="dl_ytd_asset",
                        use_container_width=True,
                    )

                st.dataframe(
                    by_asset,
                    use_container_width=True,
                    hide_index=True,
                    height=360,
                )

        else:
            st.session_state["rhub_costs_ytd_asset"] = pd.DataFrame()

    # ---------- RIGHT: Details ----------
    with right_col:
        details_scope = df_year.copy()

        if asset_col and asset_col in details_scope.columns:
            asset_opts = ["All"] + sorted(
                details_scope[asset_col].dropna().astype(str).unique().tolist()
            )
            asset_sel = st.selectbox("Asset (details filter)", options=asset_opts, index=0)
            if asset_sel != "All":
                details_scope = details_scope[details_scope[asset_col].astype(str) == asset_sel]
        else:
            asset_sel = "All"

        d = details_scope.copy()
        d[date_col] = pd.to_datetime(d[date_col], errors="coerce").dt.date
        d = d.sort_values(by="__row", kind="stable")

        cols_all = list(d.columns)
        for h in ["__row","__Year","__Month","__Day","__grp","__MonAbbr","__header_flag"]:
            if h in cols_all:
                cols_all.remove(h)

        core_order = []
        if wo_col and wo_col in d.columns:
            core_order.append(wo_col)
        for c in [title_col, "Description", date_col, po_col, cost_col]:
            if c and c in d.columns and c not in core_order:
                core_order.append(c)
        ordered = core_order + [c for c in cols_all if c not in core_order]
        details = d[ordered] if ordered else d

        right_total = pd.to_numeric(details.get(cost_col, 0), errors="coerce").fillna(0).sum()

        label = "All assets" if asset_sel == "All" else f"Asset: {asset_sel}"
        st.markdown(f"#### Details — {label}")

        # Caption + small XLSX button
        cap_col3, btn_col3 = st.columns([4, 1])
        with cap_col3:
            st.caption(f"Year: {year_i}  •  Rows: {len(details)}  •  Total: ${right_total:,.2f}")
        with btn_col3:
            safe_label = label.replace(" ", "_").replace(":", "_")
            xlsx_details = _df_to_xlsx_bytes(details, sheet_name="Details")
            st.download_button(
                "⬇️ XLSX",
                data=xlsx_details,
                file_name=f"Costs_Details_{safe_label}_{year_i}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="dl_details",
                use_container_width=True,
            )

        st.dataframe(details, use_container_width=True, hide_index=True, height=360)

    # ---------- Simple monthly chart ----------
    try:
        import altair as alt
        chart_df = (df_year.groupby(["__Month"], as_index=False)[cost_col].sum()
                             .rename(columns={cost_col: "Total"}))
        if not chart_df.empty:
            chart_df["Month"] = chart_df["__Month"].map(
                lambda m: calendar.month_abbr[int(m)] if pd.notna(m) else ""
            )
            ch = (alt.Chart(chart_df)
                    .mark_line(point=True)
                    .encode(
                        x=alt.X("Month:N",
                                sort=[calendar.month_abbr[m] for m in range(1,13)],
                                title="Month"),
                        y=alt.Y("Total:Q", title="Total Cost"),
                        tooltip=["Month", "Total"]
                    ).properties(height=220, title=f"Cost by Month — {year_i}"))
            st.altair_chart(ch, use_container_width=True)
    except Exception:
        pass



# ======================================================================
# PAGE: ↕ Transactions — Inventory vs Non-Inventory
#   Uses main filters: date window + locations
# ======================================================================
def render_transactions(df_tx, start_date, end_date, selected_locations):
    from io import BytesIO
    import pandas as pd
    import streamlit as st

    st.markdown("### ↕ Transactions — Inventory vs Non-Inventory")

    if df_tx is None or df_tx.empty:
        st.info("No transaction records for the selected filters.")
        return

    # ---------- XLSX helper ----------
    if "to_xlsx_bytes" not in globals():
        def to_xlsx_bytes(df, sheet_name=None, sheet=None):
            name = sheet or sheet_name or "Sheet1"
            buf = BytesIO()
            with pd.ExcelWriter(buf, engine="xlsxwriter") as xw:
                df.to_excel(xw, index=False, sheet_name=name)
            return buf.getvalue()
        globals()["to_xlsx_bytes"] = to_xlsx_bytes
    else:
        to_xlsx_bytes = globals()["to_xlsx_bytes"]

    # ---------- utilities ----------
    def _fmt_currency_table(df, exclude=("Location",)):
        out = df.copy()
        if out.empty:
            return out
        for col in out.columns:
            if col in exclude:
                continue
            s = pd.to_numeric(out[col], errors="coerce")
            out[col] = s.map(lambda v: f"${v:,.2f}" if pd.notna(v) else "")
        return out

    def _find(df, *names):
        names_l = [n.lower() for n in names]
        m = {str(c).strip().lower(): c for c in df.columns}
        for n in names_l:
            if n in m:
                return m[n]
        for k, orig in m.items():
            for n in names_l:
                if n.replace("_", " ") in k:
                    return orig
        return None

    def _truthy_mask(series: pd.Series | None, idx) -> pd.Series:
        if series is None:
            return pd.Series(False, index=idx)
        s = series
        if pd.api.types.is_bool_dtype(s):
            return s.fillna(False)
        num = pd.to_numeric(s, errors="coerce")
        if num.notna().any():
            return num.fillna(0).astype(int).astype(bool)
        sval = s.astype(str).str.strip().str.lower()
        return sval.isin({"true", "t", "yes", "y", "1"})

    # ---------- work on a copy ----------
    df_tx = df_tx.copy()

    # ---------- discover columns ----------
    col_tc     = _find(df_tx, "TOTAL COST", "Ext Cost", "Extended Cost", "Value")
    col_date   = _find(df_tx, "DATE", "Date")
    col_reason = _find(df_tx, "TRANSACTION REASON", "Reason")
    col_loc    = _find(df_tx, "LOCATION", "Location", "NS Location")
    col_in     = _find(df_tx, "trans_in", "Trans In", "In")
    col_out    = _find(df_tx, "trans_out", "Trans Out", "Out")

    # basic type cleanup
    if col_date in df_tx.columns:
        try:
            df_tx[col_date] = pd.to_datetime(df_tx[col_date], errors="coerce").dt.date
        except Exception:
            pass
    if col_tc in df_tx.columns:
        df_tx[col_tc] = pd.to_numeric(df_tx[col_tc], errors="coerce")

    # Apply main location filter
    if col_loc and selected_locations:
        loc_strs = [str(x) for x in selected_locations]
        df_tx = df_tx[df_tx[col_loc].astype(str).isin(loc_strs)]

    inv_in_mask_all  = _truthy_mask(df_tx[col_in]  if col_in  in df_tx.columns else None, df_tx.index)
    inv_out_mask_all = _truthy_mask(df_tx[col_out] if col_out in df_tx.columns else None, df_tx.index)

    # ---------- Page-level controls ----------
    mode = st.radio("View", ["Inventory", "Non-Inventory"], horizontal=True, key="tx_mode")

    if mode == "Non-Inventory":
        rea_all = sorted(df_tx[col_reason].dropna().astype(str).unique().tolist()) if col_reason else []
        st.caption("Reason (optional)")
        pick_reason = st.multiselect(
            "",
            rea_all,
            default=[],
            label_visibility="collapsed",
            placeholder="All reasons",
            key="tx_reason",
        )
    else:
        pick_reason = []

    details_dir = st.radio(
        "Details direction",
        ["Both", "In only", "Out only"],
        horizontal=True,
        key="tx_dir",
    )

    # Use main Reporting Hub date window
    start = start_date
    end   = end_date

    # ---------- Apply base filters ----------
    base = df_tx.copy()

    # Inventory vs Non-Inventory split
    if mode == "Inventory":
        keep_mask = inv_in_mask_all | inv_out_mask_all
    else:
        keep_mask = (~inv_in_mask_all.fillna(False)) & (~inv_out_mask_all.fillna(False))
    base = base[keep_mask].copy()

    # date window
    if col_date in base.columns:
        base = base[(base[col_date] >= start) & (base[col_date] <= end)]

    # extra Non-Inventory reason filter
    if mode == "Non-Inventory" and pick_reason and col_reason in base.columns:
        base = base[base[col_reason].astype(str).isin(pick_reason)]

    if base.empty:
        st.info("No transactions after applying filters.")
        return

    # ---------- Month labels inside window ----------
    if col_date in base.columns:
        dts = pd.to_datetime(base[col_date], errors="coerce")
        single_year = (start.year == end.year)
        month_label = dts.dt.strftime("%b") if single_year else dts.dt.strftime("%Y-%b")
        base = base.assign(
            __month=month_label,
            __month_sort=dts.dt.to_period("M").astype(int),
        )
        months_sorted = (
            base[["__month", "__month_sort"]]
            .drop_duplicates()
            .sort_values("__month_sort", ascending=False)["__month"]
            .tolist()
        )
    else:
        months_sorted = []

    # ---------- Summaries ----------
    def _pivot_io(df_in: pd.DataFrame, which: str) -> pd.DataFrame:
        label_ytd = f"YTD {which.title()}"
        cols = ["Location", label_ytd] + [f"{m} {which.title()}" for m in months_sorted]
        if df_in.empty or not col_tc or (col_loc not in df_in.columns):
            return pd.DataFrame(columns=cols)

        mask = _truthy_mask(df_in[col_in] if which == "in" else df_in[col_out], df_in.index) \
               if ((which == "in" and col_in in df_in.columns) or (which == "out" and col_out in df_in.columns)) \
               else pd.Series(False, index=df_in.index)

        sub = df_in.loc[mask].copy()
        if sub.empty:
            return pd.DataFrame(columns=cols)

        g = (
            sub.groupby([col_loc, "__month", "__month_sort"], dropna=False)[col_tc]
            .sum()
            .reset_index()
        )
        pv = g.pivot_table(
            index=col_loc,
            columns="__month",
            values=col_tc,
            aggfunc="sum",
            fill_value=0.0,
        )

        # Ensure all months included, newest→oldest
        for m in months_sorted:
            if m not in pv.columns:
                pv[m] = 0.0
        pv = pv.reindex(columns=months_sorted, fill_value=0.0)

        pv[label_ytd] = pv.sum(axis=1)
        pv = pv.rename(columns={m: f"{m} {which.title()}" for m in months_sorted})
        pv = pv.reset_index().rename(columns={col_loc: "Location"})
        out_cols = ["Location", label_ytd] + [f"{m} {which.title()}" for m in months_sorted]
        pv = pv[out_cols]

        # TOTAL row
        if not pv.empty:
            tot = {"Location": "TOTAL"}
            for c in out_cols[1:]:
                tot[c] = float(pd.to_numeric(pv[c], errors="coerce").fillna(0).sum())
            pv = pd.concat([pv, pd.DataFrame([tot])], ignore_index=True)
        return pv

    pivot_in  = _pivot_io(base, "in")
    pivot_out = _pivot_io(base, "out")

    # ---------- Header metrics (window) ----------
    if col_tc in base.columns:
        in_mask_win  = _truthy_mask(base[col_in]  if col_in  in base.columns else None, base.index)
        out_mask_win = _truthy_mask(base[col_out] if col_out in base.columns else None, base.index)
        in_sum  = float(pd.to_numeric(base.loc[in_mask_win,  col_tc], errors="coerce").fillna(0).sum())
        out_sum = float(pd.to_numeric(base.loc[out_mask_win, col_tc], errors="coerce").fillna(0).sum())
    else:
        in_sum = out_sum = 0.0

    # Treat OUT as positive and show “true difference” as IN + OUT
    net_sum = in_sum + out_sum

    m1, m2, m3, m4 = st.columns([1, 1, 1, 1])
    with m1:
        st.metric("Rows", f"{len(base):,}")
    with m2:
        st.metric("IN $ (window)",  f"${in_sum:,.2f}")
    with m3:
        st.metric("OUT $ (window)", f"${out_sum:,.2f}")
    with m4:
        st.metric("IN+OUT $ (window)", f"${net_sum:,.2f}")

    # ---------- Side-by-side summaries ----------
    left, right = st.columns([7, 5])
    with left:
        st.subheader("IN — Costs by Location (YTD + Months)")
        st.dataframe(
            _fmt_currency_table(pivot_in, exclude=("Location",)),
            use_container_width=True,
            height=360,
        )
        st.download_button(
            "⬇️ Download IN summary (XLSX)",
            data=to_xlsx_bytes(pivot_in, sheet_name="IN_Summary"),
            file_name=f"Transactions_summary_IN_{mode}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
            key="tx_dl_in",
        )
    with right:
        st.subheader("OUT — Costs by Location (YTD + Months)")
        st.dataframe(
            _fmt_currency_table(pivot_out, exclude=("Location",)),
            use_container_width=True,
            height=360,
        )
        st.download_button(
            "⬇️ Download OUT summary (XLSX)",
            data=to_xlsx_bytes(pivot_out, sheet_name="OUT_Summary"),
            file_name=f"Transactions_summary_OUT_{mode}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
            key="tx_dl_out",
        )

    # ---------- Details ----------
    if details_dir == "In only":
        base_details = base[
            _truthy_mask(base[col_in] if col_in in base.columns else None, base.index)
        ].copy()
    elif details_dir == "Out only":
        base_details = base[
            _truthy_mask(base[col_out] if col_out in base.columns else None, base.index)
        ].copy()
    else:
        base_details = base.copy()

    st.subheader("Filtered Details (all columns)")
    st.dataframe(base_details, use_container_width=True, height=420)
    st.download_button(
        "⬇️ Download details (XLSX)",
        data=to_xlsx_bytes(base_details, sheet_name="Details"),
        file_name=f"Transactions_details_{mode}_{details_dir.replace(' ', '_')}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True,
        key="tx_dl_details",
    )
    
    
# ======================================================================
# PAGE RENDER: 📈 Inventory Analysis — Baseline & Monthly (filterable)
#   Uses:
#     - df_parts  (Parts_Master.parquet)
#     - df_tx     (TRANSACTIONS-FILTER.parquet)
#     - start_date / end_date (for Custom default)
#     - locations (for default location pick)
# ======================================================================
def _render_inventory_analysis(df_parts, df_tx, start_date, end_date, locations):
    import os, io, re, json, base64, calendar
    from urllib.parse import quote
    from datetime import date, datetime

    import numpy as np
    import pandas as pd
    import streamlit as st
    import streamlit.components.v1 as components

    # ---------- basic checks ----------
    if df_parts is None or df_parts.empty:
        st.info("No Parts / inventory records available for Inventory Analysis.")
        return
    if df_tx is None or df_tx.empty:
        st.info("No transaction records (TRANSACTIONS-FILTER) available for Inventory Analysis.")
        return

    # ---------- small helpers ----------
    def _find(df, *names):
        names_l = [n.lower() for n in names]
        m = {str(c).strip().lower(): c for c in df.columns}
        for n in names_l:
            if n in m:
                return m[n]
        for k, orig in m.items():
            for n in names_l:
                if n.replace("_"," ") in k:
                    return orig
        return None

    def _to_ts(s): return pd.to_datetime(s, errors="coerce")
    def _num(s):  return pd.to_numeric(s, errors="coerce")

    def _truthy_mask(series: pd.Series | None, idx) -> pd.Series:
        if series is None:
            return pd.Series(False, index=idx)
        s = series
        if pd.api.types.is_bool_dtype(s):
            return s.fillna(False)
        num = pd.to_numeric(s, errors="coerce")
        if num.notna().any():
            return num.fillna(0).astype(int).astype(bool)
        sval = s.astype(str).str.strip().str.lower()
        return sval.isin({"true","t","yes","y","1"})

    st.markdown("### Inventory Analysis — Baseline & Monthly (filterable)")

    # ---------- KPI thresholds persistence ----------
    SETTINGS_PATH = "inv_kpi_settings.json"
    DEFAULT_THRESHOLDS = {
        "ytd_growth_pct": {"type":"band","green_min":0.0,"green_max":10.0,"warn_min":-5.0,"warn_max":20.0,"label":"YTD Inventory Growth %"},
        "avg_mom_pct": {"type":"max_good","goal":2.0,"warn_max":4.0,"label":"Avg MoM Δ %"},
        "coverage_net_pct": {"type":"min_good","goal":95.0,"warn_min":90.0,"label":"PO/WO Coverage (NET %)"},
        "coverage_abs_pct": {"type":"min_good","goal":95.0,"warn_min":90.0,"label":"PO/WO Coverage (ABS %)"},
        "gap_share_net_pct": {"type":"max_good","goal":15.0,"warn_max":20.0,"label":"Adj/Unmatched (NET %)"},
        "avg_error_pct": {"type":"max_good","goal":0.50,"warn_max":1.00,"label":"Avg Error Rate %"},
        "reopen_pct_of_wo": {"type":"max_good","goal":5.0,"warn_max":10.0,"label":"Reopen % of WO activity (ABS)"},
    }
    def load_thresholds():
        try:
            if os.path.exists(SETTINGS_PATH):
                with open(SETTINGS_PATH, "r", encoding="utf-8") as f:
                    data = json.load(f)
                merged = DEFAULT_THRESHOLDS.copy()
                for k, v in data.items():
                    if k in merged and isinstance(v, dict):
                        merged[k].update(v)
                return merged
        except Exception:
            pass
        return DEFAULT_THRESHOLDS.copy()
    def save_thresholds(thr):
        try:
            with open(SETTINGS_PATH, "w", encoding="utf-8") as f:
                json.dump(thr, f, indent=2)
            return True
        except Exception as e:
            st.warning(f"Could not save KPI defaults: {e}")
            return False
    KPI_T = load_thresholds()

    # ----- KPI tile rendering (large, high-contrast) -----
    def _classify(val: float | None, cfg: dict) -> str:
        if val is None or (isinstance(val, float) and np.isnan(val)):
            return "grey"
        t = cfg.get("type", "min_good")
        if t == "band":
            gmin, gmax = cfg["green_min"], cfg["green_max"]
            wmin, wmax = cfg["warn_min"], cfg["warn_max"]
            if wmin <= val <= wmax:
                return "green" if (gmin <= val <= gmax) else "yellow"
            return "red"
        if t == "min_good":
            goal, wmin = cfg["goal"], cfg["warn_min"]
            if val >= goal: return "green"
            if val >= wmin: return "yellow"
            return "red"
        if t == "max_good":
            goal, wmax = cfg["goal"], cfg["warn_max"]
            if val <= goal: return "green"
            if val <= wmax: return "yellow"
            return "red"
        return "grey"

    def _trend_svg(color: str, shape: str) -> str:
        color_map = {"green":"#22c55e","yellow":"#f59e0b","red":"#ef4444","grey":"#9ca3af"}
        stroke = color_map.get(color, "#9ca3af")
        if shape == "flat":
            return f'<svg width="42" height="18" viewBox="0 0 42 18"><line x1="2" y1="9" x2="40" y2="9" stroke="{stroke}" stroke-width="3" stroke-linecap="round"/></svg>'
        if shape == "up":
            return f'<svg width="42" height="18" viewBox="0 0 42 18"><polyline points="2,16 16,10 26,12 38,4" fill="none" stroke="{stroke}" stroke-width="3" stroke-linecap="round" stroke-linejoin="round"/><polygon points="38,4 37,9 42,7" fill="{stroke}"/></svg>'
        return f'<svg width="42" height="18" viewBox="0 0 42 18"><polyline points="2,2 16,8 26,6 38,14" fill="none" stroke="{stroke}" stroke-width="3" stroke-linecap="round" stroke-linejoin="round"/><polygon points="38,14 37,9 42,11" fill="{stroke}"/></svg>'

    def kpi_card(title: str, value: float | None, suffix: str, cfg_key: str):
        cfg = KPI_T[cfg_key]
        color = _classify(value, cfg)
        shape = "flat" if cfg["type"] == "band" else ("down" if cfg["type"] == "max_good" else "up")
        disp = "" if (value is None or (isinstance(value, float) and np.isnan(value))) else f"{value:.2f}{suffix}"
        badge_bg = {"green":"#064e3b","yellow":"#78350f","red":"#7f1d1d","grey":"#374151"}[color]
        badge_fg = {"green":"#a7f3d0","yellow":"#fde68a","red":"#fecaca","grey":"#e5e7eb"}[color]
        border_c = {"green":"#115e59","yellow":"#92400e","red":"#991b1b","grey":"#4b5563"}[color]
        pill_bg  = {"green":"#064e3b26","yellow":"#78350f26","red":"#7f1d1d26","grey":"#1f293726"}[color]
        html = f"""
        <div style="padding:14px 16px;border-radius:14px;background:{pill_bg};border:1px solid {border_c};
                    font-family:system-ui,-apple-system,Segoe UI,Roboto,Helvetica,Arial;">
          <div style="font-size:13px;color:#d1d5db;margin-bottom:8px;">{title}</div>
          <div style="display:flex;align-items:center;justify-content:space-between;gap:10px;">
            {_trend_svg(color, shape)}
            <div style="display:flex;align-items:center;gap:10px;">
              <div style="font-weight:700;font-size:24px;color:#f9fafb;">{disp}</div>
              <span style="padding:3px 10px;border-radius:999px;background:{badge_bg};color:{badge_fg};
                           font-size:12px;border:1px solid rgba(255,255,255,0.08);">{color.upper()}</span>
            </div>
          </div>
        </div>"""
        components.html(html, height=110)

    # ---------- parts columns ----------
    col_total   = _find(df_parts, "Total Cost","Ext Cost","Extended Cost","Value","Inventory Value","Total Value")
    col_qty     = _find(df_parts, "Quantity in Stock","On Hand","Qty On Hand","Quantity","QOH","Stock")
    col_unit    = _find(df_parts, "Unit Cost","Avg Cost","Average Cost","Last Price","Cost")
    col_created = _find(df_parts, "Created","Created Date","Date Created","Added","Added On","Date")
    col_updated = _find(df_parts, "Updated","Last Updated","Modified","Last Modified","Change Date")
    col_loc_parts_candidates = []
    for cand in ["Location", "Location2", "Company"]:
        c = _find(df_parts, cand)
        if c and c in df_parts.columns and c not in col_loc_parts_candidates:
            col_loc_parts_candidates.append(c)

    # value column for parts
    if col_total:
        parts_value_all = _num(df_parts[col_total]).fillna(0.0)
    elif col_qty and col_unit:
        parts_value_all = _num(df_parts[col_qty]).fillna(0.0) * _num(df_parts[col_unit]).fillna(0.0)
    else:
        st.error("Couldn't derive Parts value (need 'Total Cost' or 'Quantity in Stock' × 'Unit Cost').")
        return

    created_ts_all = _to_ts(df_parts[col_created]) if col_created else pd.Series(pd.NaT, index=df_parts.index)
    updated_ts_all = _to_ts(df_parts[col_updated]) if col_updated else pd.Series(pd.NaT, index=df_parts.index)

    # ---------- transactions columns ----------
    col_tc     = _find(df_tx, "TOTAL COST","Ext Cost","Extended Cost","Value","Amount","AMOUNT")
    col_date   = _find(df_tx, "DATE","Date")
    col_loc    = _find(df_tx, "LOCATION","Location","Location2")
    col_dir    = _find(df_tx, "DIRECTION","Direction")
    col_tin    = _find(df_tx, "trans_in","Trans In","In")
    col_tout   = _find(df_tx, "trans_out","Trans Out","Out")
    col_reason = _find(df_tx, "TRANSACTION REASON","Transaction Reason","Reason")
    col_wo     = _find(df_tx, "PO WORK ORDER","WO","Work Order")
    col_part   = _find(df_tx, "PART NAME","Item","Item Name")

    if col_tc:   df_tx[col_tc] = _num(df_tx[col_tc])
    if col_date: df_tx[col_date] = pd.to_datetime(df_tx[col_date], errors="coerce")

    # ---------- Location Rollup JSON ----------
    ROLLUP_JSON = "location_rollup.json"
    ROLLUP_DEFAULT = {
        "300 - Greer Lime - Stone": ["300","310","320","335"],
    }
    def _load_rollup():
        try:
            if os.path.exists(ROLLUP_JSON):
                with open(ROLLUP_JSON, "r", encoding="utf-8") as f:
                    return json.load(f)
        except Exception:
            pass
        return ROLLUP_DEFAULT
    def _save_rollup(data: dict):
        try:
            with open(ROLLUP_JSON, "w", encoding="utf-8") as f:
                json.dump(data, f, indent=2)
        except Exception as e:
            st.warning(f"Could not save {ROLLUP_JSON}: {e}")
    ROLLUP = _load_rollup()

    def _extract_code(cell) -> str | None:
        if pd.isna(cell): return None
        m = re.match(r"^\s*(\d+)", str(cell))
        return m.group(1) if m else None

    def _codes_for_parent(pick_label: str) -> set[str]:
        if pick_label in ROLLUP:
            return set(ROLLUP[pick_label])
        code = _extract_code(pick_label)
        return set([code]) if code else set()

    # ---------- LOCATION OPTIONS (driven by main-page filters) ----------
    ALL = "« All »"
    admin_tokens = {"« All locations »", "« All companies (admin) »", "(all)", "All", "ALL"}

    # locations comes from the main page selected_locations (list of parent labels)
    if locations:
        eff_locations = [str(l) for l in locations if str(l) not in admin_tokens]
    else:
        eff_locations = []

    if not eff_locations:
        loc_label_for_display = "All Locations"
    elif len(eff_locations) == 1:
        loc_label_for_display = eff_locations[0]
    else:
        loc_label_for_display = ", ".join(eff_locations)

    # Period controls (no per-page Location dropdown)
    c_loclbl, c_mode, c_month, c_year = st.columns([2, 1.4, 1, 1])

    with c_loclbl:
        st.markdown(f"**Location(s):** {loc_label_for_display}")

    with c_mode:
        period_mode = st.selectbox(
            "Period",
            ["YTD → month", "Single month (MTD)", "Rolling 12 months", "Custom range"],
            index=0,
            key="inv_period_mode"
        )

    avail_years = sorted(df_tx[col_date].dropna().dt.year.unique().tolist()) if col_date else [date.today().year]
    with c_year:
        year_pick = st.selectbox("Year", avail_years[::-1], index=0, key="inv_pick_year")
    with c_month:
        month_pick = st.selectbox("Month", list(range(1, 13)), index=min(date.today().month-1, 11), key="inv_pick_month")

    # Custom range control (defaults to main sidebar dates)
    custom_start, custom_end = None, None
    if period_mode == "Custom range":
        c_a, c_b = st.columns(2)
        with c_a:
            custom_start = st.date_input("Start", value=start_date, key="inv_custom_start")
        with c_b:
            custom_end = st.date_input("End", value=end_date, key="inv_custom_end")

    # Rollup toggle + editor (still available, but applies to main-page locations)
    use_rollup = st.toggle("Roll up sub-locations", value=True, key="inv_rollup_on")
    with st.expander("⚙️ Location rollup map (save once, reuses next time)"):
        st.caption("Keys are parent inventory labels from your dropdown; values are the list of numeric LOCATION codes that roll up under the parent.")
        st.json(ROLLUP, expanded=False)
        if st.button("💾 Save rollup map", use_container_width=True, key="save_rollup_btn"):
            _save_rollup(ROLLUP)
            st.success("Saved location_rollup.json")


    # ---------- KPI Help page ----------
    def _help_html(now_ts: str) -> str:
        def row(title, what, calc, target, improve, sustain):
            return f"""
            <section style='margin:18px 0;'>
              <h3 style='margin:0 0 6px 0;'>{title}</h3>
              <p><b>What it is:</b> {what}</p>
              <p><b>How we calculate:</b> {calc}</p>
              <p><b>Target:</b> {target}</p>
              <p><b>To improve:</b> {improve}</p>
              <p><b>To sustain:</b> {sustain}</p>
              <hr style='border:none;border-top:1px solid #e5e7eb;margin:16px 0;'/>
            </section>
            """
        css = """
        <style>
          body { font-family: system-ui,-apple-system,Segoe UI,Roboto,Helvetica,Arial; margin: 20px; color:#0f172a; }
          h1 { font-size: 22px; margin: 0 0 8px 0; }
          h2 { font-size: 18px; margin: 22px 0 8px 0; }
          h3 { font-size: 16px; }
          code { background:#f1f5f9; padding:2px 6px; border-radius:6px; }
          .note { background:#f8fafc; border:1px solid #e5e7eb; padding:10px 12px; border-radius:10px; }
          @media print { .noprint { display:none; } }
          a.button { text-decoration:none; border:1px solid #e5e7eb; padding:8px 10px; border-radius:10px; background:#f8fafc; color:#0f172a; }
        </style>
        """
        hdr = f"<h1>Inventory KPI Help</h1><div class='note'>Generated {now_ts}</div>"
        why_over_100 = """
        <div class='note'>
          <b>Why can NET Coverage exceed 100%?</b><br/>
          NET uses <i>In − Out</i> as the denominator. If outs and returns straddle month boundaries,
          net can be small while PO/WO-backed flows are sizable, pushing NET Coverage over 100%.
          Use <b>ABS Coverage</b> to verify true compliance (it’s capped at 0–100%).
        </div>
        """
        sections = []
        sections.append(row(
            "YTD Inventory Growth %",
            "How much total on-hand inventory value has grown (or shrunk) from the period start to now.",
            "<code>(End value ÷ Start value − 1) × 100%</code>",
            "Green band 0–10% (watch −5–20%).",
            "Tune reorder points to usage; standardize parts; reduce lead times/MOQs; tighter PO approvals.",
            "Monthly dead-stock review; quarterly lead-time/MOQ review."
        ))
        sections.append(row(
            "Avg MoM Δ %",
            "Average of month-to-month % changes in inventory value—stability measure.",
            "Average of <code>%Δ vs prior month</code> across months in view.",
            "≤ 2% (watch ≤ 4%).",
            "Split large POs; align receipts to actual dates; plan buys to PM schedules.",
            "Keep PM calendar accurate; revisit buffer stock quarterly."
        ))
        sections.append(row(
            "PO/WO Coverage (NET %)",
            "Share of NET movement explained by PO/WO-backed transactions (incl. REOPENED/COMPLETED).",
            "<code>NET Coverage = (NET PO/WO) ÷ (NET Direction) × 100%</code> (NET=In−Out).",
            "≥ 95% (watch ≥ 90%).",
            "Align timing (issue/return in same period); enforce WO/PO on all issues/receipts; restrict manual updates.",
            "Weekly exception list for CREATED/MANUAL/blank reasons."
        ))
        sections.append(row(
            "PO/WO Coverage (ABS %)",
            "Compliance on an absolute basis so In/Out don’t cancel.",
            "<code>ABS Coverage = (abs(PO/WO IN)+abs(PO/WO OUT)) ÷ (abs(all IN)+abs(all OUT)) × 100%</code>",
            "≥ 95% (watch ≥ 90%).",
            "Train reason codes; lock down (blank)/manual updates; barcode/QR at issue/return.",
            "Monthly audit of top 20 items by spend for 100% PO/WO linkage."
        ))
        sections.append(row(
            "Adj/Unmatched (NET %)",
            "Share of NET movement unexplained by PO/WO (adjustments/unmatched).",
            "<code>Adj/Unmatched (NET %) = Σ(NET gap) ÷ Σ(NET direction)</code>",
            "≤ 15% (watch ≤ 20%).",
            "Limit back-dating; ensure returns (REOPENED) occur same week; coach users with frequent MANUAL_UPDATE.",
            "15-min monthly variance review covering largest gaps by item & user."
        ))
        sections.append(row(
            "Avg Error Rate %",
            "Average of monthly error rates normalized to inventory value.",
            "<code>Monthly Error % = NET gap ÷ inventory value</code>; averaged across months.",
            "≤ 0.50% (stretch 0.25%; watch ≤ 1.00%).",
            "Dual-verification for high-value items; require WO closure notes where parts>$X.",
            "Keep dual-check on top 10% by cost; weekly spot audit."
        ))
        sections.append(row(
            "Reopen % of WO activity (ABS)",
            "Of all PO/WO activity (abs $), what fraction are WORK_ORDER_REOPENED returns.",
            "<code>abs($ REOPENED IN) ÷ abs($ PO/WO IN + OUT) × 100%</code>",
            "≤ 5% (watch ≤ 10%).",
            "Pre-issue checklist (asset, qty, interchange); stage kits close to start date; track returns by user/asset.",
            "Maintain kit-staging; require reason codes on every return."
        ))
        html = f"<!doctype html><html><head><meta charset='utf-8'>{css}</head><body>{hdr}{why_over_100}{''.join(sections)}<div class='noprint' style='margin-top:16px;'><a class='button' href='javascript:window.print()'>🖨️ Print</a></div></body></html>"
        return html

    c_help1, c_help2 = st.columns([1,1])
    with c_help1:
        if st.button("❓ Open KPI Help (new tab)", use_container_width=True):
            html = _help_html(datetime.now().strftime("%Y-%m-%d %H:%M"))
            data_url = "data:text/html;charset=utf-8," + quote(html)
            components.html(f"<script>window.open('{data_url}','_blank');</script>", height=0)
    with c_help2:
        st.download_button(
            "⬇️ Download KPI Help (HTML)",
            data=_help_html(datetime.now().strftime("%Y-%m-%d %H:%M")).encode("utf-8"),
            file_name="Inventory_KPI_Help.html",
            mime="text/html",
            use_container_width=True
        )

    # ---------- DATE WINDOW & MONTH CUTS ----------
    def _month_start(y, m): return pd.Timestamp(f"{y}-{m:02d}-01 00:00:00")
    def _next_month(y, m):
        ny, nm = (y+1, 1) if m == 12 else (y, m+1)
        return pd.Timestamp(f"{ny}-{nm:02d}-01 00:00:00")

    if period_mode == "YTD → month":
        window_start = _month_start(year_pick, 1)
        window_end   = _next_month(year_pick, month_pick)
        baseline_ts  = window_start
    elif period_mode == "Single month (MTD)":
        window_start = _month_start(year_pick, month_pick)
        window_end   = _next_month(year_pick, month_pick)
        baseline_ts  = window_start
    elif period_mode == "Rolling 12 months":
        window_end   = _next_month(year_pick, month_pick)
        start_dt     = window_end - pd.DateOffset(months=12)
        window_start = _month_start(start_dt.year, start_dt.month)
        baseline_ts  = window_start
    else:  # Custom range
        cs = pd.Timestamp(custom_start) if custom_start else pd.Timestamp(start_date)
        ce = pd.Timestamp(custom_end) + pd.Timedelta(days=1) if custom_end else (pd.Timestamp(end_date) + pd.Timedelta(days=1))
        window_start, window_end = cs.normalize(), ce.normalize()
        baseline_ts  = window_start

    # Cutoffs at each month boundary in the window
    cutoffs = []
    cursor = pd.Timestamp(_month_start(window_start.year, window_start.month))
    while cursor < window_end:
        nxt = _next_month(cursor.year, cursor.month)
        label = f"{calendar.month_abbr[cursor.month]} {cursor.year}"
        cutoffs.append((label, nxt))
        cursor = nxt

    # ---------- rollup-aware filtering ----------
    def _parts_for_location(dfv: pd.DataFrame, loc_labels) -> pd.DataFrame:
        """Filter parts by one or more parent locations from main page."""
        if not col_loc_parts_candidates or not loc_labels:
            return dfv
        # If any 'ALL' selection sneaks through, treat as no filter
        if any(str(l).strip() in admin_tokens or str(l).strip() == ALL for l in loc_labels):
            return dfv

        if not use_rollup:
            m = pd.Series(False, index=dfv.index)
            for c in col_loc_parts_candidates:
                m |= dfv[c].astype(str).isin(loc_labels)
            return dfv[m]

        # Rollup mode — convert parent labels to numeric codes
        target_codes: set[str] = set()
        for label in loc_labels:
            label = str(label)
            codes = _codes_for_parent(label)
            if codes:
                target_codes.update(codes)
            else:
                code = _extract_code(label)
                if code:
                    target_codes.add(code)

        if not target_codes:
            return dfv

        m = pd.Series(False, index=dfv.index)
        for c in col_loc_parts_candidates:
            m |= dfv[c].map(_extract_code).isin(target_codes)
        return dfv[m]

    def _tx_for_location(dfv: pd.DataFrame, loc_labels) -> pd.DataFrame:
        """Filter transactions by one or more parent locations from main page."""
        if not col_loc or not loc_labels:
            return dfv
        if any(str(l).strip() in admin_tokens or str(l).strip() == ALL for l in loc_labels):
            return dfv

        if not use_rollup:
            return dfv[dfv[col_loc].astype(str).isin(loc_labels)]

        target_codes: set[str] = set()
        for label in loc_labels:
            label = str(label)
            codes = _codes_for_parent(label)
            if codes:
                target_codes.update(codes)
            else:
                code = _extract_code(label)
                if code:
                    target_codes.add(code)

        if not target_codes:
            return dfv

        return dfv[dfv[col_loc].map(_extract_code).isin(target_codes)]


    # ---------- apply location filters ----------
    parts_loc = _parts_for_location(df_parts, eff_locations)
    created_loc = created_ts_all.loc[parts_loc.index]
    updated_loc = updated_ts_all.loc[parts_loc.index]
    value_loc   = parts_value_all.loc[parts_loc.index]

    def _inv_as_of(cutoff: pd.Timestamp) -> float:
        mask = created_loc.lt(cutoff)
        if col_updated:
            mask |= updated_loc.lt(cutoff)
        return float(value_loc.where(mask, 0.0).sum())

    inv_baseline = _inv_as_of(baseline_ts)

    tx_loc = _tx_for_location(df_tx, eff_locations)

    if col_date:
        tx_loc = tx_loc[(tx_loc[col_date] >= window_start) & (tx_loc[col_date] < window_end)]

    if col_dir:
        tx_dir = tx_loc[col_dir].astype(str).str.lower()
        is_dir_in_loc  = tx_dir.str.startswith("in")
        is_dir_out_loc = tx_dir.str.startswith("out")
    else:
        is_dir_in_loc  = pd.Series(False, index=tx_loc.index)
        is_dir_out_loc = pd.Series(False, index=tx_loc.index)

    is_tin_loc  = _truthy_mask(tx_loc[col_tin]  if (col_tin  in tx_loc.columns) else None, tx_loc.index)
    is_tout_loc = _truthy_mask(tx_loc[col_tout] if (col_tout in tx_loc.columns) else None, tx_loc.index)

    # ---------- OVERRIDE: map REOPENED/COMPLETED into PO/WO coverage ----------
    if col_reason and col_dir:
        reas_up = tx_loc[col_reason].astype(str).str.upper()
        is_tin_loc  = is_tin_loc  | (reas_up.eq("WORK_ORDER_REOPENED") & is_dir_in_loc)
        is_tout_loc = is_tout_loc | (reas_up.eq("WORK_ORDER_COMPLETED") & is_dir_out_loc)

    PO_WO_REASONS = {"PURCHASE_ORDER_FULFILLED","WORK_ORDER_COMPLETE","WORK_ORDER_REOPENED","RESTOCK"}

    # ---------- month-by-month summary ----------
    summary_rows = []
    prev_val = inv_baseline

    for label, cutoff in cutoffs:
        v = _inv_as_of(cutoff)
        pct_vs_prior = ((v - prev_val) / prev_val * 100.0) if (pd.notna(prev_val) and prev_val != 0) else np.nan

        # mask for that month
        month_start = pd.Timestamp(f"{label.split()[1]}-{list(calendar.month_abbr).index(label.split()[0])}-01")
        month_mask = (tx_loc[col_date] >= month_start) & (tx_loc[col_date] < cutoff)

        in_cost  = float(_num(tx_loc.loc[month_mask & is_dir_in_loc,  col_tc]).sum())
        out_cost = float(_num(tx_loc.loc[month_mask & is_dir_out_loc, col_tc]).sum())
        out_abs  = abs(out_cost)
        net_io   = in_cost - out_abs

        denom_inv = v if (pd.notna(v) and v != 0) else np.nan
        pct_io    = (net_io / denom_inv * 100.0) if pd.notna(denom_inv) else np.nan

        tin_cost  = float(_num(tx_loc.loc[month_mask & is_tin_loc,  col_tc]).sum())
        tout_cost = float(_num(tx_loc.loc[month_mask & is_tout_loc, col_tc]).sum())
        net_tr    = tin_cost - abs(tout_cost)
        pct_tr    = (net_tr / denom_inv * 100.0) if pd.notna(denom_inv) else np.nan

        gap   = net_io - net_tr
        err_r = (gap / denom_inv * 100.0) if pd.notna(denom_inv) else np.nan

        summary_rows.append({
            "Period": label,
            "Inventory Value": v,
            "%Δ Inventory vs prior": pct_vs_prior,
            "In $ (DIRECTION)": in_cost,
            "Out $ (DIRECTION)": out_abs,
            "Net (In−Out) $ (DIRECTION)": net_io,
            "% Impact on Inv (DIR)": pct_io,
            "trans_in $": tin_cost,
            "trans_out $": abs(tout_cost),
            "Net (trans_in−out) $": net_tr,
            "% Impact on Inv (trans)": pct_tr,
            "Gap (Net IO − Net Trans) $": gap,
            "Error Rate vs Inv Value %": err_r,
        })
        prev_val = v

    out = pd.DataFrame(summary_rows)

    # ====================
    # KPI calculations
    # ====================
    pct_total = avg_mom = coverage_net = gap_share_net = avg_err = None
    coverage_abs_pct = reopen_pct_of_wo = None
    start_val = end_val = None

    if not out.empty:
        for c in out.columns:
            if c != "Period":
                out[c] = pd.to_numeric(out[c], errors="coerce")
        start_val = float(out.iloc[0]["Inventory Value"]) if pd.notna(out.iloc[0]["Inventory Value"]) else np.nan
        end_val   = float(out.iloc[-1]["Inventory Value"]) if pd.notna(out.iloc[-1]["Inventory Value"]) else np.nan
        pct_total = ((end_val / start_val - 1) * 100.0) if (pd.notna(start_val) and start_val != 0) else np.nan
        avg_mom   = float(out["%Δ Inventory vs prior"].mean())

        net_dir   = float(out["Net (In−Out) $ (DIRECTION)"].sum())
        net_trans = float(out["Net (trans_in−out) $"].sum())
        gap_sum   = float(out["Gap (Net IO − Net Trans) $"].sum())
        coverage_net  = (net_trans / net_dir * 100.0) if (net_dir not in (0, np.nan)) else np.nan
        gap_share_net = (gap_sum / net_dir * 100.0) if (net_dir not in (0, np.nan)) else np.nan
        avg_err   = float(out["Error Rate vs Inv Value %"].mean())
        
        # --- Store Inventory KPIs & Monthly rollup for PDF report ---
    try:
        st.session_state["hub_inv_monthly_rollup"] = out.copy()

        st.session_state["hub_inv_kpis"] = {
            "pct_total": pct_total,
            "avg_mom": avg_mom,
            "coverage_net": coverage_net,
            "coverage_abs_pct": coverage_abs_pct,
            "gap_share_net": gap_share_net,
            "avg_err": avg_err,
            "reopen_pct_of_wo": reopen_pct_of_wo,
            "start_val": start_val,
            "end_val": end_val,
        }

        # Also store the context of the inventory filter used
        st.session_state["hub_inv_location"] = pick_loc
        st.session_state["hub_inv_period_mode"] = period_mode
    except Exception:
        # Don't let session_state issues break the page
        pass


    # =========================================================
    # Transaction Reason distribution + ABS coverage + Reopen%
    # =========================================================
    st.subheader("Transaction Reasons — distribution")
    reasons_tbl = pd.DataFrame()
    if col_reason and not tx_loc.empty:
        tx_view = tx_loc.copy()
        tx_view[col_reason] = tx_view[col_reason].fillna("")
        grp = tx_view.groupby(col_reason, dropna=False).agg(
            Count=(col_reason, "size"),
            Amount_Sum=(col_tc, "sum")
        ).reset_index().rename(columns={col_reason: "TRANSACTION REASON"})
        tot_c = grp["Count"].sum()
        tot_a = grp["Amount_Sum"].abs().sum()
        grp["% by Count"]  = (grp["Count"] / tot_c * 100.0) if tot_c else 0.0
        grp["% by Amount"] = (grp["Amount_Sum"].abs() / tot_a * 100.0) if tot_a else 0.0
        grp["PO/WO Related"] = grp["TRANSACTION REASON"].str.upper().isin({"PURCHASE_ORDER_FULFILLED","WORK_ORDER_COMPLETE","WORK_ORDER_REOPENED","RESTOCK"})
        grp["Non-PO/WO"] = ~grp["PO/WO Related"]
        reasons_tbl = grp.sort_values(by="Count", ascending=False)

        is_in_abs  = tx_view[col_dir].astype(str).str.lower().str.startswith("in")
        is_out_abs = tx_view[col_dir].astype(str).str.lower().str.startswith("out")
        abs_all = tx_view.loc[is_in_abs, col_tc].abs().sum() + tx_view.loc[is_out_abs, col_tc].abs().sum()
        wo_related_mask = tx_view[col_reason].str.upper().isin({"PURCHASE_ORDER_FULFILLED","WORK_ORDER_COMPLETE","WORK_ORDER_REOPENED","RESTOCK"})
        abs_wo = tx_view.loc[wo_related_mask & is_in_abs, col_tc].abs().sum() + tx_view.loc[wo_related_mask & is_out_abs, col_tc].abs().sum()
        coverage_abs_pct = (abs_wo / abs_all * 100.0) if abs_all else np.nan

        abs_reopen = tx_view.loc[tx_view[col_reason].str.upper().eq("WORK_ORDER_REOPENED") & is_in_abs, col_tc].abs().sum()
        reopen_pct_of_wo = (abs_reopen / abs_wo * 100.0) if abs_wo else np.nan

        pct_non_count = reasons_tbl.loc[reasons_tbl["Non-PO/WO"], "Count"].sum() / tot_c * 100.0 if tot_c else 0.0
        pct_non_amt   = reasons_tbl.loc[reasons_tbl["Non-PO/WO"], "Amount_Sum"].abs().sum() / tot_a * 100.0 if tot_a else 0.0

        st.caption(
            f"Non-PO/WO share — by count: **{pct_non_count:.2f}%**, by amount: **{pct_non_amt:.2f}%**  "
            f"• ABS PO/WO Coverage: **{(0 if coverage_abs_pct is None or np.isnan(coverage_abs_pct) else coverage_abs_pct):.2f}%**  "
            f"• Reopen % of WO activity (ABS): **{(0 if reopen_pct_of_wo is None or np.isnan(reopen_pct_of_wo) else reopen_pct_of_wo):.2f}%**"
        )
        st.dataframe(
            reasons_tbl[["TRANSACTION REASON","Count","% by Count","Amount_Sum","% by Amount"]]
              .rename(columns={"Amount_Sum":"Sum $"})
              .round(2),
            use_container_width=True
        )
    else:
        st.info("No TRANSACTION REASON column found; skipping reason distribution.")

    # =========================
    # KPI strip (large tiles)
    # =========================
    st.subheader("KPI Overview")
    cA, cB, cC, cD, cE, cF = st.columns(6)
    with cA: kpi_card(KPI_T["ytd_growth_pct"]["label"], pct_total, "%", "ytd_growth_pct")
    with cB: kpi_card(KPI_T["avg_mom_pct"]["label"],   avg_mom,   "%", "avg_mom_pct")
    with cC: kpi_card(KPI_T["coverage_net_pct"]["label"], coverage_net, "%", "coverage_net_pct")
    with cD: kpi_card(KPI_T["coverage_abs_pct"]["label"], coverage_abs_pct, "%", "coverage_abs_pct")
    with cE: kpi_card(KPI_T["gap_share_net_pct"]["label"], gap_share_net, "%", "gap_share_net_pct")
    with cF: kpi_card(KPI_T["avg_error_pct"]["label"], avg_err, "%", "avg_error_pct")

    cG, _ = st.columns([1,3])
    with cG: kpi_card(KPI_T["reopen_pct_of_wo"]["label"], reopen_pct_of_wo, "%", "reopen_pct_of_wo")

    # ===========================
    # WO lifecycle-normalized view
    # ===========================
    with st.expander("WO Normalized Consumption (Net_WO$ & Abs_WO$ per month)"):
        if col_wo and col_part and col_reason and col_dir and col_tc and col_date:
            txn = tx_loc.copy()
            txn[col_date] = pd.to_datetime(txn[col_date], errors="coerce")
            txn["month"] = txn[col_date].dt.to_period("M").dt.to_timestamp()
            is_reopen   = txn[col_reason].astype(str).str.upper().eq("WORK_ORDER_REOPENED")
            is_complete = txn[col_reason].astype(str).str.upper().eq("WORK_ORDER_COMPLETED")
            is_in_m     = txn[col_dir].astype(str).str.lower().str.startswith("in")
            is_out_m    = txn[col_dir].astype(str).str.lower().str.startswith("out")
            grp_keys = [col_wo, col_part, "month"]
            agg = txn.groupby(grp_keys, dropna=False).apply(
                lambda g: pd.Series({
                    "Issue_OUT$": abs(g.loc[is_complete & is_out_m, col_tc]).sum(),
                    "Return_IN$": abs(g.loc[is_reopen  & is_in_m,  col_tc]).sum()
                })
            ).reset_index()
            if not agg.empty:
                agg["Net_WO$"] = agg["Issue_OUT$"] - agg["Return_IN$"]
                agg["Abs_WO$"] = agg["Issue_OUT$"] + agg["Return_IN$"]
                st.dataframe(agg.sort_values(["month", col_wo, "Abs_WO$"], ascending=[True, True, False]).round(2),
                             use_container_width=True)
            else:
                st.info("No WO reopen/complete activity in current filter.")
        else:
            st.info("Missing columns for WO normalization (need WO, PART, DATE, DIRECTION, TRANSACTION REASON, TOTAL COST).")

    # ===========================
    # Monthly roll-up (pretty)
    # ===========================
    show = out.copy()
    money_cols = [
        "Inventory Value","In $ (DIRECTION)","Out $ (DIRECTION)","Net (In−Out) $ (DIRECTION)",
        "trans_in $","trans_out $","Net (trans_in−out) $","Gap (Net IO − Net Trans) $"
    ]
    pct_cols = ["%Δ Inventory vs prior","% Impact on Inv (DIR)","% Impact on Inv (trans)","Error Rate vs Inv Value %"]
    for c in money_cols:
        show[c] = show[c].map(lambda v: "" if pd.isna(v) else f"${v:,.2f}")
    for c in pct_cols:
        show[c] = show[c].map(lambda v: "" if pd.isna(v) else f"{v:,.2f}%")
    st.subheader("Monthly Roll-up")
    st.dataframe(show, use_container_width=True, height=520)

    # ---------- XLSX export ----------
    if "to_xlsx_bytes" not in globals():
        from io import BytesIO
        def to_xlsx_bytes(df, sheet_name=None, sheet=None):
            name = sheet or sheet_name or "Sheet1"
            buf = BytesIO()
            with pd.ExcelWriter(buf, engine="xlsxwriter") as xw:
                df.to_excel(xw, index=False, sheet_name=name)
            return buf.getvalue()
        globals()["to_xlsx_bytes"] = to_xlsx_bytes

    period_tag = period_mode.replace(" ", "_").replace("→","to")
    safe_loc_tag = loc_label_for_display.replace(" ", "_").replace(",", "_")
    st.download_button(
        "⬇️ Download Inventory Analysis (XLSX)",
        data=to_xlsx_bytes(out, sheet_name=f"Inv_Analysis_{period_tag}"),
        file_name=f"Inventory_Analysis_{safe_loc_tag}_{period_tag}.xlsx",
    )


    # ---------- Word export ----------
    DOCX_AVAILABLE = False
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        DOCX_AVAILABLE = True
    except Exception:
        DOCX_AVAILABLE = False

    if DOCX_AVAILABLE and not out.empty:
        outliers = out.loc[out["Error Rate vs Inv Value %"].abs() >= 1,
                           ["Period","Error Rate vs Inv Value %","Gap (Net IO − Net Trans) $"]] \
                    .sort_values(by="Gap (Net IO − Net Trans) $", key=lambda s: s.abs(), ascending=False)
        doc = Document()
        ttl = doc.add_paragraph()
        rn = ttl.add_run(f"Inventory KPI Report — {loc_label_for_display} — {period_mode}")



        doc.add_paragraph(
            f"PO/WO coverage (NET): {0 if coverage_net is None or np.isnan(coverage_net) else coverage_net:.2f}%. "
            f"PO/WO coverage (ABS): {0 if coverage_abs_pct is None or np.isnan(coverage_abs_pct) else coverage_abs_pct:.2f}%. "
            f"Adjustments/unmatched (NET): {0 if gap_share_net is None or np.isnan(gap_share_net) else gap_share_net:.2f}%. "
            f"Average error rate: {0 if avg_err is None or np.isnan(avg_err) else avg_err:.2f}%."
        )
        if (reopen_pct_of_wo is not None) and not np.isnan(reopen_pct_of_wo):
            doc.add_paragraph(f"Reopen % of WO activity (ABS): {reopen_pct_of_wo:.2f}%.")

        t = doc.add_table(rows=1, cols=4); h = t.rows[0].cells
        h[0].text, h[1].text, h[2].text, h[3].text = "KPI", "Current", "Target", "Flag"
        rows = [
            ("YTD / Period Growth %", f"{0 if pct_total is None or np.isnan(pct_total) else pct_total:.2f}%", "0–10% (plan)", "Red if < -5% or > 20%"),
            ("Avg MoM Δ %", f"{0 if avg_mom is None or np.isnan(avg_mom) else avg_mom:.2f}%", "≤ 2%", "Red if > 4%"),
            ("PO/WO Coverage (NET %)", f"{0 if coverage_net is None or np.isnan(coverage_net) else coverage_net:.2f}%", "≥ 95%", "Red if < 90%"),
            ("PO/WO Coverage (ABS %)", f"{0 if coverage_abs_pct is None or np.isnan(coverage_abs_pct) else coverage_abs_pct:.2f}%", "≥ 95%", "Red if < 90%"),
            ("Adj/Unmatched (NET %)", f"{0 if gap_share_net is None or np.isnan(gap_share_net) else gap_share_net:.2f}%", "≤ 15%", "Red if > 20%"),
            ("Avg Error Rate %", f"{0 if avg_err is None or np.isnan(avg_err) else avg_err:.2f}%", "≤ 0.50% (stretch 0.25%)", "Red if ≥ 1.00%"),
        ]
        for k, cur, tgt, flg in rows:
            r = t.add_row().cells; r[0].text, r[1].text, r[2].text, r[3].text = k, cur, tgt, flg
        doc.add_paragraph(""); doc.add_paragraph("Outlier Months (Error Rate ≥ 1.00%):")
        if not outliers.empty:
            t2 = doc.add_table(rows=1, cols=3); hh = t2.rows[0].cells
            hh[0].text, hh[1].text, hh[2].text = "Period", "Error Rate %", "Gap $"
            for _, rr in outliers.iterrows():
                rw = t2.add_row().cells
                rw[0].text = str(rr["Period"])
                rw[1].text = f"{float(rr['Error Rate vs Inv Value %']):.3f}%"
                rw[2].text = f"${float(rr['Gap (Net IO − Net Trans) $']):,.2f}"
        else:
            doc.add_paragraph("None.")
        if not reasons_tbl.empty:
            doc.add_paragraph(""); doc.add_paragraph("Transaction Reasons (current filter):")
            t3 = doc.add_table(rows=1, cols=5); hh = t3.rows[0].cells
            hh[0].text, hh[1].text, hh[2].text, hh[3].text, hh[4].text = "Reason","Count","% by Count","Sum $","% by Amount"
            for _, rr in reasons_tbl.iterrows():
                rw = t3.add_row().cells
                rw[0].text = str(rr["TRANSACTION REASON"] or "(blank)")
                rw[1].text = f"{int(rr['Count'])}"
                rw[2].text = f"{float(rr['% by Count']):.2f}%"
                rw[3].text = f"${float(rr['Sum $'] if 'Sum $' in reasons_tbl.columns else rr['Amount_Sum']):,.2f}"
                rw[4].text = f"{float(rr['% by Amount']):.2f}%"
        bio = io.BytesIO(); doc.save(bio)
        # --- Word download (robust) ---
        import io  # ensure available in this scope

        def _docx_to_bytes(_doc):
            """Return docx bytes safely."""
            _bio = io.BytesIO()
            _doc.save(_bio)
            _bio.seek(0)
            return _bio.getvalue()

        try:
            docx_bytes = _docx_to_bytes(doc)
            safe_loc = (pick_loc or "All").replace(" ", "_")
            st.download_button(
                "Download KPI Word Report (current filter)",
                data=docx_bytes,
                file_name=f"Inventory_KPI_Report_{safe_loc}_{period_tag}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key=f"inv_kpi_docx_{safe_loc}_{period_tag}",
            )
        except Exception as e:
            st.error(f"Word export failed: {type(e).__name__}: {e}")

    elif not DOCX_AVAILABLE:
        st.caption("Word export unavailable (python-docx not installed).")

    # ---------- KPI thresholds editor ----------
    with st.expander("KPI Threshold Settings — adjust & Save as defaults"):
        st.caption("Green = good, Yellow = watch, Red = fail. Saved to inv_kpi_settings.json.")
        def num(key, subkey, step=0.1):
            val = KPI_T[key].get(subkey)
            return st.number_input(f"{KPI_T[key]['label']} — {subkey}", value=float(val), step=step, key=f"{key}_{subkey}")

        c1, c2, c3, c4 = st.columns(4)
        with c1: KPI_T["ytd_growth_pct"]["green_min"] = num("ytd_growth_pct","green_min",0.5)
        with c2: KPI_T["ytd_growth_pct"]["green_max"] = num("ytd_growth_pct","green_max",0.5)
        with c3: KPI_T["ytd_growth_pct"]["warn_min"]  = num("ytd_growth_pct","warn_min",0.5)
        with c4: KPI_T["ytd_growth_pct"]["warn_max"]  = num("ytd_growth_pct","warn_max",0.5)

        c1, c2 = st.columns(2)
        with c1:
            KPI_T["avg_mom_pct"]["goal"] = num("avg_mom_pct","goal",0.1)
            KPI_T["gap_share_net_pct"]["goal"] = num("gap_share_net_pct","goal",0.5)
            KPI_T["avg_error_pct"]["goal"] = num("avg_error_pct","goal",0.05)
            KPI_T["reopen_pct_of_wo"]["goal"] = num("reopen_pct_of_wo","goal",0.5)
        with c2:
            KPI_T["avg_mom_pct"]["warn_max"] = num("avg_mom_pct","warn_max",0.1)
            KPI_T["gap_share_net_pct"]["warn_max"] = num("gap_share_net_pct","warn_max",0.5)
            KPI_T["avg_error_pct"]["warn_max"] = num("avg_error_pct","warn_max",0.1)
            KPI_T["reopen_pct_of_wo"]["warn_max"] = num("reopen_pct_of_wo","warn_max",0.5)

        c1, c2 = st.columns(2)
        with c1:
            KPI_T["coverage_net_pct"]["goal"] = num("coverage_net_pct","goal",0.5)
            KPI_T["coverage_abs_pct"]["goal"] = num("coverage_abs_pct","goal",0.5)
        with c2:
            KPI_T["coverage_net_pct"]["warn_min"] = num("coverage_net_pct","warn_min",0.5)
            KPI_T["coverage_abs_pct"]["warn_min"] = num("coverage_abs_pct","warn_min",0.5)

        if st.button("💾 Save as defaults", use_container_width=True):
            st.success("Saved KPI defaults.") if save_thresholds(KPI_T) else st.error("Failed to save KPI defaults.")
            
            
def render_expected_service(
    df_expected: pd.DataFrame,
    selected_locations: list[str] | None = None,
    admin_all_tokens=("« All locations »", "« All companies (admin) »", "(all)", "All", "ALL"),
):
    """
    Expected Service page for Reporting Hub.

    - Uses df_expected passed in (Expected.parquet already loaded by the hub)
    - Applies the main location filter (selected_locations from the sidebar)
    - Breaks output into separate sections instead of tabs:
        • All
        • Needs Service
        • Expected  (sorted by DueDate: soonest → farthest)
        • New Reading Needed
        • Overdue   (sorted by DueDate: oldest → newest)
        • Due
    - "Next Service" is always shown as plain text (no date coercion)

    Extra section (bottom of page):
        • Loads WO-DB.parquet
        • Applies same main location filter (Location2)
        • Filters to Procedure Template ID = 2132428 AND IsOpen = True
        • Splits into:
            A) Without Start/Due date (table used in PDF)
            B) With Start or Due date
        • Builds 7-day and 30-day calendars (Expected DueDate + Open Proc WOs)
          with separate PDF exports.

    PDF content:
        Page 1: 7/30-day calendars
        Page 2+: Tables
            - Open WOs (Proc 2132428) — No Start/Due Date
            - Needs Service       (Name, Last Service Type, Date of Last service)
            - New Reading Needed  (Name, Last Reading, Date)
            - Overdue             (Name, Next Service Type, Remaining)
            - Due                 (Name, Next Service Type, Remaining)
    """
    from io import BytesIO
    from pathlib import Path
    from datetime import date as _date, timedelta
    import math
    import html

    import pandas as pd
    import streamlit as st

    # ---- tiny_xlsx_button shim (shared with other pages) ----
    if "tiny_xlsx_button" not in globals():
        def tiny_xlsx_button(df: pd.DataFrame, sheet: str = "Sheet1",
                             filename: str = "export.xlsx", key_suffix: str = "dl"):
            buf = BytesIO()
            with pd.ExcelWriter(buf, engine="xlsxwriter") as xw:
                (df if isinstance(df, pd.DataFrame) else pd.DataFrame(df)).to_excel(
                    xw, index=False, sheet_name=sheet
                )
            st.download_button(
                "⬇️ Download (XLSX)",
                data=buf.getvalue(),
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key=f"dl_{key_suffix}"
            )
        globals()["tiny_xlsx_button"] = tiny_xlsx_button
    else:
        tiny_xlsx_button = globals()["tiny_xlsx_button"]

    st.markdown("### 🛠 Expected Service")

    if df_expected is None or df_expected.empty:
        st.warning("The 'Expected' dataset is empty.")
        return

    # ---------- small helper: pick column ----------
    def _pick_col(df: pd.DataFrame, candidates: list[str]) -> str | None:
        """Return the first matching column from candidates (case-insensitive)."""
        if df is None or df.empty:
            return None
        lowmap = {str(c).lower(): c for c in df.columns}
        for c in candidates:
            if c in df.columns:
                return c
            lc = str(c).lower()
            if lc in lowmap:
                return lowmap[lc]
        return None

    # ---------- helper: sort by DueDate (ascending, NaNs last) ----------
    def _sort_by_due(df_in: pd.DataFrame) -> pd.DataFrame:
        if df_in.empty or "DueDate" not in df_in.columns:
            return df_in
        s = pd.to_datetime(df_in["DueDate"], errors="coerce")
        df_tmp = df_in.copy()
        df_tmp["_DueSort"] = s
        df_tmp = df_tmp.sort_values("_DueSort", ascending=True, na_position="last")
        return df_tmp.drop(columns=["_DueSort"])

    def _safe_str_cell(val) -> str:
        """Return '' for NA/NaN/None, else string, without boolean 'or' on NA."""
        if val is None:
            return ""
        try:
            if pd.isna(val):
                return ""
        except Exception:
            pass
        return str(val)

    base = df_expected.copy()

    # ---------- Apply main location filter to Expected ----------
    loc_col = _pick_col(base, ["Location", "Location2", "Company", "Site"])

    apply_filter = (
        loc_col is not None
        and selected_locations is not None
        and len(selected_locations) > 0
    )

    if apply_filter:
        # Normalize both sides and ignore any explicit "ALL" tokens
        allowed = {
            s.strip().casefold()
            for s in selected_locations
            if isinstance(s, str) and s.strip() and s.strip() not in admin_all_tokens
        }

        if allowed:
            loc_norm = base[loc_col].astype(str).str.strip().str.casefold()
            mask = loc_norm.isin(allowed)
            filtered = base[mask].copy()
            if filtered.empty:
                st.warning(
                    "No rows after applying the main location filter to Expected.parquet.\n\n"
                    "Showing all rows instead (location labels may not match 1:1)."
                )
                df = base
            else:
                df = filtered
        else:
            # All tokens / admin case => no filter
            df = base
    else:
        df = base

    # Normalize strings (strip)
    for c in df.columns:
        if df[c].dtype == object:
            df[c] = df[c].astype("string").str.strip()

    # ---------- Column mapping (source → output headers) ----------
    # Prefer explicit display column for Next Service if present.
    next_service_display_src = _pick_col(df, ["Next Service (display)", "NextServiceDisplay", "Next Service Display"])
    next_service_src_fallback = _pick_col(df, ["Next Service", "NextService", "Next Service At", "Next Due Reading"])

    col_map = {
        "Name":                        _pick_col(df, ["Name","Asset","Asset Name","Line Name"]),
        "Last Reading":                _pick_col(df, ["Last Reading","LastReading","Last Meter","Last Meter Reading","Last Meter Value","MReading"]),
        "Date":                        _pick_col(df, ["Date","Reading Date","Last Reading Date","LastReadDate"]),
        "Last Service Type":           _pick_col(df, ["Last Service Type","LastServiceType","Prev Service Type","Previous Service Type"]),
        "Date of Last service":        _pick_col(df, ["Date of Last service","Last Service Date","Prev Service Date","Previous Service Date"]),
        "Hours/Miles at Last Service": _pick_col(df, ["Hours/Miles at Last Service","Prev Service Reading","Previous Service Reading","Reading at Last Service"]),
        "Next Service Type":           _pick_col(df, ["Next Service Type","NextServiceType","Upcoming Service Type"]),
        # "Next Service" handled separately to force plain-text display
        "Remaining":                   _pick_col(df, ["Remaining","Remaining to Service","Remaining Reading","Remaining Hours","Remaining Miles"]),
        "Status":                      _pick_col(df, ["Status","Service Status"]),
        "DueDate":                     _pick_col(df, ["DueDate","Due Date","Expected Due Date","Next Due Date"]),
    }

    out_cols = [
        "Name","Last Reading","Date","Last Service Type","Date of Last service",
        "Hours/Miles at Last Service","Next Service Type","Next Service",
        "Remaining","Status","DueDate"
    ]

    slim = pd.DataFrame(index=df.index)

    # Copy mapped columns (except "Next Service")
    for k, src in col_map.items():
        if k == "Next Service":
            continue
        if src is not None and src in df.columns:
            slim[k] = df[src]
        else:
            slim[k] = pd.NA

    # ---------- Next Service: ALWAYS verbatim text ----------
    if next_service_display_src and next_service_display_src in df.columns:
        slim["Next Service"] = df[next_service_display_src].astype(str)
    elif next_service_src_fallback and next_service_src_fallback in df.columns:
        slim["Next Service"] = df[next_service_src_fallback].astype(str)
    else:
        slim["Next Service"] = pd.NA

    # ---------- Coerce true date columns ONLY (NOT "Next Service") ----------
    for dcol in ["Date", "Date of Last service", "DueDate"]:
        if dcol in slim.columns:
            s = pd.to_datetime(slim[dcol], errors="coerce")
            slim[dcol] = s.dt.date.where(~s.isna(), slim[dcol])

    # ---------- Status-based views (no tabs; stacked sections) ----------
    has_status = "Status" in slim.columns and slim["Status"].notna().any()

    def status_view(df_in: pd.DataFrame, label: str) -> pd.DataFrame:
        if not has_status:
            return df_in.iloc[0:0]
        return df_in[
            df_in["Status"]
            .astype(str)
            .str.strip()
            .str.casefold()
            == label.strip().casefold()
        ]

    # --- Section: All ---
    st.markdown("#### 📋 All Assets in Expected Matrix")
    view_all = slim[out_cols]
    st.caption(f"Rows: {len(view_all)}")
    st.dataframe(view_all, use_container_width=True, hide_index=True)
    tiny_xlsx_button(
        view_all,
        sheet="Expected_All",
        filename="Expected_All.xlsx",
        key_suffix="expected_all",
    )

    st.markdown("---")

    # --- Section: Needs Service ---
    st.markdown("#### 🔧 Needs Service")
    view_ns = status_view(slim, "Needs Service")[out_cols]
    st.caption(f"Rows: {len(view_ns)}")
    st.dataframe(view_ns, use_container_width=True, hide_index=True)
    if not view_ns.empty:
        tiny_xlsx_button(
            view_ns,
            sheet="Expected_NeedsService",
            filename="Expected_NeedsService.xlsx",
            key_suffix="expected_ns",
        )

    st.markdown("---")

    # --- Section: Expected (sorted soonest → farthest by DueDate) ---
    st.markdown("#### 📆 Expected")
    view_exp = status_view(slim, "Expected")[out_cols]
    view_exp = _sort_by_due(view_exp)
    st.caption(f"Rows: {len(view_exp)} (sorted by DueDate: soonest → farthest)")
    st.dataframe(view_exp, use_container_width=True, hide_index=True)
    if not view_exp.empty:
        tiny_xlsx_button(
            view_exp,
            sheet="Expected_Expected",
            filename="Expected_Expected.xlsx",
            key_suffix="expected_expected",
        )

    st.markdown("---")

    # --- Section: New Reading Needed ---
    st.markdown("#### 📝 New Reading Needed")
    view_newr = status_view(slim, "New Reading Needed")[out_cols]
    st.caption(f"Rows: {len(view_newr)}")
    st.dataframe(view_newr, use_container_width=True, hide_index=True)
    if not view_newr.empty:
        tiny_xlsx_button(
            view_newr,
            sheet="Expected_NewReadingNeeded",
            filename="Expected_NewReadingNeeded.xlsx",
            key_suffix="expected_newr",
        )

    st.markdown("---")

    # --- Section: Overdue (sorted oldest → newest by DueDate) ---
    st.markdown("#### ⏰ Overdue")
    view_over = status_view(slim, "Overdue")[out_cols]
    view_over = _sort_by_due(view_over)
    st.caption(f"Rows: {len(view_over)} (sorted by DueDate: oldest → newest)")
    st.dataframe(view_over, use_container_width=True, hide_index=True)
    if not view_over.empty:
        tiny_xlsx_button(
            view_over,
            sheet="Expected_Overdue",
            filename="Expected_Overdue.xlsx",
            key_suffix="expected_over",
        )

    st.markdown("---")

    # --- Section: Due ---
    st.markdown("#### ✅ Due")
    view_due = status_view(slim, "Due")[out_cols]
    st.caption(f"Rows: {len(view_due)}")
    st.dataframe(view_due, use_container_width=True, hide_index=True)
    if not view_due.empty:
        tiny_xlsx_button(
            view_due,
            sheet="Expected_Due",
            filename="Expected_Due.xlsx",
            key_suffix="expected_due",
        )

    # ======================================================================
    # Extra Section: Open WOs for Procedure Template ID 2132428 (from WO-DB)
    # ======================================================================
    st.markdown("---")
    st.markdown("#### 🔍 Open Work Orders for Procedure Template 2132428")

    from pathlib import Path as _Path

    # Locate and load WO-DB.parquet
    parq_dir = _Path(
        globals().get("PARQUET_DIR", st.session_state.get("parquet_dir", ""))  # type: ignore[arg-type]
        or (_Path.cwd() / "parquet_db")
    )

    wo_candidates = [
        parq_dir / "WO-DB.parquet",
        parq_dir / "WO_DB.parquet",
        parq_dir / "WO_DB.parquest",
        parq_dir / "WO-DB.parquest",
    ]
    wo_path = None
    for p in wo_candidates:
        if p.is_file():
            wo_path = p
            break

    if wo_path is None:
        st.info(
            "Could not find **WO-DB.parquet** in the parquet folder — "
            "skipping extra open-WO tables."
        )
        return

    df_wo = pd.read_parquet(wo_path)
    if df_wo.empty:
        st.info("WO-DB.parquet is empty — no work orders to analyze for extra tables.")
        return

    # Apply main location filter to WO_DB, using Location2
    loc_col_wo = _pick_col(df_wo, ["Location2"])  # force Location2 as the filter
    if loc_col_wo and apply_filter:
        allowed = {
            s.strip().casefold()
            for s in selected_locations
            if isinstance(s, str) and s.strip() and s.strip() not in admin_all_tokens
        }
        if allowed:
            loc_norm_wo = df_wo[loc_col_wo].astype(str).str.strip().str.casefold()
            mask_wo = loc_norm_wo.isin(allowed)
            df_wo_filt = df_wo[mask_wo].copy()
        else:
            df_wo_filt = df_wo.copy()
    else:
        df_wo_filt = df_wo.copy()

    if df_wo_filt.empty:
        st.info("No work orders after applying the main location filter to WO-DB.parquet.")
        return

    # Pick key columns in WO_DB
    proc_col = _pick_col(
        df_wo_filt,
        [
            "Procedure Template ID",
            "Procedure_Template_ID",
            "procedure_template_id",
            "ProcedureTemplateId",
            "Procedure Template Id",
        ],
    )
    isopen_col = _pick_col(df_wo_filt, ["IsOpen", "Is Open", "is_open", "ISOPEN"])
    start_col = _pick_col(
        df_wo_filt,
        ["Planned Start Date", "Start Date", "Planned Start", "start_date"],
    )
    due_col = _pick_col(df_wo_filt, ["Due date", "Due Date", "Due_Date", "due_date"])

    if proc_col is None or isopen_col is None:
        st.info(
            "WO-DB.parquet does not contain the expected columns for "
            "`Procedure Template ID` and/or `IsOpen` — skipping extra tables."
        )
        return

    # Filter: Procedure Template ID = 2132428 AND IsOpen = True
    target_proc = 2132428

    # Procedure filter supports numeric or string storage
    proc_series = df_wo_filt[proc_col]
    try:
        proc_num = pd.to_numeric(proc_series, errors="coerce")
        mask_proc = proc_num.eq(target_proc)
    except Exception:
        mask_proc = proc_series.astype(str).str.strip() == str(target_proc)

    # IsOpen filter supports bool or numeric
    isopen_series = df_wo_filt[isopen_col]
    if isopen_series.dtype == bool:
        mask_open = isopen_series.fillna(False)
    else:
        isopen_num = pd.to_numeric(isopen_series, errors="coerce")
        mask_open = isopen_num.fillna(0).ne(0)

    df_wo_proc_open = df_wo_filt[mask_proc & mask_open].copy()

    if df_wo_proc_open.empty:
        st.info(
            "No open work orders found for **Procedure Template ID = 2132428** "
            "after location filtering."
        )
        return

    # Helper: "has value" for start/due date fields
    def _has_value(series: pd.Series) -> pd.Series:
        if series is None:
            return pd.Series(False, index=df_wo_proc_open.index)
        s = series
        # If datetime-like, non-null is enough; otherwise, non-empty string
        if pd.api.types.is_datetime64_any_dtype(s):
            return s.notna()
        s_str = s.astype(str)
        return s_str.notna() & s_str.str.strip().ne("")

    has_start = _has_value(df_wo_proc_open[start_col]) if start_col else pd.Series(False, index=df_wo_proc_open.index)
    has_due = _has_value(df_wo_proc_open[due_col]) if due_col else pd.Series(False, index=df_wo_proc_open.index)

    mask_with_dates = has_start | has_due
    mask_without_dates = ~mask_with_dates

    # Table A: without start or due date
    df_no_dates = df_wo_proc_open[mask_without_dates].copy()

    # Table B: with start or due date
    df_with_dates = df_wo_proc_open[mask_with_dates].copy()

    # --- Base display of extra tables (full, for reference) ---
    st.markdown("##### A) Open WOs for Proc 2132428 — **No Start/Due Date** (raw)")
    st.caption(f"Rows: {len(df_no_dates)}")
    st.dataframe(df_no_dates, use_container_width=True, hide_index=True)
    if not df_no_dates.empty:
        tiny_xlsx_button(
            df_no_dates,
            sheet="WO_Proc2132428_NoDates",
            filename="WO_Proc2132428_NoStartDue.xlsx",
            key_suffix="wo_proc213_no_dates",
        )

    st.markdown("---")

    st.markdown("##### B) Open WOs for Proc 2132428 — **With Start or Due Date** (raw)")
    st.caption(f"Rows: {len(df_with_dates)}")
    st.dataframe(df_with_dates, use_container_width=True, hide_index=True)
    if not df_with_dates.empty:
        tiny_xlsx_button(
            df_with_dates,
            sheet="WO_Proc2132428_WithDates",
            filename="WO_Proc2132428_WithStartDue.xlsx",
            key_suffix="wo_proc213_with_dates",
        )

    # ======================================================================
    # Calendars: Expected (DueDate) + Open Proc 2132428 (Start/Due)
    # ======================================================================
    st.markdown("---")
    st.markdown("#### 📅 Expected & Open Proc 2132428 Calendars")

    today = _date.today()
    cal_start = st.date_input(
        "Calendar start date (Expected DueDate + Open Proc 2132428)",
        value=today,
        key="expected_cal_start",
    )

    # ---------- Collect day items: date -> list[(type, text)] ----------
    day_items: dict[_date, list] = {}

    # 1) Expected table (DueDate), show Name - Next Service Type - Remaining
    if "view_exp" in locals():
        exp_df = view_exp.copy()
    else:
        exp_df = pd.DataFrame()

    if not exp_df.empty and "DueDate" in exp_df.columns:
        due_ser = pd.to_datetime(exp_df["DueDate"], errors="coerce").dt.date
        for idx, row in exp_df.iterrows():
            d = due_ser.get(idx, None)
            if d is None or pd.isna(d):
                continue

            name      = _safe_str_cell(row.get("Name", ""))
            next_type = _safe_str_cell(row.get("Next Service Type", ""))
            remaining = _safe_str_cell(row.get("Remaining", ""))

            parts = [name, next_type, remaining]
            parts = [p for p in parts if p]
            if not parts:
                continue

            label = " - ".join(parts)
            day_items.setdefault(d, []).append(("exp", label))

    # 2) Open WOs for Proc 2132428 WITH Start/Due (df_with_dates)
    from reportlab.lib.pagesizes import letter, landscape
    from reportlab.pdfgen import canvas

    # Prefer the true WORKORDER column if present
    if "WORKORDER" in df_wo_proc_open.columns:
        wo_col = "WORKORDER"
    elif "Workorder" in df_wo_proc_open.columns:
        wo_col = "Workorder"
    else:
        wo_col = _pick_col(
            df_wo_proc_open,
            [
                "Work Order",
                "Work Order #",
                "WO_ID",
                "ID",
                "WO",  # last resort, often a flag
            ],
        )

    title_col = _pick_col(
        df_wo_proc_open,
        ["TITLE", "Title", "Summary", "Description"]
    )

    # Asset (Name) and Remaining columns in WO DB (for tables later)
    asset_col_wo = _pick_col(
        df_wo_proc_open,
        ["Name", "Asset", "Asset Name", "Line Name"]
    )
    remaining_col_wo = _pick_col(
        df_wo_proc_open,
        ["Remaining", "Remaining to Service", "Remaining Reading", "Remaining Hours", "Remaining Miles"]
    )

    if not df_with_dates.empty:
        # Build schedule date series: prefer Due, else Start
        sched_raw = pd.Series(pd.NaT, index=df_with_dates.index, dtype="object")
        if due_col:
            sched_raw = df_with_dates[due_col]
        if start_col:
            sched_raw = sched_raw.where(pd.notna(sched_raw), df_with_dates[start_col])

        sched_dt = pd.to_datetime(sched_raw, errors="coerce").dt.date

        for idx, row in df_with_dates.iterrows():
            d = sched_dt.get(idx, None)
            if d is None or pd.isna(d):
                continue

            # Decide whether this date is Due or Start for label
            due_val = row.get(due_col, None) if due_col else None
            start_val = row.get(start_col, None) if start_col else None
            kind = "Due"
            try:
                if pd.isna(due_val) and not pd.isna(start_val):
                    kind = "Start"
                else:
                    dv = pd.to_datetime(due_val).date() if pd.notna(due_val) else None
                    sv = pd.to_datetime(start_val).date() if pd.notna(start_val) else None
                    if sv == d and dv != d:
                        kind = "Start"
            except Exception:
                pass

            wo_raw = row.get(wo_col, "") if wo_col else ""
            wo_txt = ""
            if pd.notna(wo_raw) and wo_raw != "":
                try:
                    wo_txt = str(int(float(wo_raw)))
                except Exception:
                    wo_txt = str(wo_raw)

            title_txt = _safe_str_cell(row.get(title_col, "")) if title_col else ""
            label = f"{wo_txt} - {title_txt} - {d} ({kind})".strip(" -")
            day_items.setdefault(d, []).append(("wo", label))

    # ---------- Condensed table for open WOs with NO Start/Due ----------
    table_cols = []
    if wo_col and wo_col in df_no_dates.columns:
        table_cols.append(wo_col)
    if title_col and title_col in df_no_dates.columns:
        table_cols.append(title_col)
    if asset_col_wo and asset_col_wo in df_no_dates.columns:
        table_cols.append(asset_col_wo)
    if remaining_col_wo and remaining_col_wo in df_no_dates.columns:
        table_cols.append(remaining_col_wo)

    if not table_cols:
        df_no_dates_view = df_no_dates.copy()
    else:
        df_no_dates_view = df_no_dates[table_cols].copy()

    # Clean WO number formatting in this view (strip .0)
    if wo_col and wo_col in df_no_dates_view.columns:
        def _fmt_wo(v):
            if v is None or (isinstance(v, str) and not v.strip()):
                return ""
            try:
                return str(int(float(v)))
            except Exception:
                return str(v)
        df_no_dates_view[wo_col] = df_no_dates_view[wo_col].apply(_fmt_wo)

    st.markdown("##### 📋 Open WOs (Proc 2132428) — No Start/Due Date (Table for PDF)")
    st.caption(f"Rows: {len(df_no_dates_view)}")
    st.dataframe(df_no_dates_view, use_container_width=True, hide_index=True)

    # ---------- Build small tables for PDF from status views ----------
    def _pdf_subset(src_df: pd.DataFrame, cols: list[str]) -> pd.DataFrame:
        if src_df is None or src_df.empty:
            return src_df.iloc[0:0].copy()
        avail = [c for c in cols if c in src_df.columns]
        if not avail:
            return src_df.iloc[0:0].copy()
        return src_df[avail].copy()

    # Needs Service: Name - Last Service Type - Date of Last service
    tbl_needs = _pdf_subset(view_ns, ["Name", "Last Service Type", "Date of Last service"])

    # New Reading Needed: Name - Last Reading - Date
    tbl_newr = _pdf_subset(view_newr, ["Name", "Last Reading", "Date"])

    # Overdue: Name - Next Service Type - Remaining
    tbl_over = _pdf_subset(view_over, ["Name", "Next Service Type", "Remaining"])

    # Due: Name - Next Service Type - Remaining
    tbl_due = _pdf_subset(view_due, ["Name", "Next Service Type", "Remaining"])

    # ---------- Calendar HTML styling ----------
    style_block = """
    <style>
    table.exp-cal {
        border-collapse: collapse;
        width: 100%;
        font-size: 0.8rem;
    }
    table.exp-cal th, table.exp-cal td {
        border: 1px solid #ccc;
        vertical-align: top;
        padding: 4px;
    }
    table.exp-cal th {
        background-color: #f0f0f0;
        text-align: center;
        font-weight: bold;
    }
    table.exp-cal ul {
        margin: 0;
        padding-left: 16px;
    }
    .exp-item {
        color: #118811;
        font-weight: 600;
    }
    .wo-item {
        color: #222222;
    }
    </style>
    """

    # ---------- Helper: build calendar HTML (aligned by weekday) ----------
    def _build_calendar_grid_html(title_txt: str, start_day: _date, num_days: int) -> str:
        end_day = start_day + timedelta(days=num_days - 1)

        # Align dates under correct weekday: 0 = Monday
        days = [start_day + timedelta(days=i) for i in range(num_days)]
        first_wd = start_day.weekday()  # 0=Mon
        cells: list[_date | None] = [None] * first_wd + days
        weeks = math.ceil(len(cells) / 7)

        parts = []
        parts.append(f"<h3>{html.escape(title_txt)} ({start_day} to {end_day})</h3>")
        parts.append("<table class='exp-cal'>")

        # Header row (Mon–Sun)
        parts.append("<tr>")
        for name in ["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"]:
            parts.append(f"<th>{name}</th>")
        parts.append("</tr>")

        idx = 0
        for _w in range(weeks):
            parts.append("<tr>")
            for _d in range(7):
                if idx < len(cells) and cells[idx] is not None:
                    day = cells[idx]
                    idx += 1
                    items = day_items.get(day, [])
                    cell_html = f"<div><strong>{day.strftime('%Y-%m-%d')}</strong></div>"
                    if items:
                        lis = []
                        for raw in items:
                            # Support both (type, text) and plain string
                            if isinstance(raw, (tuple, list)) and len(raw) == 2:
                                item_type, text = raw
                            else:
                                item_type, text = "wo", raw
                            cls = "exp-item" if item_type == "exp" else "wo-item"
                            lis.append(
                                f"<li><span class='{cls}'>{html.escape(str(text))}</span></li>"
                            )
                        cell_html += "<ul>" + "".join(lis) + "</ul>"
                    parts.append(f"<td>{cell_html}</td>")
                else:
                    idx += 1
                    parts.append("<td>&nbsp;</td>")
            parts.append("</tr>")

        parts.append("</table>")
        return "".join(parts)

    # ---------- Helper: build calendar + tables PDF ----------
    def _build_calendar_pdf_bytes(
        title_txt: str,
        start_day: _date,
        num_days: int,
        table_no_dates: pd.DataFrame,
        tbl_needs: pd.DataFrame,
        tbl_new: pd.DataFrame,
        tbl_over: pd.DataFrame,
        tbl_due: pd.DataFrame,
    ) -> bytes:
        from reportlab.platypus import Table, TableStyle, Paragraph
        from reportlab.lib.styles import getSampleStyleSheet

        buf = BytesIO()
        c = canvas.Canvas(buf, pagesize=landscape(letter))
        width, height = landscape(letter)

        margin = 36  # half-inch
        title_height = 24

        # Styles for wrapped table cells
        styles = getSampleStyleSheet()
        cell_style = styles["Normal"]
        cell_style.fontName = "Helvetica"
        cell_style.fontSize = 7
        cell_style.leading = 8
        cell_style.wordWrap = "CJK"  # allow wrapping

        MAX_ROWS_PER_PAGE = 20  # limit rows per page so table doesn't run off

        # ----------------- Page 1: Calendar (aligned by weekday) -----------------
        c.setFont("Helvetica-Bold", 14)
        c.drawString(
            margin,
            height - margin,
            f"{title_txt} ({start_day} to {start_day + timedelta(days=num_days-1)})",
        )

        grid_top = height - margin - title_height
        grid_bottom = margin
        grid_height = grid_top - grid_bottom

        cols = 7
        # align by weekday
        days = [start_day + timedelta(days=i) for i in range(num_days)]
        first_wd = start_day.weekday()
        cal_cells: list[_date | None] = [None] * first_wd + days
        rows = math.ceil(len(cal_cells) / cols)

        cell_w = (width - 2 * margin) / cols
        cell_h = grid_height / rows

        # Headers (Mon–Sun)
        c.setFont("Helvetica-Bold", 10)
        for col_idx, name in enumerate(["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"]):
            x = margin + col_idx * cell_w + 2
            y = grid_top + 4
            c.drawString(x, y, name)

        # Calendar cells
        idx = 0
        for row_i in range(rows):
            for col_i in range(cols):
                x = margin + col_i * cell_w
                y_top = grid_top - row_i * cell_h
                y_bottom = y_top - cell_h

                c.rect(x, y_bottom, cell_w, cell_h, stroke=1, fill=0)

                if idx < len(cal_cells) and cal_cells[idx] is not None:
                    day = cal_cells[idx]
                    idx += 1

                    # date
                    c.setFont("Helvetica-Bold", 8)
                    c.drawString(x + 2, y_top - 10, day.strftime("%Y-%m-%d"))

                    # items
                    items = day_items.get(day, [])  # from outer scope
                    c.setFont("Helvetica", 7)
                    text_y = y_top - 22
                    max_lines = int((cell_h - 18) / 9)
                    line_count = 0
                    max_width = cell_w - 8

                    for raw in items:
                        if line_count >= max_lines:
                            c.setFillColor(colors.black)
                            c.drawString(x + 4, text_y, "... more ...")
                            break

                        # support ("exp", text) or ("wo", text) or plain text
                        if isinstance(raw, (tuple, list)) and len(raw) == 2:
                            item_type, item_text = raw
                        else:
                            item_type, item_text = "wo", raw

                        # green for Expected, black for WO
                        c.setFillColor(colors.green if item_type == "exp" else colors.black)

                        s = str(item_text)
                        words = s.split()
                        line = ""

                        for w in words:
                            test = (line + " " + w).strip()
                            if c.stringWidth(test, "Helvetica", 7) <= max_width:
                                line = test
                            else:
                                if line_count >= max_lines:
                                    c.setFillColor(colors.black)
                                    c.drawString(x + 4, text_y, "... more ...")
                                    line_count += 1
                                    break
                                c.drawString(x + 4, text_y, line)
                                text_y -= 9
                                line_count += 1
                                line = w

                        if line and line_count < max_lines:
                            c.drawString(x + 4, text_y, line)
                            text_y -= 9
                            line_count += 1

                        if line_count >= max_lines:
                            break

                    c.setFillColor(colors.black)
                else:
                    idx += 1  # advance even for empty cell

        # Calendar page done
        c.showPage()

        # ---------- Helper: paginated, wrapped table pages ----------
        def _draw_table_pages(title: str, df_in: pd.DataFrame):
            if df_in is None or df_in.empty:
                return

            df_str = df_in.astype(str).reset_index(drop=True)
            headers = list(df_str.columns)
            ncols = max(1, len(headers))

            col_width = (width - 2 * margin) / ncols
            avail_width = width - 2 * margin
            avail_height = height - margin - title_height - 12

            total_rows = len(df_str)
            start_row = 0

            while start_row < total_rows:
                chunk = df_str.iloc[start_row:start_row + MAX_ROWS_PER_PAGE]

                # Build data with Paragraphs for wrapping
                data = []
                # header row
                header_cells = [Paragraph(str(h), cell_style) for h in headers]
                data.append(header_cells)
                # body rows
                for _, row in chunk.iterrows():
                    row_cells = [Paragraph(str(v), cell_style) for v in row]
                    data.append(row_cells)

                table = Table(
                    data,
                    colWidths=[col_width] * ncols,
                )
                table.setStyle(
                    TableStyle(
                        [
                            ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                            ("FONTSIZE", (0, 0), (-1, -1), 7),
                            ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ]
                    )
                )

                c.setPageSize(landscape(letter))
                c.setFont("Helvetica-Bold", 14)
                c.drawString(margin, height - margin, title)

                tw, th = table.wrapOn(c, avail_width, avail_height)
                table.drawOn(c, margin, height - margin - title_height - th)

                c.showPage()
                start_row += MAX_ROWS_PER_PAGE

        # ---------- Pages for each table ----------

        # 1) All Assets — Expected Services Matrix (use view_all from outer scope)
        try:
            if view_all is not None and not view_all.empty:
                _draw_table_pages(
                    "All Assets — Expected Services Matrix",
                    view_all,
                )
        except NameError:
            # view_all not defined in this run — just skip it
            pass

        # 2) Open WOs (Proc 2132428) — No Start/Due Date
        if not table_no_dates.empty:
            _draw_table_pages(
                "Open WOs (Proc 2132428) — No Start/Due Date",
                table_no_dates,
            )

        # 3) Needs Service
        if not tbl_needs.empty:
            _draw_table_pages("Needs Service", tbl_needs)

        # 4) New Reading Needed
        if not tbl_new.empty:
            _draw_table_pages("New Reading Needed", tbl_new)

        # 5) Overdue
        if not tbl_over.empty:
            _draw_table_pages("Overdue", tbl_over)

        # 6) Due
        if not tbl_due.empty:
            _draw_table_pages("Due", tbl_due)

        c.save()
        buf.seek(0)
        return buf.getvalue()

    # ---------- Build and display calendars (HTML) ----------
    cal7_html = _build_calendar_grid_html("Next 7 Days", cal_start, 7)
    cal30_html = _build_calendar_grid_html("Next 30 Days", cal_start, 30)

    st.markdown(style_block, unsafe_allow_html=True)
    col7, col30 = st.columns(2)
    with col7:
        st.markdown(cal7_html, unsafe_allow_html=True)
    with col30:
        st.markdown(cal30_html, unsafe_allow_html=True)

    # ---------- Build PDFs (calendar + all tables) ----------
    pdf_7 = _build_calendar_pdf_bytes(
        "Expected & Open Proc 2132428 — Next 7 Days",
        cal_start,
        7,
        df_no_dates_view,
        tbl_needs,
        tbl_newr,
        tbl_over,
        tbl_due,
    )
    pdf_30 = _build_calendar_pdf_bytes(
        "Expected & Open Proc 2132428 — Next 30 Days",
        cal_start,
        30,
        df_no_dates_view,
        tbl_needs,
        tbl_newr,
        tbl_over,
        tbl_due,
    )

    st.download_button(
        "⬇️ 7-Day Expected/Proc Calendar + Tables (PDF)",
        data=pdf_7,
        file_name=f"Expected_Proc2132428_Calendar_7day_{cal_start}.pdf",
        mime="application/pdf",
        key="dl_expected_cal_7_pdf",
    )

    st.download_button(
        "⬇️ 30-Day Expected/Proc Calendar + Tables (PDF)",
        data=pdf_30,
        file_name=f"Expected_Proc2132428_Calendar_30day_{cal_start}.pdf",
        mime="application/pdf",
        key="dl_expected_cal_30_pdf",
    )



# =========================
# WO REPORT TAB (STATUS SNAPSHOT)
# =========================
from pathlib import Path
from datetime import date as _date, timedelta
import math
import html  # for safe HTML escaping
import re
import io
from reportlab.lib.pagesizes import letter, landscape
from reportlab.pdfgen import canvas
from reportlab.lib import colors
import pandas as pd
import streamlit as st


def render_wo_report(
    df_wo: pd.DataFrame | None,  # kept for signature compatibility, but not used
    start_date,
    end_date,
    selected_locations,
):
    """
    Work Order Report — status snapshot from WO_Trans_non_DB.parquet.

    • Filters base set by selected_locations (main page)
    • Builds 4+1 tables with a common column set:

        WORKORDER, TITLE, Description, step_progress, PO, P/N,
        Completed on, assignee, team_assigned, Created on,
        Planned Start Date, Due date, IsOverDue, DaysOverDue,
        ComOverDue, DaysOverDue_Com, PrePlaDat, ReaPlanDat,
        KPI_Pre, KPI_Rea, KPI_Oth, IsOpen, DaysOpen,
        total_days, make, model, serial, year, rea_pre

    Tables:
      1) Completed in selected time frame (by Completed on)
      2) IsOverDue = True
      3) ComOverDue = True (AND Completed on in selected time frame, if range given)
      4) Currently Open (IsOpen = True AND (Due/Start <= today))
      5) Upcoming Open (IsOpen = True AND today < (Due/Start))
         + printable 7-day and 30-day calendar views
    """

    st.markdown("### 🧾 Work Orders — Status Report")

    # ---------- Locate WO_Trans_non_DB parquet ----------
    parq_dir = Path(
        globals().get("PARQUET_DIR", st.session_state.get("parquet_dir", ""))  # type: ignore[arg-type]
        or (Path.cwd() / "parquet_db")
    )

    wo_candidates = [
        parq_dir / "WO_Trans_non_DB.parquet",
        parq_dir / "WO-Trans_non_DB.parquet",
        parq_dir / "WO-Trans_no_DB.parquet",
    ]
    wo_path = None
    for p in wo_candidates:
        if p.is_file():
            wo_path = p
            break

    if wo_path is None:
        st.error(
            "Could not find **WO_Trans_non_DB.parquet**.\n\n"
            "Tried:\n" + "\n".join(str(p) for p in wo_candidates)
        )
        return

    df = pd.read_parquet(wo_path)
    if df.empty:
        st.info("WO_Trans_non_DB.parquet is empty — no work orders to show.")
        return

    # ---------- Column locators ----------
    loc_col = _find(df, "Location", "location")
    comp_src_col = _find(df, "COMPLETED ON", "Completed On", "Completed on")

    # Normalize Completed On to datetime
    if comp_src_col:
        df[comp_src_col] = pd.to_datetime(df[comp_src_col], errors="coerce")

    # ---------- Apply location filter from main page ----------
    if loc_col and selected_locations:
        df = df[df[loc_col].isin(selected_locations)]

    if df.empty:
        st.info("No work orders match the current **location** filters.")
        return

    # ---------- Desired output columns + candidate names ----------
    desired_cols = [
        "WORKORDER",
        "TITLE",
        "Description",
        "step_progress",
        "PO",
        "P/N",
        "Completed on",
        "assignee",
        "team_assigned",
        "Created on",
        "Planned Start Date",
        "Due date",
        "IsOverDue",
        "DaysOverDue",
        "ComOverDue",
        "DaysOverDue_Com",
        "PrePlaDat",
        "ReaPlanDat",
        "KPI_Pre",
        "KPI_Rea",
        "KPI_Oth",
        "IsOpen",
        "DaysOpen",
        "total_days",
        "make",
        "model",
        "serial",
        "year",
        "rea_pre",
    ]

    col_spec = {
        "WORKORDER": ["WORKORDER", "Workorder", "workorder"],
        "TITLE": ["TITLE", "Title"],
        "Description": ["Description", "DESC"],
        "step_progress": ["step_progress", "STATUSstep_progress"],
        "PO": ["PO"],
        "P/N": ["P/N", "P N", "PN"],
        "Completed on": ["COMPLETED ON", "Completed On", "Completed on"],
        "assignee": ["assignee", "Assigned To"],
        "team_assigned": ["team_assigned", "team assigned", "team"],
        "Created on": ["Created on", "Created On", "Created date", "Created Date"],
        "Planned Start Date": ["Planned Start Date", "Start Date", "Planned Start"],
        "Due date": ["Due date", "Due Date", "Due"],
        "IsOverDue": ["IsOverDue"],
        "DaysOverDue": ["DaysOverDue"],
        "ComOverDue": ["ComOverDue"],
        "DaysOverDue_Com": ["DaysOverDue_Com"],
        "PrePlaDat": ["PrePlaDat"],
        "ReaPlanDat": ["ReaPlanDat"],
        "KPI_Pre": ["KPI_Pre"],
        "KPI_Rea": ["KPI_Rea"],
        "KPI_Oth": ["KPI_Oth"],
        "IsOpen": ["IsOpen"],
        "DaysOpen": ["DaysOpen"],
        "total_days": ["total_days", "TotalDays"],
        "make": ["make", "Make", "Manufacturer"],
        "model": ["model", "Model"],
        "serial": ["serial", "Serial", "Serial Number"],
        "year": ["year", "Year"],
        "rea_pre": ["rea_pre", "Rea_Pre", "REA_PRE"],
    }

    rename_map: dict[str, str] = {}
    for out_name, cands in col_spec.items():
        actual = _find(df, *cands)
        if actual:
            rename_map[actual] = out_name

    if not rename_map:
        st.error("None of the expected WO columns were found in WO_Trans_non_DB.parquet.")
        return

    df_sel = df[list(rename_map.keys())].rename(columns=rename_map)

    # ensure display order is exactly as requested, but only keep ones we actually have
    display_order = [c for c in desired_cols if c in df_sel.columns]

    # Normalize key dates into df_sel (date only)
    if "Completed on" in df_sel.columns:
        df_sel["Completed on"] = pd.to_datetime(
            df_sel["Completed on"], errors="coerce"
        ).dt.date
    if "Planned Start Date" in df_sel.columns:
        df_sel["Planned Start Date"] = pd.to_datetime(
            df_sel["Planned Start Date"], errors="coerce"
        ).dt.date
    if "Due date" in df_sel.columns:
        df_sel["Due date"] = pd.to_datetime(
            df_sel["Due date"], errors="coerce"
        ).dt.date
    if "Created on" in df_sel.columns:
        df_sel["Created on"] = pd.to_datetime(
            df_sel["Created on"], errors="coerce"
        ).dt.date

    # ---------- User & Team filters (single-select, split multi-values) ----------
    df_base = df_sel.copy()
    filt_cols = st.columns(2)

    # --- User filter (assignee) ---
    if "assignee" in df_base.columns:
        raw_users = df_base["assignee"].dropna().astype(str)
        user_tokens: set[str] = set()
        for s in raw_users:
            for part in re.split(r"[;,]", s):
                name = part.strip()
                if name:
                    user_tokens.add(name)

        user_opts = ["All"] + sorted(user_tokens)
        with filt_cols[0]:
            user_sel = st.selectbox("User filter", user_opts, index=0)

        if user_sel != "All":
            pattern = rf"(^|[,;])\s*{re.escape(user_sel)}\s*(?=,|;|$)"
            mask_user = df_base["assignee"].astype(str).str.contains(
                pattern, regex=True, na=False
            )
            df_base = df_base[mask_user]
    else:
        user_sel = "All"

    # --- Team filter (team_assigned) ---
    if "team_assigned" in df_base.columns:
        raw_teams = df_base["team_assigned"].dropna().astype(str)
        team_tokens: set[str] = set()
        for s in raw_teams:
            for part in re.split(r"[;,]", s):
                name = part.strip()
                if name:
                    team_tokens.add(name)

        team_opts = ["All"] + sorted(team_tokens)
        with filt_cols[1]:
            team_sel = st.selectbox("Team filter", team_opts, index=0)

        if team_sel != "All":
            pattern = rf"(^|[,;])\s*{re.escape(team_sel)}\s*(?=,|;|$)"
            mask_team = df_base["team_assigned"].astype(str).str.contains(
                pattern, regex=True, na=False
            )
            df_base = df_base[mask_team]
    else:
        team_sel = "All"

    if df_base.empty:
        st.info("No work orders match the current **location / user / team** filters.")
        return

    # Recompute display_order based on df_base columns (same logical order)
    display_order = [c for c in desired_cols if c in df_base.columns]

    # Trim helper columns: everything to the RIGHT of "Due date"
    if "Due date" in display_order:
        cut_idx = display_order.index("Due date")
        display_order_trim = display_order[: cut_idx + 1]
    else:
        display_order_trim = display_order  # no Due date: keep all

    # ---------- Helper: boolean masks for KPI flags ----------
    def _bool_mask(df_in: pd.DataFrame, col_name: str) -> pd.Series:
        if col_name not in df_in.columns:
            return pd.Series(False, index=df_in.index)
        s = df_in[col_name]
        if s.dtype == bool:
            return s.fillna(False)
        num = pd.to_numeric(s, errors="coerce")
        return num.fillna(0) != 0

    # Completed in selected time frame (using "Completed on" in df_base)
    if "Completed on" in df_base.columns and start_date and end_date:
        comp_dates = df_base["Completed on"]
        mask_completed_range_df = comp_dates.between(start_date, end_date)
    else:
        mask_completed_range_df = pd.Series(False, index=df_base.index)

    # ---------- Build main masks/tables from df_base ----------

    # 1) Completed in selected time frame (by Completed on)
    df_completed = df_base[mask_completed_range_df].copy()

    # 2) IsOverDue (all overdue after filters — not date-restricted)
    df_overdue = df_base[_bool_mask(df_base, "IsOverDue")].copy()

    # 3) ComOverDue (date-restricted if range exists)
    base_com_mask = _bool_mask(df_base, "ComOverDue")
    if "Completed on" in df_base.columns and start_date and end_date:
        df_com_overdue = df_base[base_com_mask & mask_completed_range_df].copy()
    else:
        df_com_overdue = df_base[base_com_mask].copy()

    # 4) IsOpen (all open after filters — for splitting "currently open" vs "upcoming")
    open_mask = _bool_mask(df_base, "IsOpen")

    # ---------- Build schedule dates (Due or Planned Start) ----------
    sched_raw = pd.Series(pd.NaT, index=df_base.index, dtype="object")
    if "Due date" in df_base.columns:
        sched_raw = df_base["Due date"]
    if "Planned Start Date" in df_base.columns:
        sched_raw = sched_raw.where(pd.notna(sched_raw), df_base["Planned Start Date"])

    today = _date.today()
    sched_dt = pd.to_datetime(sched_raw, errors="coerce").dt.date

    mask_has_sched = sched_dt.notna()
    mask_past_or_today = mask_has_sched & (sched_dt <= today)
    mask_future = mask_has_sched & (sched_dt > today)

    # Currently open (due or started already)
    mask_open_due_started = open_mask & mask_past_or_today
    df_open_due_started = df_base[mask_open_due_started].copy()

    # Upcoming open (not due yet — used for calendars)
    upcoming_mask = open_mask & mask_future
    df_upcoming = df_base[upcoming_mask].copy()
    sched_up = sched_dt[upcoming_mask]

    # ---------- Layout / base caption ----------
    loc_label = (
        ", ".join(map(str, selected_locations)) if selected_locations else "All"
    )
    st.caption(
        f"Base WO rows after filters: **{len(df_base)}**  "
        f"(Locations: {loc_label} • User: {user_sel} • Team: {team_sel})"
    )

    # --- Completed in time frame ---
    st.markdown("#### ✅ Completed in Selected Time Frame")
    if start_date and end_date:
        st.caption(f"Completed between **{start_date}** and **{end_date}** (by *Completed on*).")

    df_view = df_completed[display_order_trim] if not df_completed.empty else df_completed
    st.dataframe(df_view, use_container_width=True, hide_index=True)
    if not df_completed.empty and display_order_trim:
        st.download_button(
            "⬇️ Download Completed WOs (XLSX)",
            data=_to_xlsx_bytes(df_view),
            file_name="WO_completed_in_range.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="dl_wo_completed",
        )

    st.markdown("---")

    # --- IsOverDue ---
    st.markdown("#### ⏰ Overdue Work Orders (IsOverDue = True)")
    df_view = df_overdue[display_order_trim] if not df_overdue.empty else df_overdue
    st.dataframe(df_view, use_container_width=True, hide_index=True)
    if not df_overdue.empty and display_order_trim:
        st.download_button(
            "⬇️ Download Overdue WOs (XLSX)",
            data=_to_xlsx_bytes(df_view),
            file_name="WO_IsOverDue.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="dl_wo_overdue",
        )

    st.markdown("---")

    # --- ComOverDue (date-restricted if range exists) ---
    st.markdown("#### 📉 Completed Overdue (ComOverDue = True)")
    if start_date and end_date and "Completed on" in df_base.columns:
        st.caption(
            f"Completed between **{start_date}** and **{end_date}** "
            f"AND ComOverDue = True."
        )

    df_view = df_com_overdue[display_order_trim] if not df_com_overdue.empty else df_com_overdue
    st.dataframe(df_view, use_container_width=True, hide_index=True)
    if not df_com_overdue.empty and display_order_trim:
        st.download_button(
            "⬇️ Download Completed-Overdue WOs (XLSX)",
            data=_to_xlsx_bytes(df_view),
            file_name="WO_ComOverDue.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="dl_wo_comoverdue",
        )

    st.markdown("---")

    # --- Currently Open (Due/Started) ---
    st.markdown("#### 🟡 Currently Open (Due or Started) — IsOpen = True AND (Due/Start ≤ today)")
    df_view = df_open_due_started[display_order_trim] if not df_open_due_started.empty else df_open_due_started
    st.dataframe(df_view, use_container_width=True, hide_index=True)
    if not df_open_due_started.empty and display_order_trim:
        st.download_button(
            "⬇️ Download Currently Open (Due/Started) WOs (XLSX)",
            data=_to_xlsx_bytes(df_view),
            file_name="WO_CurrentlyOpen_DueStarted.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="dl_wo_open_due_started",
        )

    st.markdown("---")

    # --- Upcoming Open (Not Overdue) ---
    st.markdown("#### 📅 Upcoming Open (Not Due Yet)")

    if df_upcoming.empty or (
        "IsOpen" not in df_base.columns or
        ("Due date" not in df_base.columns and "Planned Start Date" not in df_base.columns)
    ):
        st.info("No upcoming open work orders (not overdue) for the calendar.")
        return

    df_upcoming_view = df_upcoming[display_order_trim] if display_order_trim else df_upcoming
    st.caption(f"Rows: {len(df_upcoming_view)}")
    st.dataframe(df_upcoming_view, use_container_width=True, hide_index=True)

    # ---------- Calendar start date picker ----------
    st.markdown("#### 🖨️ Printable Calendars")
    today = _date.today()
    cal_start = st.date_input(
        "Calendar start date",
        value=today,
        key="wo_calendar_start",
    )

    # ------- Build printable calendar data (WO - Title - Start/Due) -------
    day_items: dict[_date, list[str]] = {}

    for idx, sched_date in sched_up.items():
        if sched_date is None or pd.isna(sched_date):
            continue
        if not isinstance(sched_date, _date):
            try:
                sched_date = pd.to_datetime(sched_date).date()
            except Exception:
                continue

        row = df_upcoming.loc[idx]

        wo_raw = row.get("WORKORDER", "")
        wo = ""
        if pd.notna(wo_raw) and wo_raw != "":
            try:
                wo = str(int(float(wo_raw)))
            except Exception:
                wo = str(wo_raw)
        title = str(row.get("TITLE", "") or "")

        due_val = row.get("Due date", None)
        start_val = row.get("Planned Start Date", None)

        label_kind = "Due"
        if pd.isna(due_val) and not pd.isna(start_val):
            label_kind = "Start"
        elif isinstance(due_val, _date) and sched_date == due_val:
            label_kind = "Due"
        elif isinstance(start_val, _date) and sched_date == start_val:
            label_kind = "Start"

        text = f"{wo} - {title} - {label_kind}"
        day_items.setdefault(sched_date, []).append(text)

    # ---------- Normalize Overdue / Open tables (IsOpen = True, proper Overdue) ----------
    today = _date.today()

    def _pick_col_wo(df_in: pd.DataFrame, candidates: list[str]) -> str | None:
        if df_in is None or df_in.empty:
            return None
        lowmap = {str(c).lower(): c for c in df_in.columns}
        for c in candidates:
            if c in df_in.columns:
                return c
            lc = str(c).lower()
            if lc in lowmap:
                return lowmap[lc]
        return None

    # Work from whichever df is non-empty to locate shared column names
    base_df_for_cols = df_overdue if not df_overdue.empty else df_open_due_started

    due_col_all = _pick_col_wo(
        base_df_for_cols,
        ["Due date", "Due Date", "Due_Date", "due_date"],
    )
    isopen_col_all = _pick_col_wo(
        base_df_for_cols,
        ["IsOpen", "Is Open", "is_open", "ISOPEN"],
    )
    status_col_all = _pick_col_wo(
        base_df_for_cols,
        ["wo_status", "Status", "Status Name"],
    )

    def _is_open_mask(df_in: pd.DataFrame, isopen_col: str) -> pd.Series:
        s = df_in[isopen_col]
        if s.dtype == bool:
            return s.fillna(False)
        # treat non-zero numeric or truthy strings as open
        s_num = pd.to_numeric(s, errors="coerce")
        mask_num = s_num.fillna(0).ne(0)
        # also allow "true"/"yes" strings
        s_str = s.astype(str).str.strip().str.lower()
        mask_str = s_str.isin(["true", "yes", "y"])
        return mask_num | mask_str

    # --- Rebuild df_overdue: IsOpen = True AND DueDate < today ---
    if due_col_all and isopen_col_all and due_col_all in df_overdue.columns and isopen_col_all in df_overdue.columns:
        # Normalize due dates to midnight timestamps so we can safely compare (avoids datetime64 vs date TypeError)
        due_dt = pd.to_datetime(df_overdue[due_col_all], errors="coerce")
        today_ts = pd.Timestamp(today)
        mask_open_over = _is_open_mask(df_overdue, isopen_col_all)
        mask_due_past = due_dt.dt.normalize() < today_ts.normalize()
        df_overdue = df_overdue[mask_open_over & mask_due_past].copy()

    # --- Rebuild df_open_due_started: IsOpen = True (optionally Status in [Due, Started]) ---
    if isopen_col_all and isopen_col_all in df_open_due_started.columns:
        mask_open = _is_open_mask(df_open_due_started, isopen_col_all)

        if status_col_all and status_col_all in df_open_due_started.columns:
            status_norm = (
                df_open_due_started[status_col_all]
                .astype(str)
                .str.strip()
                .str.lower()
            )
            # keep only Due / Started if the column exists
            mask_status = status_norm.isin(["due", "started"])
            df_open_due_started = df_open_due_started[mask_open & mask_status].copy()
        else:
            df_open_due_started = df_open_due_started[mask_open].copy()

    # ---------- Style block for calendar tables ----------
    style_block = """
    <style>
    table.wo-cal {
        border-collapse: collapse;
        width: 100%;
        font-size: 0.8rem;
    }
    table.wo-cal th, table.wo-cal td {
        border: 1px solid #ccc;
        vertical-align: top;
        padding: 4px;
    }
    table.wo-cal th {
        background-color: #f0f0f0;
        text-align: center;
        font-weight: bold;
    }
    table.wo-cal ul {
        margin: 0;
        padding-left: 16px;
    }
    </style>
    """

    # ---------- Helper: build a calendar grid (HTML) ----------
    def _build_calendar_grid_html(title_txt: str, start_day: _date, num_days: int) -> str:
        end_day = start_day + timedelta(days=num_days - 1)
        days = [start_day + timedelta(days=i) for i in range(num_days)]
        weeks = math.ceil(num_days / 7)

        parts = []
        parts.append(f"<h3>{html.escape(title_txt)} ({start_day} to {end_day})</h3>")
        parts.append("<table class='wo-cal'>")

        # Header row (Mon–Sun)
        parts.append("<tr>")
        for name in ["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"]:
            parts.append(f"<th>{name}</th>")
        parts.append("</tr>")

        idx = 0
        for _w in range(weeks):
            parts.append("<tr>")
            for _d in range(7):
                if idx < num_days:
                    day = days[idx]
                    idx += 1
                    items = day_items.get(day, [])
                    cell_html = f"<div><strong>{day.strftime('%Y-%m-%d')}</strong></div>"

                    if items:
                        lis = [f"<li>{html.escape(str(text))}</li>" for text in items]
                        cell_html += "<ul>" + "".join(lis) + "</ul>"

                    parts.append(f"<td>{cell_html}</td>")
                else:
                    parts.append("<td>&nbsp;</td>")
            parts.append("</tr>")

        parts.append("</table>")
        return "".join(parts)

    # ---------- Helpers for PDF table pages (with pagination) ----------
    def _wrap_row_to_lines(c, text: str, max_width: float) -> list[str]:
        """Return wrapped lines for a row using stringWidth to respect max_width."""
        words = str(text).split()
        if not words:
            return [""]

        lines: list[str] = []
        current = ""
        for w in words:
            candidate = (current + " " + w).strip()
            if c.stringWidth(candidate, "Helvetica", 7) <= max_width:
                current = candidate
            else:
                if current:
                    lines.append(current)
                current = w
        if current:
            lines.append(current)
        return lines

    # ---------- Helper: build a calendar + 2 tables PDF ----------
    def _build_calendar_pdf_bytes(
        title_txt: str,
        start_day: _date,
        num_days: int,
        tbl_overdue: pd.DataFrame,
        tbl_open: pd.DataFrame,
    ) -> bytes:
        """
        Build a 7/30-day calendar PDF + paginated table pages (Overdue / Open).

        - Calendar header: Mon–Sun
        - Each date is placed under its *real* weekday column.
        - Tables use wrapped text and MAX_ROWS_PER_PAGE pagination.
        """
        from reportlab.platypus import Table, TableStyle, Paragraph
        from reportlab.lib.styles import ParagraphStyle

        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=landscape(letter))
        width, height = landscape(letter)

        margin = 36  # half-inch
        title_height = 24
        MAX_ROWS_PER_PAGE = 20

        # -------------------------------------------------
        # Page 1: Calendar
        # -------------------------------------------------
        c.setFont("Helvetica-Bold", 14)
        c.drawString(
            margin,
            height - margin,
            f"{title_txt} ({start_day} to {start_day + timedelta(days=num_days - 1)})",
        )

        # Grid area
        grid_top = height - margin - title_height
        grid_bottom = margin
        grid_height = grid_top - grid_bottom

        cols = 7  # Mon–Sun
        # number of calendar rows needed (weeks)
        rows = math.ceil((start_day.weekday() + num_days) / cols)
        cell_w = (width - 2 * margin) / cols
        cell_h = grid_height / rows

        # Weekday headers
        c.setFont("Helvetica-Bold", 10)
        weekday_labels = ["Mon", "Tue", "Wed", "Thu", "Fri", "Sat", "Sun"]
        for col_idx, name in enumerate(weekday_labels):
            x = margin + col_idx * cell_w + 2
            y = grid_top + 4
            c.drawString(x, y, name)

        # Build list of days
        days = [start_day + timedelta(days=i) for i in range(num_days)]

        # Draw grid + fill dates based on real weekday
        for row_i in range(rows):
            for col_i in range(cols):
                x = margin + col_i * cell_w
                y_top = grid_top - row_i * cell_h
                y_bottom = y_top - cell_h
                c.rect(x, y_bottom, cell_w, cell_h, stroke=1, fill=0)

        c.setFont("Helvetica", 7)
        for offset, day in enumerate(days):
            eff_idx = start_day.weekday() + offset  # 0 = Monday
            row_i = eff_idx // cols
            col_i = eff_idx % cols

            x = margin + col_i * cell_w
            y_top = grid_top - row_i * cell_h

            # Date label
            c.setFont("Helvetica-Bold", 8)
            c.drawString(x + 2, y_top - 10, day.strftime("%Y-%m-%d"))

            # Items for this day
            items = day_items.get(day, [])
            c.setFont("Helvetica", 7)
            text_y = y_top - 22
            max_lines = int((cell_h - 18) / 9)
            line_count = 0
            max_width_cell = cell_w - 8

            for raw in items:
                if line_count >= max_lines:
                    c.setFillColor(colors.black)
                    c.drawString(x + 4, text_y, "... more ...")
                    break

                c.setFillColor(colors.black)
                s = str(raw)
                words = s.split()
                line = ""

                for w in words:
                    test = (line + " " + w).strip()
                    if c.stringWidth(test, "Helvetica", 7) <= max_width_cell:
                        line = test
                    else:
                        if line_count >= max_lines:
                            c.setFillColor(colors.black)
                            c.drawString(x + 4, text_y, "... more ...")
                            line_count += 1
                            break
                        c.drawString(x + 4, text_y, line)
                        text_y -= 9
                        line_count += 1
                        line = w

                if line and line_count < max_lines:
                    c.drawString(x + 4, text_y, line)
                    text_y -= 9
                    line_count += 1

                if line_count >= max_lines:
                    break

            c.setFillColor(colors.black)

        c.showPage()

        # -------------------------------------------------
        # Paginated wrapped tables on subsequent pages
        # -------------------------------------------------
        para_style = ParagraphStyle(
            "tbl",
            fontName="Helvetica",
            fontSize=7,
            leading=8,
        )

        def _draw_table_pages(title: str, df: pd.DataFrame):
            if df is None or df.empty:
                return

            df_loc = df.copy()
            headers = list(df_loc.columns)
            n_cols = len(headers)
            col_widths = [(width - 2 * margin) / max(1, n_cols)] * n_cols
            avail_width = width - 2 * margin
            avail_height = height - margin - title_height - 12

            for start_idx in range(0, len(df_loc), MAX_ROWS_PER_PAGE):
                chunk = df_loc.iloc[start_idx:start_idx + MAX_ROWS_PER_PAGE]

                c.setFont("Helvetica-Bold", 14)
                c.drawString(margin, height - margin, title)

                data: list[list] = []
                header_cells = [Paragraph(str(h), para_style) for h in headers]
                data.append(header_cells)

                for _, row in chunk.iterrows():
                    cells = [Paragraph("" if pd.isna(v) else str(v), para_style) for v in row.tolist()]
                    data.append(cells)

                table = Table(data, colWidths=col_widths)
                table.setStyle(
                    TableStyle(
                        [
                            ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                            ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                            ("FONTSIZE", (0, 0), (-1, -1), 7),
                            ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ]
                    )
                )

                tw, th = table.wrapOn(c, avail_width, avail_height)
                table.drawOn(c, margin, height - margin - title_height - th)
                c.showPage()

        # -------------------------------------------------
        # Pages 2–3+: Overdue / Open tables
        # -------------------------------------------------
        _draw_table_pages("Overdue Work Orders", tbl_overdue)
        _draw_table_pages("Open Work Orders (IsOpen = True)", tbl_open)

        c.save()
        buf.seek(0)
        return buf.getvalue()

    # ---------- Build 7-day and 30-day calendars from selected start date ----------
    cal7_html = _build_calendar_grid_html("Next 7 Days", cal_start, 7)
    cal30_html = _build_calendar_grid_html("Next 30 Days", cal_start, 30)

    st.markdown(style_block, unsafe_allow_html=True)
    col7, col30 = st.columns(2)
    with col7:
        st.markdown(cal7_html, unsafe_allow_html=True)
    with col30:
        st.markdown(cal30_html, unsafe_allow_html=True)

    # ---------- Prep subset tables for PDF ----------
    overdue_cols_pdf = [
        c for c in
        ["WORKORDER", "TITLE", "Description", "Due date", "assignee", "team_assigned"]
        if c in df_overdue.columns
    ]
    open_cols_pdf = [
        c for c in
        ["WORKORDER", "TITLE", "Description", "Created on", "Due date", "assignee", "team_assigned"]
        if c in df_open_due_started.columns
    ]

    tbl_overdue_pdf = df_overdue[overdue_cols_pdf].copy() if overdue_cols_pdf else df_overdue.copy()
    tbl_open_pdf = df_open_due_started[open_cols_pdf].copy() if open_cols_pdf else df_open_due_started.copy()

    # ---------- Build PDFs and expose as two separate downloads ----------
    pdf_7 = _build_calendar_pdf_bytes("Next 7 Days", cal_start, 7, tbl_overdue_pdf, tbl_open_pdf)
    pdf_30 = _build_calendar_pdf_bytes("Next 30 Days", cal_start, 30, tbl_overdue_pdf, tbl_open_pdf)

    st.download_button(
        "⬇️ 7-Day Calendar + Tables (PDF)",
        data=pdf_7,
        file_name=f"WO_Calendar_7day_{cal_start}.pdf",
        mime="application/pdf",
        key="dl_wo_cal_7_pdf",
    )

    st.download_button(
        "⬇️ 30-Day Calendar + Tables (PDF)",
        data=pdf_30,
        file_name=f"WO_Calendar_30day_{cal_start}.pdf",
        mime="application/pdf",
        key="dl_wo_cal_30_pdf",
    )


import io
import re
import os
from datetime import date, datetime
from pathlib import Path
from typing import Optional, List

import numpy as np
import pandas as pd
import streamlit as st

from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.platypus import Table, TableStyle
from reportlab.lib.utils import ImageReader


# =====================================================
# Globals for logo (set inside _render_wo_forms_pdf_page)
# =====================================================

LOGO_PATH: Optional[Path] = None


# ========================================
# Helpers
# ========================================

def _safe_str(val) -> str:
    """Return '' for None / NaN / NA, else string."""
    if val is None:
        return ""
    try:
        if pd.isna(val):
            return ""
    except TypeError:
        pass
    return str(val)


def _find(df: pd.DataFrame, *cands: str) -> Optional[str]:
    """
    Case-insensitive, space/underscore-insensitive column finder.
    Used with explicit candidate names only (no fuzzy 'status' searching).
    """
    if df is None or df.empty:
        return None

    lower_map = {str(c).strip().lower(): c for c in df.columns}

    # exact lowercase match
    for cand in cands:
        if not cand:
            continue
        key = cand.strip().lower()
        if key in lower_map:
            return lower_map[key]

    # ignore spaces/underscores
    norm_map = {
        str(c).strip().lower().replace(" ", "").replace("_", ""): c
        for c in df.columns
    }
    for cand in cands:
        if not cand:
            continue
        key = cand.strip().lower().replace(" ", "").replace("_", "")
        if key in norm_map:
            return norm_map[key]

    return None


def _to_xlsx_bytes(df: pd.DataFrame, sheet_name: str = "WorkOrders") -> bytes:
    """Convert a DataFrame to XLSX bytes."""
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False, sheet_name=sheet_name)
    output.seek(0)
    return output.getvalue()


def _get_step_progress_col(df: pd.DataFrame) -> Optional[str]:
    """
    Helper for the renamed STATUS column.
    Avoids grabbing generic 'STATUS' to keep it distinct from wo_status.
    """
    col = _find(df, "step_progress")
    if col:
        return col
    # Legacy fallback if it ever appears
    col = _find(df, "STATUSstep_progress")
    return col


def _extract_unique_teams(series: pd.Series) -> List[str]:
    """
    Split multi-team strings (comma/semicolon) into unique trimmed names.
    """
    teams = set()
    for v in series.dropna():
        text = str(v).replace(";", ",")
        for part in text.split(","):
            t = part.strip()
            if t:
                teams.add(t)
    return sorted(teams)


@st.cache_data(show_spinner=True)
def load_data(path: Path | str) -> tuple[pd.DataFrame, Optional[str]]:
    """Load the main WO/PO/Trans parquet and normalize the Completed On date column."""
    df = pd.read_parquet(path)
    date_col = _find(df, "COMPLETED ON", "Completed On", "completed_on")
    if date_col:
        df[date_col] = pd.to_datetime(df[date_col], errors="coerce")
    return df, date_col


@st.cache_data(show_spinner=False)
def load_users(path: Optional[Path] | str) -> pd.DataFrame:
    """Load Users.parquet if it exists; otherwise return empty DF."""
    if not path:
        return pd.DataFrame()
    p = Path(path)
    if not p.is_file():
        return pd.DataFrame()
    return pd.read_parquet(p)


def _bool_flag(dff: pd.DataFrame, col: str) -> bool:
    """Coerce a KPI column to a single boolean flag (any row true / >0)."""
    if col not in dff.columns:
        return False
    s = dff[col]
    if s.dtype == bool:
        return bool(s.any())
    # treat non-zero as True
    vals = pd.to_numeric(s, errors="coerce").fillna(0)
    return bool((vals != 0).any())


def _max_int(dff: pd.DataFrame, col: str) -> int:
    """Return max integer value for a KPI column (or 0)."""
    if col not in dff.columns:
        return 0
    vals = pd.to_numeric(dff[col], errors="coerce").fillna(0)
    if len(vals) == 0:
        return 0
    return int(vals.max())


def _has_real_date(dff: pd.DataFrame, col: str) -> Optional[bool]:
    """
    Return True if any value in `col` is a real date (not blank/0/1900-01-01).
    Return False if present but all blanks/invalid.
    Return None if column not present.
    """
    if col not in dff.columns:
        return None

    s = pd.to_datetime(dff[col], errors="coerce")
    if s.isna().all():
        return False

    # treat 1900-01-01 and earlier as "no date" (Power Query 0 / 1/0/1900 junk)
    good = s.notna() & (s.dt.year > 1900)
    return bool(good.any())


def _summarize_kpis(dff: pd.DataFrame) -> list[str]:
    """
    Build a list of human-readable KPI lines for a workorder/group.

    Uses:
      • rea_pre  -> Work type (Reactive / Preventive / Other)
      • IsOverDue / DaysOverDue
      • ComOverDue / DaysOverDue_Com
      • Planned Start Date
      • Due date
      • IsOpen / DaysOpen
    """
    lines: list[str] = []

    # ---------- Work type from rea_pre ----------
    type_label = None

    rea_pre_col = _find(dff, "rea_pre", "Rea_Pre", "REA_PRE")
    if rea_pre_col and rea_pre_col in dff.columns:
        raw = str(dff[rea_pre_col].iloc[0]).strip().lower()
        if raw.startswith("rea"):
            type_label = "Reactive"
        elif raw.startswith("pre"):
            type_label = "Preventive"
        elif raw:
            type_label = raw.capitalize()

    # mild fallback if rea_pre missing/blank
    if not type_label:
        wt_col = _find(dff, "Work Type", "Type")
        if wt_col and wt_col in dff.columns:
            wt_val = str(dff[wt_col].iloc[0]).strip()
            if wt_val:
                type_label = wt_val

    if type_label:
        lines.append(f"• Work type: {type_label}")

    # ---------- Overdue / Completed overdue ----------
    if "IsOverDue" in dff.columns:
        is_overdue = _bool_flag(dff, "IsOverDue")
        days_over = _max_int(dff, "DaysOverDue") if "DaysOverDue" in dff.columns else 0
        if is_overdue:
            lines.append(
                f"• Overdue? YES — {days_over} days past due"
                if days_over else "• Overdue? YES"
            )
        else:
            lines.append("• Overdue? No")

    if "ComOverDue" in dff.columns:
        com_over = _bool_flag(dff, "ComOverDue")
        days_com_over = (
            _max_int(dff, "DaysOverDue_Com")
            if "DaysOverDue_Com" in dff.columns
            else 0
        )
        if com_over:
            lines.append(
                f"• Completed overdue? YES — finished {days_com_over} days late"
                if days_com_over else "• Completed overdue? YES"
            )
        else:
            lines.append("• Completed overdue? No")

    # ---------- Planned start date ----------
    planned_col = _find(dff, "Planned Start Date", "Planned Start", "Start Date")
    has_start = _has_real_date(dff, planned_col) if planned_col else None
    if has_start is not None:
        lines.append(
            "• Has planned start date? Yes" if has_start else "• Has planned start date? No"
        )

    # ---------- Due date ----------
    due_col = _find(dff, "Due date", "Due Date", "Due")
    has_due = _has_real_date(dff, due_col) if due_col else None
    if has_due is not None:
        lines.append("• Has due date? Yes" if has_due else "• Has due date? No")

    # ---------- Currently open ----------
    if "IsOpen" in dff.columns:
        is_open = _bool_flag(dff, "IsOpen")
        days_open = _max_int(dff, "DaysOpen") if "DaysOpen" in dff.columns else 0
        if is_open:
            lines.append(
                f"• Currently open? YES — open for {days_open} days"
                if days_open else "• Currently open? YES"
            )
        else:
            lines.append(
                f"• Currently open? No — was open {days_open} days"
                if days_open else "• Currently open? No"
            )

    # Ensure only plain strings are returned
    return [str(x) for x in lines if isinstance(x, (str, int, float))]


def _wrap_text(text: str, max_width: float, font_name: str = "Helvetica", font_size: float = 10) -> list[str]:
    """Wrap a string into multiple lines so each line fits within max_width."""
    if not text:
        return []
    words = str(text).split()
    lines: list[str] = []
    current = ""

    for w in words:
        trial = f"{current} {w}".strip()
        width = pdfmetrics.stringWidth(trial, font_name, font_size)
        if width <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = w
    if current:
        lines.append(current)
    return lines


def _format_date(val) -> str:
    """
    Nice date formatting; handles pandas Timestamps, strings, NaN, 0, 1/0/1900, etc.

    Output: mm/dd/yyyy or '' if it's effectively blank.
    """

    # Treat obvious nulls as blank
    if val is None:
        return ""

    # NaN / NA / NaT
    try:
        if pd.isna(val):
            return ""
    except TypeError:
        pass

    # If it's already a Timestamp or datetime-like
    if isinstance(val, (pd.Timestamp, datetime, np.datetime64)):
        ts = pd.to_datetime(val, errors="coerce")
        if ts is pd.NaT:
            return ""
        # Treat anything before year 2000 as "no real date" in this system
        if ts.year < 2000:
            return ""
        return ts.strftime("%m/%d/%Y")

    # Numbers (Excel/placeholder 0, etc.) -> treat as blank
    if isinstance(val, (int, float)):
        if val <= 0:
            return ""
        return ""

    # Strings: clean and try to parse
    s = str(val).strip()
    if not s:
        return ""
    if s in {"NaT", "0", "1/0/1900"}:
        return ""

    ts = pd.to_datetime(s, errors="coerce")
    if ts is not pd.NaT and ts.year >= 2000:
        return ts.strftime("%m/%d/%Y")

    # Fallback
    return s


def _get_asset_info(row: pd.Series) -> dict:
    """
    Asset tag info now comes directly from the WO row itself.
    """

    def pick(*cols: str) -> str:
        """Return the first non-empty, non-NA string from the given columns."""
        for col in cols:
            if col in row.index:
                val = row[col]
                try:
                    if pd.isna(val):
                        continue
                except TypeError:
                    if val is None:
                        continue
                s = str(val).strip()
                if s != "":
                    return s
        return ""

    info = {
        "Asset": pick("ASSET", "Asset", "asset", "Name"),
        "Manufacturer": pick("make", "Make", "Manufacturer"),
        "Model": pick("model", "Model"),
        "Serial": pick("serial", "Serial", "Serial Number"),
        "Year": pick("year", "Year"),
        "Type": pick("Type", "types"),
        "Department": pick("Department", "Dept"),
        "Location": pick("Location"),
    }
    return info


# ========================================
# PDF builder — summary + per-WO pages
# ========================================

def _make_multi_wo_pdf(df: pd.DataFrame) -> bytes:
    """
    Build a multi-page PDF:
      • Page 1: Summary (with logo, grouped blocks, clickable to detail pages)
      • Subsequent pages: one page per Workorder_group/WORKORDER
    """
    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter

    margin_left = 40
    margin_right = 40
    text_width = width - margin_left - margin_right

    def _wrap_cell(val, col_width_pts, font_name="Helvetica", font_size=8) -> str:
        txt = _safe_str(val)
        if not txt:
            return ""
        lines = _wrap_text(txt, col_width_pts - 4, font_name, font_size)
        return "\n".join(lines)

    # Decide grouping: Workorder_group if present, else WORKORDER
    group_col = _find(df, "Workorder_group", "workorder_group")
    if group_col and df[group_col].notna().any():
        groups = df[group_col].dropna().unique().tolist()

        def subset(gval):
            return df[df[group_col] == gval].copy()
    else:
        group_col = None
        groups = df["WORKORDER"].dropna().unique().tolist()

        def subset(gval):
            return df[df["WORKORDER"] == gval].copy()

    # Group workorders by step_progress for summary dividers
    step_group_col = _get_step_progress_col(df)
    status_to_groups: dict[str, list] = {}

    for gval in groups:
        dff = subset(gval)
        if dff.empty:
            continue
        if step_group_col and step_group_col in dff.columns:
            step_val = _safe_str(dff[step_group_col].iloc[0])
        else:
            step_val = ""
        status_to_groups.setdefault(step_val, []).append(gval)

    preferred_order = ["open", "in progress", "on hold",
                       "done", "closed", "complete", "completed"]

    def _status_sort_key(s):
        s_norm = str(s).strip().lower()
        if s_norm in preferred_order:
            return (0, preferred_order.index(s_norm))
        return (1, s_norm)

    ordered_statuses = sorted(status_to_groups.keys(), key=_status_sort_key)

    # Common column locators
    title_col     = _find(df, "TITLE", "Title")
    desc_col      = _find(df, "Description", "DESC")
    asset_col     = _find(df, "ASSET", "Asset")
    assignee_col  = _find(df, "assignee", "Assigned To")
    team_col      = _find(df, "team_assigned", "team", "Team")
    wo_status_col = _find(df, "wo_status", "WO_STATUS")

    created_col   = _find(df, "Created on", "Created On", "Created date", "Created Date")
    start_col     = _find(df, "Planned Start Date", "Start Date", "Planned Start")
    due_col       = _find(df, "Due date", "Due Date", "Due")
    completed_col = _find(df, "COMPLETED ON", "Completed on", "Completed On")

    # ---------- PAGE 1: SUMMARY ----------
    c.setTitle("Work Orders Summary")
    summary_bookmark = "SUMMARY"
    c.bookmarkPage(summary_bookmark)

    # Logo on summary
    try:
        if LOGO_PATH and os.path.exists(LOGO_PATH):
            logo_img = ImageReader(str(LOGO_PATH))
            c.drawImage(
                logo_img,
                width - 1.6 * inch,
                height - 0.8 * inch,
                width=1.4 * inch,
                preserveAspectRatio=True,
                mask="auto",
            )
    except Exception as e:
        st.write("Error drawing logo on summary:", repr(e))

    c.setFont("Helvetica-Bold", 16)
    c.drawString(margin_left, height - 60, "Work Orders Summary")

    y = height - 90
    c.setFont("Helvetica", 10)

    for status_val in ordered_statuses:
        groups_in_status = status_to_groups[status_val]

        # Divider line for this step_progress group
        if status_val:
            if y < 100:
                c.showPage()
                c.bookmarkPage(summary_bookmark)
                try:
                    if LOGO_PATH and os.path.exists(LOGO_PATH):
                        logo_img = ImageReader(str(LOGO_PATH))
                        c.drawImage(
                            logo_img,
                            width - 1.6 * inch,
                            height - 0.8 * inch,
                            width=1.4 * inch,
                            preserveAspectRatio=True,
                            mask="auto",
                        )
                except Exception:
                    pass
                c.setFont("Helvetica-Bold", 14)
                c.drawString(margin_left, height - 60, "Work Orders Summary (cont.)")
                y = height - 90

            c.setFont("Helvetica-Bold", 12)
            c.drawString(margin_left, y, f"Step Progress: {status_val}")
            y -= 16
            c.setFont("Helvetica", 10)

        for gval in groups_in_status:
            dff = subset(gval)
            if dff.empty:
                continue
            header = dff.iloc[0]

            wo       = _safe_str(header.get("WORKORDER"))
            title    = _safe_str(header.get(title_col)) if title_col else ""
            asset    = _safe_str(header.get(asset_col)) if asset_col else ""
            assignee = _safe_str(header.get(assignee_col)) if assignee_col else ""
            team     = _safe_str(header.get(team_col)) if team_col else ""
            status   = _safe_str(header.get(wo_status_col)) if wo_status_col else ""

            # PO summary
            po_col   = _find(dff, "PO")
            step_col = _get_step_progress_col(dff)
            po_done_col = completed_col
            po_status_str = "-"

            if po_col and po_col in dff.columns:
                po_rows = dff[dff[po_col].notna()].copy()
                if not po_rows.empty:
                    po_step = ""
                    if step_col and step_col in po_rows.columns:
                        vals = (
                            po_rows[step_col]
                            .dropna()
                            .astype(str)
                            .str.strip()
                        )
                        if len(vals):
                            po_step = vals.iloc[0]

                    po_done_str = ""
                    if po_done_col and po_done_col in po_rows.columns:
                        dates = pd.to_datetime(
                            po_rows[po_done_col], errors="coerce"
                        ).dropna()
                        if not dates.empty:
                            po_done_str = _format_date(dates.max())

                    if po_step and po_done_str:
                        po_status_str = f"{po_step} — {po_done_str}"
                    elif po_step:
                        po_status_str = po_step
                    elif po_done_str:
                        po_status_str = po_done_str

            if y < 100:
                c.showPage()
                c.bookmarkPage(summary_bookmark)
                try:
                    if LOGO_PATH and os.path.exists(LOGO_PATH):
                        logo_img = ImageReader(str(LOGO_PATH))
                        c.drawImage(
                            logo_img,
                            width - 1.6 * inch,
                            height - 0.8 * inch,
                            width=1.4 * inch,
                            preserveAspectRatio=True,
                            mask="auto",
                        )
                except Exception:
                    pass
                c.setFont("Helvetica-Bold", 14)
                c.drawString(margin_left, height - 60, "Work Orders Summary (cont.)")
                y = height - 90
                c.setFont("Helvetica", 10)

            # Line 1: WO — Title (WO clickable)
            c.setFont("Helvetica-Bold", 11)
            header_line = f"{wo} — {title}"
            c.drawString(margin_left, y, header_line)
            wo_bookmark = f"WO_{wo}"
            wo_width = pdfmetrics.stringWidth(wo, "Helvetica-Bold", 11)
            c.linkRect(
                "",
                wo_bookmark,
                (margin_left, y - 2, margin_left + wo_width + 4, y + 10),
                relative=0,
                thickness=0,
            )
            y -= 14

            # Line 2: Asset
            c.setFont("Helvetica", 10)
            for line in _wrap_text(f"Asset: {asset or '-'}", text_width, "Helvetica", 10):
                c.drawString(margin_left + 10, y, line)
                y -= 12

            # Line 3: Assignee / Team
            line3 = f"Assignee: {assignee or '-'}    Team: {team or '-'}"
            for line in _wrap_text(line3, text_width, "Helvetica", 10):
                c.drawString(margin_left + 10, y, line)
                y -= 12

            # Line 4: WO Status / PO summary
            line4 = f"WO Status: {status or '-'}    PO: {po_status_str or '-'}"
            for line in _wrap_text(line4, text_width, "Helvetica", 10):
                c.drawString(margin_left + 10, y, line)
                y -= 12

            y -= 6  # spacer between WOs

    # ---------- DETAIL PAGES PER WO ----------
    for gval in groups:
        dff = subset(gval)
        if dff.empty:
            continue
        header = dff.iloc[0]

        c.showPage()

        # Bookmark for this workorder detail page
        wo = _safe_str(header.get("WORKORDER"))
        wo_bookmark = f"WO_{wo}"
        c.bookmarkPage(wo_bookmark)

        # Logo top-right
        try:
            if LOGO_PATH and os.path.exists(LOGO_PATH):
                logo_img = ImageReader(str(LOGO_PATH))
                c.drawImage(
                    logo_img,
                    width - 1.6 * inch,
                    height - 0.8 * inch,
                    width=1.4 * inch,
                    preserveAspectRatio=True,
                    mask="auto",
                )
        except Exception as e:
            st.write("Error drawing logo on detail page:", repr(e))

        # "Return to Summary" text under logo with link
        c.setFont("Helvetica-Bold", 9)
        return_label = "Return to Summary"
        return_x2 = width - margin_right
        return_y = height - 85
        c.drawRightString(return_x2, return_y, return_label)

        text_w = pdfmetrics.stringWidth(return_label, "Helvetica-Bold", 9)
        c.linkRect(
            "",
            summary_bookmark,
            (
                return_x2 - text_w - 4,
                return_y - 2,
                return_x2 + 2,
                return_y + 10,
            ),
            relative=0,
            thickness=0,
        )

        y = height - 110

        # Header: Work Order, Title, Description
        title = _safe_str(header.get(title_col)) if title_col else ""
        desc  = _safe_str(header.get(desc_col)) if desc_col else ""

        c.setFont("Helvetica-Bold", 12)
        c.drawString(margin_left, y, f"Work Order: {wo}")
        y -= 16

        c.setFont("Helvetica-Bold", 11)
        c.drawString(margin_left, y, "Title:")
        c.setFont("Helvetica", 11)
        for line in _wrap_text(title, text_width, "Helvetica", 11):
            c.drawString(margin_left + 45, y, line)
            y -= 14
        y -= 4

        c.setFont("Helvetica-Bold", 11)
        c.drawString(margin_left, y, "Description:")
        y -= 14
        c.setFont("Helvetica", 10)
        for line in _wrap_text(desc, text_width, "Helvetica", 10):
            if y < 80:
                c.showPage()
                y = height - 60
                c.setFont("Helvetica", 10)
            c.drawString(margin_left, y, line)
            y -= 12
        y -= 8

        # Dates line
        created = _format_date(header.get(created_col)) if created_col else ""
        start   = _format_date(header.get(start_col)) if start_col else ""
        due     = _format_date(header.get(due_col)) if due_col else ""

        c.setFont("Helvetica", 10)
        date_line = f"Created: {created or '-'}      Start: {start or '-'}      Due: {due or '-'}"
        c.drawString(margin_left, y, date_line)
        y -= 18

        # Asset tag 2-column
        asset_info = _get_asset_info(header)

        left_lines = [
            f"Asset: {asset_info['Asset']}",
            f"Manufacturer: {asset_info['Manufacturer']}",
            f"Model: {asset_info['Model']}",
            f"Serial Number: {asset_info['Serial']}",
            f"Year: {asset_info['Year']}",
        ]
        right_lines = [
            f"Type: {asset_info['Type']}",
            f"Department: {asset_info['Department']}",
            f"Location: {asset_info['Location']}",
        ]

        c.setFont("Helvetica-Bold", 11)
        c.drawString(margin_left, y, "Asset Tag")
        y -= 14
        c.setFont("Helvetica", 10)

        left_x = margin_left
        right_x = margin_left + (text_width / 2) + 10
        max_left_width = (text_width / 2) - 10
        max_right_width = (text_width / 2) - 10

        y_left = y
        for ln in left_lines:
            for wline in _wrap_text(ln, max_left_width, "Helvetica", 10):
                c.drawString(left_x, y_left, wline)
                y_left -= 12

        y_right = y
        for ln in right_lines:
            for wline in _wrap_text(ln, max_right_width, "Helvetica", 10):
                c.drawString(right_x, y_right, wline)
                y_right -= 12

        y = min(y_left, y_right) - 10

        # Assigned / Team
        assignee = _safe_str(header.get(assignee_col)) if assignee_col else ""
        team     = _safe_str(header.get(team_col)) if team_col else ""

        c.setFont("Helvetica", 10)
        c.drawString(margin_left, y, f"Assigned: {assignee or '-'}")
        y -= 12
        c.drawString(margin_left, y, f"Team: {team or '-'}")
        y -= 18

        # Status & KPI Summary
        c.setFont("Helvetica-Bold", 11)
        c.drawString(margin_left, y, "Status & KPI Summary")
        y -= 14
        c.setFont("Helvetica", 10)

        kpi_lines = _summarize_kpis(dff)
        completed = _format_date(header.get(completed_col)) if completed_col else ""
        patched_lines: list[str] = []
        for line in kpi_lines:
            if line.strip().lower().startswith("• currently open?"):
                if completed:
                    patched_lines.append(f"{line}    Completed Date: {completed}")
                else:
                    patched_lines.append(line)
            else:
                patched_lines.append(line)

        for line in patched_lines:
            for wline in _wrap_text(line, text_width, "Helvetica", 10):
                if y < 80:
                    c.showPage()
                    y = height - 60
                    c.setFont("Helvetica", 10)
                c.drawString(margin_left, y, wline)
                y -= 12
        y -= 10

        # ----- Determine rank-based masks for PO vs Transactions -----
        rank_col = _find(dff, "RANK", "Rank", "rank", "_rank")
        if rank_col and rank_col in dff.columns:
            rank_s = dff[rank_col]
            rank_s_str = rank_s.astype(str).str.strip().str.lower()
            rank_s_num = pd.to_numeric(rank_s, errors="coerce")

            # PO rows: text "po" OR numeric 2
            po_mask = (rank_s_str == "po") | (rank_s_num == 2)

            # Transactions rows: text starting with "tran" OR numeric 3
            tx_mask = rank_s_str.str.startswith("tran") | (rank_s_num == 3)
        else:
            po_mask = pd.Series(True, index=dff.index)
            tx_mask = pd.Series(True, index=dff.index)

        pn_col_global = _find(dff, "P/N", "P N", "PN")

        # ---------- PO table ----------
        po_col       = _find(dff, "PO")
        step_col     = _find(dff, "step_progress", "STATUSstep_progress", "Step", "Status")
        qty_col      = _find(dff, "QUANTITY RECEIVED", "Quantity received", "Qty")
        vendor_col   = _find(dff, "Vendors", "Vendor")
        cost_col     = _find(dff, "Total Parts Cost", "ITEM COST", "TOTAL ITEM COST", "Total cost")
        po_done_col  = completed_col

        po_rows = dff[po_mask & (dff[po_col].notna() if po_col else True)] if po_col else dff[po_mask]

        if po_col and not po_rows.empty:
            po_col_widths = [50, 80, 55, 35, 190, 60, 62]

            if y < 80:
                c.showPage()
                y = height - 60
            c.setFont("Helvetica-Bold", 11)
            c.drawString(margin_left, y, "Purchase Orders")
            y -= 16

            po_data = [["PO", "P/N", "Step", "Qty", "Vendor", "Cost", "Completed Date"]]

            for _, row in po_rows.iterrows():
                po_val   = _wrap_cell(row.get(po_col),        po_col_widths[0])
                pn_val   = _wrap_cell(row.get(pn_col_global), po_col_widths[1]) if pn_col_global else ""
                step_val = _wrap_cell(row.get(step_col),      po_col_widths[2]) if step_col   else ""
                qty_val  = _safe_str(row.get(qty_col)) if qty_col else ""
                vend_val = _wrap_cell(row.get(vendor_col),    po_col_widths[4]) if vendor_col else ""

                if cost_col:
                    cost_raw = row.get(cost_col)
                    if pd.notna(cost_raw):
                        try:
                            cost_val = f"{float(cost_raw):,.2f}"
                        except Exception:
                            cost_val = _safe_str(cost_raw)
                    else:
                        cost_val = ""
                else:
                    cost_val = ""

                done_val = _format_date(row.get(po_done_col)) if po_done_col else ""

                po_data.append([
                    po_val,
                    pn_val,
                    step_val,
                    qty_val,
                    vend_val,
                    cost_val,
                    done_val,
                ])

            po_table = Table(
                po_data,
                colWidths=po_col_widths,
            )
            po_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.green),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                ("ALIGN", (3, 1), (3, -1), "RIGHT"),
                ("ALIGN", (5, 1), (5, -1), "RIGHT"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("LEADING", (0, 0), (-1, -1), 9),
            ]))

            tw, th = po_table.wrapOn(c, width, height)
            if y - th < 40:
                c.showPage()
                y = height - 60
            po_table.drawOn(c, margin_left, y - th)
            y = y - th - 12

        # ---------- Transactions table ----------
        pn_col        = pn_col_global
        part_loc_col  = _find(dff, "part_loc", "Part Location")
        part_area_col = _find(dff, "part_area", "Area")
        tx_qty_col    = qty_col

        tx_rows = dff[tx_mask & (dff[pn_col].notna() if pn_col else True)] if pn_col else dff[tx_mask]

        if pn_col and not tx_rows.empty:
            tx_col_widths = [210, 40, 130, 130]

            if y < 80:
                c.showPage()
                y = height - 60
            c.setFont("Helvetica-Bold", 11)
            c.drawString(margin_left, y, "Transactions")
            y -= 16

            tx_data = [["P/N", "Qty", "Part Loc", "Part Area"]]

            for _, row in tx_rows.iterrows():
                pn_val   = _wrap_cell(row.get(pn_col),       tx_col_widths[0])
                qty_val  = _safe_str(row.get(tx_qty_col)) if tx_qty_col else ""
                loc_val  = _wrap_cell(row.get(part_loc_col), tx_col_widths[2]) if part_loc_col  else ""
                area_val = _wrap_cell(row.get(part_area_col),tx_col_widths[3]) if part_area_col else ""

                tx_data.append([pn_val, qty_val, loc_val, area_val])

            tx_table = Table(
                tx_data,
                colWidths=tx_col_widths,
            )
            tx_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.green),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
                ("ALIGN", (1, 1), (1, -1), "RIGHT"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("LEADING", (0, 0), (-1, -1), 9),
            ]))

            tw, th = tx_table.wrapOn(c, width, height)
            if y - th < 40:
                c.showPage()
                y = height - 60
            tx_table.drawOn(c, margin_left, y - th)
            y = y - th - 12

    c.save()
    buf.seek(0)
    return buf.getvalue()


def _make_single_wo_pdf(df: pd.DataFrame, workorder: Optional[str] = None) -> bytes:
    """
    Build a PDF for a *single* work order (no summary page).
    Layout matches the detail pages from _make_multi_wo_pdf.
    """

    if df is None or df.empty:
        raise ValueError("No data provided to _make_single_wo_pdf.")

    dff = df.copy()

    if workorder is not None:
        wo_col = "WORKORDER"
        if wo_col not in dff.columns:
            wo_col = _find(dff, "WORKORDER", "Workorder", "workorder")

        if not wo_col or wo_col not in dff.columns:
            raise ValueError("Cannot find WORKORDER column for single-WO PDF.")

        dff = dff[dff[wo_col].astype(str) == str(workorder)].copy()

        if dff.empty:
            raise ValueError(f"No rows found for workorder {workorder!r}.")

    group_col = _find(dff, "Workorder_group", "workorder_group")
    if group_col and not dff[group_col].isna().all():
        grp_val = dff[group_col].iloc[0]
        dff = df[df[group_col] == grp_val].copy()
        if dff.empty:
            raise ValueError(f"No rows found for workorder group {grp_val!r}.")

    if dff is None or dff.empty:
        raise ValueError("No rows in DataFrame for single WO PDF.")

    header = dff.iloc[0]

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter

    margin_left = 40
    margin_right = 40
    text_width = width - margin_left - margin_right

    title_col     = _find(dff, "TITLE", "Title")
    desc_col      = _find(dff, "Description", "DESC")
    assignee_col  = _find(dff, "assignee", "Assigned To")
    team_col      = _find(dff, "team_assigned", "team", "Team")
    created_col   = _find(dff, "Created on", "Created On", "Created date", "Created Date")
    start_col     = _find(dff, "Planned Start Date", "Start Date", "Planned Start")
    due_col       = _find(dff, "Due date", "Due Date", "Due")
    completed_col = _find(dff, "COMPLETED ON", "Completed on", "Completed On")

    def _wrap_cell(val, col_width_pts, font_name="Helvetica", font_size=8) -> str:
        txt = _safe_str(val)
        if not txt:
            return ""
        lines = _wrap_text(txt, col_width_pts - 4, font_name, font_size)
        return "\n".join(lines)

    c.setTitle("Work Order Detail")

    try:
        if LOGO_PATH and os.path.exists(LOGO_PATH):
            logo_img = ImageReader(str(LOGO_PATH))
            c.drawImage(
                logo_img,
                width - 1.6 * inch,
                height - 0.8 * inch,
                width=1.4 * inch,
                preserveAspectRatio=True,
                mask="auto",
            )
    except Exception as e:
        st.write("Error drawing logo on single-WO detail page:", repr(e))

    y = height - 110

    wo = _safe_str(header.get("WORKORDER"))
    title = _safe_str(header.get(title_col)) if title_col else ""
    desc  = _safe_str(header.get(desc_col)) if desc_col else ""

    c.setFont("Helvetica-Bold", 12)
    c.drawString(margin_left, y, f"Work Order: {wo}")
    y -= 16

    c.setFont("Helvetica-Bold", 11)
    c.drawString(margin_left, y, "Title:")
    c.setFont("Helvetica", 11)
    for line in _wrap_text(title, text_width, "Helvetica", 11):
        c.drawString(margin_left + 45, y, line)
        y -= 14
    y -= 4

    c.setFont("Helvetica-Bold", 11)
    c.drawString(margin_left, y, "Description:")
    y -= 14
    c.setFont("Helvetica", 10)
    for line in _wrap_text(desc, text_width, "Helvetica", 10):
        if y < 80:
            c.showPage()
            y = height - 60
            c.setFont("Helvetica", 10)
        c.drawString(margin_left, y, line)
        y -= 12
    y -= 8

    created = _format_date(header.get(created_col)) if created_col else ""
    start   = _format_date(header.get(start_col)) if start_col else ""
    due     = _format_date(header.get(due_col)) if due_col else ""

    c.setFont("Helvetica", 10)
    date_line = f"Created: {created or '-'}      Start: {start or '-'}      Due: {due or '-'}"
    c.drawString(margin_left, y, date_line)
    y -= 18

    asset_info = _get_asset_info(header)

    left_lines = [
        f"Asset: {asset_info['Asset']}",
        f"Manufacturer: {asset_info['Manufacturer']}",
        f"Model: {asset_info['Model']}",
        f"Serial Number: {asset_info['Serial']}",
        f"Year: {asset_info['Year']}",
    ]
    right_lines = [
        f"Type: {asset_info['Type']}",
        f"Department: {asset_info['Department']}",
        f"Location: {asset_info['Location']}",
    ]

    c.setFont("Helvetica-Bold", 11)
    c.drawString(margin_left, y, "Asset Tag")
    y -= 14
    c.setFont("Helvetica", 10)

    left_x = margin_left
    right_x = margin_left + (text_width / 2) + 10
    max_left_width = (text_width / 2) - 10
    max_right_width = (text_width / 2) - 10

    y_left = y
    for ln in left_lines:
        for wline in _wrap_text(ln, max_left_width, "Helvetica", 10):
            c.drawString(left_x, y_left, wline)
            y_left -= 12

    y_right = y
    for ln in right_lines:
        for wline in _wrap_text(ln, max_right_width, "Helvetica", 10):
            c.drawString(right_x, y_right, wline)
            y_right -= 12

    y = min(y_left, y_right) - 10

    assignee = _safe_str(header.get(assignee_col)) if assignee_col else ""
    team     = _safe_str(header.get(team_col)) if team_col else ""

    c.setFont("Helvetica", 10)
    c.drawString(margin_left, y, f"Assigned: {assignee or '-'}")
    y -= 12
    c.drawString(margin_left, y, f"Team: {team or '-'}")
    y -= 18

    c.setFont("Helvetica-Bold", 11)
    c.drawString(margin_left, y, "Status & KPI Summary")
    y -= 14
    c.setFont("Helvetica", 10)

    kpi_lines = _summarize_kpis(dff)
    completed = _format_date(header.get(completed_col)) if completed_col else ""
    patched_lines: list[str] = []
    for line in kpi_lines:
        if line.strip().lower().startswith("• currently open?"):
            if completed:
                patched_lines.append(f"{line}    Completed Date: {completed}")
            else:
                patched_lines.append(line)
        else:
            patched_lines.append(line)

    for line in patched_lines:
        for wline in _wrap_text(line, text_width, "Helvetica", 10):
            if y < 80:
                c.showPage()
                y = height - 60
                c.setFont("Helvetica", 10)
            c.drawString(margin_left, y, wline)
            y -= 12
    y -= 10

    rank_col = _find(dff, "RANK", "Rank", "rank", "_rank")
    if rank_col and rank_col in dff.columns:
        rank_s = dff[rank_col]
        rank_s_str = rank_s.astype(str).str.strip().str.lower()
        rank_s_num = pd.to_numeric(rank_s, errors="coerce")

        po_mask = (rank_s_str == "po") | (rank_s_num == 2)
        tx_mask = rank_s_str.str.startswith("tran") | (rank_s_num == 3)
    else:
        po_mask = pd.Series(True, index=dff.index)
        tx_mask = pd.Series(True, index=dff.index)

    pn_col_global = _find(dff, "P/N", "P N", "PN")

    # ---------- PO table ----------
    po_col       = _find(dff, "PO")
    step_col     = _find(dff, "step_progress", "STATUSstep_progress", "Step", "Status")
    qty_col      = _find(dff, "QUANTITY RECEIVED", "Quantity received", "Qty")
    vendor_col   = _find(dff, "Vendors", "Vendor")
    cost_col     = _find(dff, "Total Parts Cost", "ITEM COST", "TOTAL ITEM COST", "Total cost")
    po_done_col  = completed_col

    po_rows = dff[po_mask & (dff[po_col].notna() if po_col else True)] if po_col else dff[po_mask]

    if po_col and not po_rows.empty:
        po_col_widths = [50, 80, 55, 35, 190, 60, 62]

        if y < 80:
            c.showPage()
            y = height - 60
        c.setFont("Helvetica-Bold", 11)
        c.drawString(margin_left, y, "Purchase Orders")
        y -= 16

        po_data = [["PO", "P/N", "Step", "Qty", "Vendor", "Cost", "Completed Date"]]

        for _, row in po_rows.iterrows():
            po_val   = _wrap_cell(row.get(po_col),        po_col_widths[0])
            pn_val   = _wrap_cell(row.get(pn_col_global), po_col_widths[1]) if pn_col_global else ""
            step_val = _wrap_cell(row.get(step_col),      po_col_widths[2]) if step_col   else ""
            qty_val  = _safe_str(row.get(qty_col)) if qty_col else ""
            vend_val = _wrap_cell(row.get(vendor_col),    po_col_widths[4]) if vendor_col else ""

            if cost_col:
                cost_raw = row.get(cost_col)
                if pd.notna(cost_raw):
                    try:
                        cost_val = f"{float(cost_raw):,.2f}"
                    except Exception:
                        cost_val = _safe_str(cost_raw)
                else:
                    cost_val = ""
            else:
                cost_val = ""

            done_val = _format_date(row.get(po_done_col)) if po_done_col else ""

            po_data.append([
                po_val,
                pn_val,
                step_val,
                qty_val,
                vend_val,
                cost_val,
                done_val,
            ])

        po_table = Table(
            po_data,
            colWidths=po_col_widths,
        )
        po_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.green),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
            ("ALIGN", (3, 1), (3, -1), "RIGHT"),
            ("ALIGN", (5, 1), (5, -1), "RIGHT"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("LEADING", (0, 0), (-1, -1), 9),
        ]))

        tw, th = po_table.wrapOn(c, width, height)
        if y - th < 40:
            c.showPage()
            y = height - 60
        po_table.drawOn(c, margin_left, y - th)
        y = y - th - 12

    # ---------- Transactions table ----------
    pn_col        = pn_col_global
    part_loc_col  = _find(dff, "part_loc", "Part Location")
    part_area_col = _find(dff, "part_area", "Area")
    tx_qty_col    = qty_col

    tx_rows = dff[tx_mask & (dff[pn_col].notna() if pn_col else True)] if pn_col else dff[tx_mask]

    if pn_col and not tx_rows.empty:
        tx_col_widths = [210, 40, 130, 130]

        if y < 80:
            c.showPage()
            y = height - 60
        c.setFont("Helvetica-Bold", 11)
        c.drawString(margin_left, y, "Transactions")
        y -= 16

        tx_data = [["P/N", "Qty", "Part Loc", "Part Area"]]

        for _, row in tx_rows.iterrows():
            pn_val   = _wrap_cell(row.get(pn_col),       tx_col_widths[0])
            qty_val  = _safe_str(row.get(tx_qty_col)) if tx_qty_col else ""
            loc_val  = _wrap_cell(row.get(part_loc_col), tx_col_widths[2]) if part_loc_col  else ""
            area_val = _wrap_cell(row.get(part_area_col),tx_col_widths[3]) if part_area_col else ""

            tx_data.append([pn_val, qty_val, loc_val, area_val])

        tx_table = Table(
            tx_data,
            colWidths=tx_col_widths,
        )
        tx_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.green),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
            ("ALIGN", (1, 1), (1, -1), "RIGHT"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("LEADING", (0, 0), (-1, -1), 9),
        ]))

        tw, th = tx_table.wrapOn(c, width, height)
        if y - th < 40:
            c.showPage()
            y = height - 60
        tx_table.drawOn(c, margin_left, y - th)
        y = y - th - 12

    c.save()
    buf.seek(0)
    return buf.getvalue()




def _render_wo_forms_pdf_page():
    """
    🧾 Work Orders — Forms, Filters & PDF Export

    Data source:
      - WO-Trans_non_DB.parquet in PARQUET_DIR / parquet_db.

    Filters:
      • Date & Location are taken from the main sidebar globals:
          - start_date, end_date
          - selected_locations
      • This tab adds extra drill-down filters:
          - Assignee
          - Team Assigned
          - WO Status (with group logic)
    """

    global LOGO_PATH

    # ========================================
    # Locate parquet dir + files
    # ========================================
    parquet_dir_global = globals().get("PARQUET_DIR", None)
    if parquet_dir_global:
        parquet_dir = Path(parquet_dir_global)
    else:
        APP_DIR = Path(__file__).resolve().parent if "__file__" in globals() else Path.cwd()
        parquet_dir = APP_DIR / "parquet_db"

    PARQUET_PATH = parquet_dir / "WO-Trans_non_DB.parquet"

    if not PARQUET_PATH.is_file():
        st.error(
            "Could not find WO-Trans_non_DB.parquet.\n\n"
            f"Looked at: {PARQUET_PATH}"
        )
        return

    USERS_PATH = parquet_dir / "Users.parquet"
    if not USERS_PATH.is_file():
        USERS_PATH = None

    LOGO_PATH = parquet_dir / "greer_logo.png"

    st.markdown("## 🧾 Work Orders — Forms, Filters & PDF Export")

    # ========================================
    # Load data
    # ========================================
    df_raw, date_col = load_data(PARQUET_PATH)

    if df_raw is None or df_raw.empty:
        st.error("No data found in WO-Trans_non_DB.parquet.")
        return

    # ========================================
    # Apply Date + Location filters from main sidebar
    # (these are defined in the global filters block)
    # ========================================
    global_start = globals().get("start_date", None)
    global_end = globals().get("end_date", None)
    global_locations = globals().get("selected_locations", None)

    # ---- Location filter ----
    loc_col = _find(df_raw, "Location", "location")
    if loc_col and global_locations:
        df_raw = df_raw[df_raw[loc_col].isin(global_locations)]

    # ---- Date filter on Completed On (if we have both date range + date column) ----
    if date_col and isinstance(global_start, date) and isinstance(global_end, date):
        df_raw = df_raw[df_raw[date_col].dt.date.between(global_start, global_end)]

    if df_raw.empty:
        st.info("No work orders match the current **date/location** filters.")
        return

    # Quick sanity check so you can see it’s filtered
    st.caption(f"Rows after main Date/Location filters: {len(df_raw)}")

    # ========================================
    # Local drill-down filters (Assignee / Team / WO Status)
    # ========================================
    st.markdown("### 🔍 Additional Filters (on top of main sidebar filters)")

    col1, col2 = st.columns(2)

    wo_status_col = _find(df_raw, "wo_status")
    group_col = _find(df_raw, "Workorder_group", "workorder_group")

    # ---- Assignee filter ----
    with col1:
        assignee_col = _find(df_raw, "assignee")
        if assignee_col:
            if USERS_PATH is not None:
                users_df = load_users(USERS_PATH)
                user_col = _find(users_df, "User")
                if user_col:
                    assignee_vals = sorted(users_df[user_col].dropna().unique())
                else:
                    assignee_vals = sorted(df_raw[assignee_col].dropna().unique())
            else:
                assignee_vals = sorted(df_raw[assignee_col].dropna().unique())

            sel_assignee = st.multiselect("Assignee", assignee_vals)
            if sel_assignee:
                mask_assignee = df_raw[assignee_col].isin(sel_assignee)
            else:
                mask_assignee = pd.Series(True, index=df_raw.index)
        else:
            mask_assignee = pd.Series(True, index=df_raw.index)

    # ---- Team + WO Status filters ----
    with col2:
        team_col = _find(df_raw, "team_assigned", "team assigned")
        if team_col:
            team_vals = _extract_unique_teams(df_raw[team_col])
            sel_team = st.multiselect("Team Assigned", team_vals)
            if sel_team:
                pattern = "|".join(re.escape(t) for t in sel_team)
                mask_team = df_raw[team_col].fillna("").str.contains(
                    pattern, case=False, regex=True
                )
            else:
                mask_team = pd.Series(True, index=df_raw.index)
        else:
            mask_team = pd.Series(True, index=df_raw.index)

        if wo_status_col:
            wo_stat_vals = sorted(df_raw[wo_status_col].dropna().unique())
            sel_wo_status = st.multiselect(
                "WO Status (multi-select)",
                wo_stat_vals,
                default=wo_stat_vals,
            )

            if sel_wo_status:
                base_mask = df_raw[wo_status_col].isin(sel_wo_status)
                if group_col:
                    groups_match = (
                        df_raw.loc[base_mask, group_col].dropna().unique()
                    )
                    mask_group_rows = df_raw[group_col].isin(groups_match)
                    mask_orphan_rows = base_mask & df_raw[group_col].isna()
                    mask_wo_status = mask_group_rows | mask_orphan_rows
                else:
                    mask_wo_status = base_mask
            else:
                mask_wo_status = pd.Series(True, index=df_raw.index)
        else:
            mask_wo_status = pd.Series(True, index=df_raw.index)

    # Combine extra filters only (Date/Location already applied above)
    mask_all = mask_assignee & mask_team & mask_wo_status
    df_filt = df_raw[mask_all].copy()

    # ========================================
    # Filtered table + lazy XLSX export
    # ========================================
    st.markdown("---")
    st.markdown("### 📋 Filtered Work Orders Table")

    step_col_disp = _get_step_progress_col(df_filt)
    if step_col_disp:
        sort_by_step = st.checkbox("Sort filtered table by Step Progress", value=False)
        if sort_by_step:
            if "WORKORDER" in df_filt.columns:
                df_filt = df_filt.sort_values(by=[step_col_disp, "WORKORDER"])
            else:
                df_filt = df_filt.sort_values(by=[step_col_disp])

    total_rows = len(df_raw)
    total_unique_wo = df_raw["WORKORDER"].nunique() if "WORKORDER" in df_raw.columns else None
    curr_rows = len(df_filt)
    curr_unique_wo = df_filt["WORKORDER"].nunique() if "WORKORDER" in df_filt.columns else None

    if total_unique_wo is not None and curr_unique_wo is not None:
        st.caption(
            f"Showing **{curr_rows}** rows from **{total_rows}** after Date/Location filters. "
            f"Unique Work Orders in view: **{curr_unique_wo}** "
            f"(of **{total_unique_wo}** after main filters)."
        )
    else:
        st.caption(f"Showing **{curr_rows}** rows from **{total_rows}** after main filters.")

    # Optional: cap displayed rows for speed
    max_rows_display = 500
    df_display = df_filt.head(max_rows_display)
    if len(df_filt) > max_rows_display:
        st.caption(f"Showing first {max_rows_display} rows (of {len(df_filt)}).")

    st.dataframe(df_display, use_container_width=True, hide_index=True)

    # -------- Lazy XLSX generation (only when requested) --------
    st.markdown("#### 📥 Export Filtered Table")

    if st.button("Generate XLSX file from filtered table", key="btn_make_wo_xlsx"):
        st.session_state["wo_filtered_xlsx"] = _to_xlsx_bytes(df_filt)

    if "wo_filtered_xlsx" in st.session_state:
        st.download_button(
            label="⬇️ Download filtered table as XLSX",
            data=st.session_state["wo_filtered_xlsx"],
            file_name="WO_filtered.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key="dl_xlsx_filtered",
        )

    st.markdown("---")

    # ========================================
    # Work Order detail + PDF exports
    # ========================================
    if "WORKORDER" in df_filt.columns and df_filt["WORKORDER"].notna().any():
        st.markdown("### 🔎 Work Order Detail & PDF Export")
        col_sel, col_single, col_multi = st.columns([1.4, 1, 1])

        with col_sel:
            wo_list = sorted(df_filt["WORKORDER"].dropna().unique())
            selected_wo = st.selectbox("Select Work Order", wo_list)

        with col_single:
            if st.button("🧾 Export selected WO as PDF", use_container_width=True):
                try:
                    pdf_bytes = _make_single_wo_pdf(df_filt, selected_wo)
                    st.download_button(
                        "📥 Download single WO PDF",
                        data=pdf_bytes,
                        file_name=f"WO_{selected_wo}.pdf",
                        mime="application/pdf",
                        key="dl_single_wo_pdf",
                    )
                except Exception as e:
                    st.error(f"Error creating single WO PDF: {e}")

        with col_multi:
            if st.button("🧾 Export ALL filtered WOs as one PDF", use_container_width=True):
                try:
                    pdf_bytes_all = _make_multi_wo_pdf(df_filt)
                    st.download_button(
                        "📥 Download multi-WO PDF",
                        data=pdf_bytes_all,
                        file_name="WO_filtered_summary.pdf",
                        mime="application/pdf",
                        key="dl_multi_wo_pdf_bottom",
                    )
                except Exception as e:
                    st.error(f"Error creating multi-WO PDF: {e}")

        # ---- Detail view layout ----
        st.markdown("---")
        st.markdown(f"### 📄 Work Order {selected_wo} — Detail View")

        dff = df_filt[df_filt["WORKORDER"] == selected_wo].copy()
        if group_col and not dff[group_col].isna().all():
            grp_val = dff[group_col].iloc[0]
            dff = df_filt[df_filt[group_col] == grp_val].copy()

        step_col = _get_step_progress_col(df_filt)
        qty_col = _find(df_filt, "QUANTITY RECEIVED", "Quantity received")

        header_row = dff.iloc[0]

        colA, colB = st.columns(2)
        with colA:
            st.markdown(f"**Title:** {_safe_str(header_row.get('TITLE'))}")
            st.markdown(f"**Description:** {_safe_str(header_row.get('Description'))}")
            if step_col:
                st.markdown(f"**Step Progress:** {_safe_str(header_row.get(step_col))}")
        with colB:
            st.markdown(f"**Assignee:** {_safe_str(header_row.get('assignee'))}")
            st.markdown(f"**Team Assigned:** {_safe_str(header_row.get('team_assigned'))}")
            st.markdown(f"**Asset:** {_safe_str(header_row.get('ASSET'))}")

        kpi_lines = _summarize_kpis(dff)
        if kpi_lines:
            st.markdown("##### 📊 KPI Summary")
            for line in kpi_lines:
                st.markdown(line)

        if "PO" in dff.columns:
            po_rows = dff[dff["PO"].notna()].copy()
        else:
            po_rows = pd.DataFrame()

        if not po_rows.empty:
            st.markdown("##### 📦 PO Lines")
            cols_po = ["PO"]
            if step_col:
                cols_po.append(step_col)
            cols_po.append("P/N")
            if qty_col and qty_col in po_rows.columns:
                po_rows = po_rows.rename(columns={qty_col: "Qty"})
                cols_po.append("Qty")
            cols_po.append("Vendors")
            cols_po_final = [c for c in cols_po if c in po_rows.columns]
            st.dataframe(po_rows[cols_po_final], use_container_width=True, hide_index=True)

        tx_mask = (
            (dff.get("part_loc").notna() if "part_loc" in dff.columns else False)
            | (dff.get("part_area").notna() if "part_area" in dff.columns else False)
        )
        tx_rows = dff[tx_mask].copy()
        if not tx_rows.empty:
            st.markdown("##### ↕️ Transactions")
            cols_tx = ["P/N"]
            if qty_col and qty_col in tx_rows.columns:
                tx_rows = tx_rows.rename(columns={qty_col: "Qty"})
                cols_tx.append("Qty")
            cols_tx.extend(["part_loc", "part_area"])
            cols_tx_final = [c for c in cols_tx if c in tx_rows.columns]
            st.dataframe(tx_rows[cols_tx_final], use_container_width=True, hide_index=True)

    else:
        st.info(
            "No work orders available in the filtered data to show detail or export PDFs. "
            "Check your main sidebar filters (date/location) and the additional filters above."
        )
        
        
def _render_hours_worksheet_page():
    """
    Hours Reading Entry — simple type-ahead dropdown

    • Asset chosen from a selectbox (search-as-you-type) from Ass_DB.parquet
    • When selected, we look up the Asset ID
    • Reading is optional (entered as text, parsed on Save)
    • Staging table columns: Asset, ID, Reading
    • Download CURRENT entries + BLANK worksheet
    """
    from pathlib import Path
    import pandas as pd
    import numpy as np
    import streamlit as st

    st.markdown("### ⏱ Hours Reading Entry")

    # -----------------------------
    # Locate Ass_DB.parquet
    # -----------------------------
    parq_dir = Path(
        globals().get("PARQUET_DIR", st.session_state.get("parquet_dir", ""))
        or (Path.cwd() / "parquet_db")
    )
    ass_path = parq_dir / "Ass_DB.parquet"
    if not ass_path.is_file():
        st.error(
            "Could not find **Ass_DB.parquet**.\n\n"
            f"Expected at: {ass_path}"
        )
        return

    # -----------------------------
    # Cached loader for assets
    # -----------------------------
    @st.cache_data(show_spinner=False)
    def _hours_load_assets(path_str: str):
        path = Path(path_str)
        try:
            try:
                df_ass = load_parquet(path)  # type: ignore[name-defined]
            except Exception:
                df_ass = pd.read_parquet(path)
        except Exception:
            return pd.DataFrame(), None, None, []

        if df_ass is None or df_ass.empty:
            return pd.DataFrame(), None, None, []

        # "Name" is what the user types / sees
        name_col = _find(df_ass, "Name", "Asset Name", "Asset")  # type: ignore[name-defined]
        id_col   = _find(df_ass, "ID", "Asset ID", "Asset_ID", "AssetID")  # type: ignore[name-defined]

        if name_col is None or id_col is None:
            return pd.DataFrame(), None, None, []

        df_assets = df_ass[[name_col, id_col]].copy()
        df_assets = df_assets.dropna(subset=[name_col]).drop_duplicates(subset=[name_col])

        options = sorted(df_assets[name_col].astype(str).tolist())
        return df_assets, name_col, id_col, options

    df_assets, name_col, id_col, asset_options = _hours_load_assets(str(ass_path))

    if df_assets.empty or name_col is None or id_col is None or not asset_options:
        st.error("No usable assets found in Ass_DB.parquet (need Name + ID columns).")
        return

    # -----------------------------
    # Session state
    # -----------------------------
    ss = st.session_state
    ss.setdefault("hours_entries", pd.DataFrame(columns=["Asset", "ID", "Reading"]))
    ss.setdefault("hours_selected_asset", "")
    ss.setdefault("hours_inputs", {"reading": ""})
    ss.setdefault("hours_focus_asset", False)  # for refocus after save
    df_entries = ss["hours_entries"]

    # -----------------------------
    # JS focus helper (like pre-op)
    # -----------------------------
    def _focus_asset_select_now():
        if ss.get("hours_focus_asset", False):
            try:
                import streamlit.components.v1 as components
                components.html(
                    """
                    <script>
                    const t=setInterval(()=>{
                      const inputs=parent.document.querySelectorAll('input');
                      for(const el of inputs){
                        const aria=el.getAttribute('aria-label')||'';
                        if(aria.includes('Asset')){ 
                          el.focus(); 
                          el.select(); 
                          clearInterval(t); 
                          break;
                        }
                      }
                    },50);
                    </script>
                    """,
                    height=0,
                )
            except Exception:
                pass
            ss["hours_focus_asset"] = False

    st.caption(
        "Start typing the asset in the dropdown below — it will auto-complete. "
        "Once you pick one, we pull the Asset ID from Ass_DB. Reading is optional."
    )

    # -----------------------------
    # Asset dropdown (type-ahead)
    # -----------------------------
    all_opts = [""] + asset_options
    current_sel = ss.get("hours_selected_asset", "")
    try:
        default_idx = all_opts.index(current_sel) if current_sel in all_opts else 0
    except ValueError:
        default_idx = 0

    selected_asset = st.selectbox(
        "Asset",
        options=all_opts,
        index=default_idx,
        help="Click and start typing to search by asset name/number.",
    )
    ss["hours_selected_asset"] = selected_asset

    # Try to focus the Asset control when needed
    _focus_asset_select_now()

    # Look up ID for the selected asset
    if selected_asset:
        row = df_assets[df_assets[name_col].astype(str) == str(selected_asset)].iloc[0]
        asset_id_val = str(row[id_col])
    else:
        asset_id_val = ""

    # Show the ID so they can see what will go into the table
    st.markdown(
        f"""
        <div style="padding:8px;border:1px solid #333;border-radius:8px;background:#111;max-width:320px;">
          <div style="opacity:.7;font-size:.8rem;">Asset ID</div>
          <div style="font-weight:800;font-size:1.1rem;">{asset_id_val}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    # -----------------------------
    # Reading (optional, text input)
    # -----------------------------
    reading_str = st.text_input(
        "Reading (optional)",
        value=ss["hours_inputs"].get("reading", ""),
    )
    ss["hours_inputs"]["reading"] = reading_str

    # -----------------------------
    # Single Save button
    # -----------------------------
    def _append_row_and_reset():
        if not selected_asset:
            st.warning("Please select an Asset from the dropdown.")
            return

        # Parse reading only on save
        reading_val = None
        rs = reading_str.strip()
        if rs:
            try:
                reading_val = float(rs)
            except ValueError:
                st.warning("Reading must be numeric (or leave it blank).")
                return

        row = {
            "Asset": selected_asset,
            "ID": asset_id_val,
            "Reading": reading_val,
        }

        ss["hours_entries"] = pd.concat(
            [ss["hours_entries"], pd.DataFrame([row])],
            ignore_index=True,
        )
        st.success("Reading row added.")

        # Reset to allow immediate new entry
        ss["hours_selected_asset"] = ""
        ss["hours_inputs"]["reading"] = ""
        ss["hours_focus_asset"] = True  # ask JS to focus Asset next run

        # Rerun to clear UI & focus asset
        if hasattr(st, "rerun"):
            st.rerun()
        elif hasattr(st, "experimental_rerun"):
            st.experimental_rerun()

    if st.button("Save", use_container_width=True):
        _append_row_and_reset()

    # -----------------------------
    # Current entries table
    # -----------------------------
    st.markdown("#### Current entries")
    df_entries = ss["hours_entries"]

    if df_entries.empty:
        st.info("No readings added yet.")
    else:
        with st.expander("Show current entries", expanded=False):
            st.dataframe(df_entries, hide_index=True, use_container_width=True)
            if st.button("🧹 Clear CURRENT table", use_container_width=True, key="hours_clear_table"):
                ss["hours_entries"] = pd.DataFrame(columns=["Asset", "ID", "Reading"])
                if hasattr(st, "rerun"):
                    st.rerun()
                elif hasattr(st, "experimental_rerun"):
                    st.experimental_rerun()

    # -----------------------------
    # Blank sheet rows (for printing/manual)
    # -----------------------------
    st.markdown("---")
    num_blank_rows = st.number_input(
        "Blank sheet rows (for manual fill)",
        min_value=5,
        max_value=200,
        value=30,
        step=5,
    )

    df_blank = pd.DataFrame(
        {
            "Asset":   ["" for _ in range(num_blank_rows)],
            "ID":      ["" for _ in range(num_blank_rows)],
            "Reading": ["" for _ in range(num_blank_rows)],
        }
    )

    # -----------------------------
    # Helper: DataFrame → XLSX bytes (global-safe)
    # -----------------------------
    global to_xlsx_bytes
    try:
        _ = to_xlsx_bytes  # type: ignore[name-defined]
    except NameError:
        from io import BytesIO

        def to_xlsx_bytes(df, sheet_name="Sheet1"):
            buf = BytesIO()
            with pd.ExcelWriter(buf, engine="xlsxwriter") as xw:
                df.to_excel(xw, index=False, sheet_name=sheet_name)
            buf.seek(0)
            return buf.getvalue()

        globals()["to_xlsx_bytes"] = to_xlsx_bytes

    # -----------------------------
    # Download buttons
    # -----------------------------
    st.markdown("#### 📥 Downloads")
    col_filled, col_blank = st.columns(2)

    with col_filled:
        st.download_button(
            "⬇️ Download CURRENT entries (XLSX)",
            data=to_xlsx_bytes(df_entries, sheet_name="Hours_Readings"),
            file_name="Hours_Readings_Current.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            disabled=df_entries.empty,
            use_container_width=True,
        )

    with col_blank:
        st.download_button(
            "⬇️ Download BLANK reading sheet (XLSX)",
            data=to_xlsx_bytes(df_blank, sheet_name="Hours_Readings_Blank"),
            file_name="Hours_Readings_Blank.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )

    st.caption(
        "Type to search for an asset, select it, optionally enter a reading, "
        "and build a staging list you can download as Excel."
    )




# =========================
# MAIN PAGE ROUTER (render ONE page at a time)
# =========================

PAGE_LABELS = [
    "💵 Costs & Trends",
    "📈 Inventory Analysis",
    "↕ Transactions",
    "🧾 Work Orders",
    "🛠 Expected Service",
    "🧾 WO Forms & PDFs",
    "⏱ Hours Worksheet",
    "📄 PDF Report",
]

current_page = st.sidebar.radio(
    "📑 Reporting Hub Page",
    PAGE_LABELS,
    index=0,
)

# --- Costs & Trends ---
if current_page == "💵 Costs & Trends":
    _render_costs_trends(
        df_all=dfs.get("costs_trends", pd.DataFrame()),
        global_start=start_date,
        global_end=end_date,
        selected_locations=selected_locations,
    )

# --- Inventory Analysis ---
elif current_page == "📈 Inventory Analysis":
    _render_inventory_analysis(
        df_parts=dfs.get("parts", pd.DataFrame()),
        df_tx=dfs.get("transactions", pd.DataFrame()),
        start_date=start_date,
        end_date=end_date,
        locations=selected_locations,
    )


# --- Transactions ---
elif current_page == "↕ Transactions":
    render_transactions(
        df_tx=dfs.get("transactions", pd.DataFrame()),
        start_date=start_date,
        end_date=end_date,
        selected_locations=selected_locations,
    )

# --- Work Orders Report ---
elif current_page == "🧾 Work Orders":
    render_wo_report(
        df_wo=dfs.get("workorders", pd.DataFrame()),
        start_date=start_date,
        end_date=end_date,
        selected_locations=selected_locations,
    )

# --- Expected Service ---
elif current_page == "🛠 Expected Service":
    render_expected_service(
        df_expected=dfs.get("expected", pd.DataFrame()),
        selected_locations=selected_locations,
    )

# --- WO Forms & PDFs (your new tab) ---
elif current_page == "🧾 WO Forms & PDFs":
    _render_wo_forms_pdf_page()

# --- Hours Worksheet (new page) ---
elif current_page == "⏱ Hours Worksheet":
    _render_hours_worksheet_page()

# --- PDF Report ---
elif current_page == "📄 PDF Report":
    st.markdown("### 📄 PDF Report")

    def _filter_df_to_window_for_pdf(df: pd.DataFrame,
                                     start: date,
                                     end: date,
                                     date_cols: list[str]) -> pd.DataFrame:
        """Apply date range to the first existing date_col in df."""
        if df is None or df.empty or not date_cols:
            return df
        date_col = None
        for c in date_cols:
            if c in df.columns:
                date_col = c
                break
        if date_col is None:
            return df
        s = pd.to_datetime(df[date_col], errors="coerce")
        mask = s.dt.date.between(start, end)
        return df.loc[mask].copy()

    def _build_ytd_by_location_from_workorders(df_in: pd.DataFrame, year_i: int) -> pd.DataFrame:
        import calendar
        import pandas as pd

        def _first_present_local(df: pd.DataFrame, candidates):
            for c in candidates:
                if c in df.columns:
                    return c
            return None

        if df_in is None or df_in.empty:
            return pd.DataFrame()

        dfw = df_in.copy()

        # detect columns (same candidates as the Costs & Trends tab)
        date_col = _first_present_local(dfw, ["Completed on","COMPLETED ON","Completed On","Date","Service date","Created on","Created On","Due date","Started on"])
        loc_col  = _first_present_local(dfw, ["Location","NS Location","location","Location2"])
        if not date_col or not loc_col:
            return pd.DataFrame()

        dfw[date_col] = pd.to_datetime(dfw[date_col], errors="coerce")

        def _num(s):
            return pd.to_numeric(s, errors="coerce").fillna(0.0)

        item_col = _first_present_local(dfw, ["Total Item Cost","TOTAL ITEM COST","total item cost","item total","item_total","TotalItemCost"])
        totl_col = _first_present_local(dfw, ["Total cost","Total Cost","TOTAL COST","total cost","cost total","cost_total","TotalCost"])

        if item_col and totl_col:
            dfw["__cost_sum"] = (_num(dfw[item_col]) + _num(dfw[totl_col])).astype(float)
            cost_col = "__cost_sum"
        elif "_Cost" in dfw.columns:
            cost_col = "_Cost"
        elif "__cost" in dfw.columns:
            cost_col = "__cost"
        else:
            dfw["__cost"] = (_num(dfw[item_col]) if item_col else 0.0) + (_num(dfw[totl_col]) if totl_col else 0.0)
            cost_col = "__cost"

        dfw[cost_col] = pd.to_numeric(dfw[cost_col], errors="coerce").fillna(0.0)

        # year scope
        dfw = dfw[dfw[date_col].dt.year == year_i].copy()
        if dfw.empty:
            return pd.DataFrame()

        dfw["__Month"] = dfw[date_col].dt.month

        g = dfw.groupby([loc_col, "__Month"], as_index=False)[cost_col].sum()
        g["Mon"] = g["__Month"].map(lambda m: calendar.month_abbr[int(m)] if pd.notna(m) else "")
        pv = g.pivot(index=loc_col, columns="Mon", values=cost_col).reset_index()

        mon_desc = [calendar.month_abbr[m] for m in range(12, 0, -1)]
        month_cols = [c for c in mon_desc if c in pv.columns]
        for c in month_cols:
            pv[c] = pd.to_numeric(pv[c], errors="coerce").fillna(0.0)

        pv["YTD Total"] = pv[month_cols].sum(axis=1, skipna=True)
        ordered_cols = [loc_col, "YTD Total"] + month_cols
        pv = pv[ordered_cols].sort_values("YTD Total", ascending=False)

        return pv


    # ---------- base frames from dfs ----------
    # Workorders.parquet lives in dfs["costs_trends"] in your app
    df_wo_pdf = dfs.get("costs_trends", pd.DataFrame())
    df_parts_pdf = dfs.get("parts", pd.DataFrame())
    df_tx_pdf    = dfs.get("transactions", pd.DataFrame())
    df_expected_pdf = dfs.get("expected", pd.DataFrame())

    # ---------- apply *location* filter ----------
    df_wo_pdf       = _filter_by_locations(df_wo_pdf, selected_locations)
    df_parts_pdf    = _filter_by_locations(df_parts_pdf, selected_locations)
    df_tx_pdf       = _filter_by_locations(df_tx_pdf, selected_locations)
    df_expected_pdf = _filter_by_locations(df_expected_pdf, selected_locations)

    # ---------- build Costs table for PDF (YTD by location) ----------
    year_i = int(end_date.year) if end_date else int(date.today().year)
    df_costs_pdf = _build_ytd_by_location_from_workorders(df_wo_pdf, year_i)

    # cache (optional)
    try:
        st.session_state["rhub_costs_ytd_loc"] = df_costs_pdf.copy() if df_costs_pdf is not None else pd.DataFrame()
    except Exception:
        pass

    # ---------- apply *date range* (where it makes sense) ----------
    df_tx_pdf = _filter_df_to_window_for_pdf(
        df_tx_pdf,
        start_date,
        end_date,
        ["Date", "Trans Date", "Created On", "Completed On"],
    )

    df_expected_pdf = _filter_df_to_window_for_pdf(
        df_expected_pdf,
        start_date,
        end_date,
        ["DueDate", "Expected Date", "Date"],
    )

    filtered_dfs_pdf = {
        "workorders": df_wo_pdf,          # raw Workorders rows (Workorders.parquet)
        "costs_trends": df_costs_pdf,     # pivot table YTD by location
        "parts": df_parts_pdf,
        "transactions": df_tx_pdf,
        "expected": df_expected_pdf,
    }

    st.caption(f"PDF inputs — workorders rows: {len(df_wo_pdf):,} • parts: {len(df_parts_pdf):,} • tx: {len(df_tx_pdf):,} • expected: {len(df_expected_pdf):,}")
    st.caption("PDF uses the current Reporting Window + selected Locations.")

    if st.button("Generate PDF report", type="primary"):
        with st.spinner("Building PDF report..."):

            # --- SANITY: show which keys exist + row counts ---
            try:
                st.write("PDF input keys:", list(filtered_dfs_pdf.keys()))
                for k, v in filtered_dfs_pdf.items():
                    if hasattr(v, "shape"):
                        st.write(f"• {k}: {v.shape[0]:,} rows × {v.shape[1]:,} cols")
                    else:
                        st.write(f"• {k}: {type(v)}")
            except Exception:
                pass

            pdf_bytes = b""
            try:
                pdf_bytes = build_reporting_hub_pdf(
                    filtered_dfs=filtered_dfs_pdf,
                    date_range=(start_date, end_date),
                    locations=selected_locations,
                )
            except Exception as e:
                st.exception(e)
                pdf_bytes = b""

            # --- GUARANTEE bytes ---
            if pdf_bytes is None:
                pdf_bytes = b""

            # If the builder returned empty, show why and stop
            if len(pdf_bytes) < 50:  # a real PDF is usually much larger than this
                st.error("PDF generation returned no data (empty bytes). This usually means the PDF builder hit an early return or never called c.save().")
            else:
                st.download_button(
                    "⬇️ Download PDF",
                    data=pdf_bytes,
                    file_name=f"Reporting_Hub_Report_{date.today():%Y-%m-%d}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                )















